from __future__ import annotations

import json
import os
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "04_最终成品" / "01_毕业设计说明书"
INPUT = OUT_DIR / "张淑鑫_浙江平湖二级油库工艺设计_毕业设计初稿_目录补漏完善版_2026-07-29.docx"
OUTPUT = OUT_DIR / "张淑鑫_浙江平湖二级油库工艺设计_毕业设计初稿_规范与一致性修订版_2026-07-30.docx"
REPORT = Path(__file__).with_name("revision_audit_20260730.json")


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace paragraph content while preserving its paragraph properties/style."""
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.add_run(text)


def replace_text(paragraph: Paragraph, old: str, new: str) -> None:
    """Replace text robustly, keeping run formatting when the match is in one run."""
    if old not in paragraph.text:
        raise ValueError(f"Text not found: {old!r}")
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return
    # Fallback for a match split across runs.
    set_paragraph_text(paragraph, paragraph.text.replace(old, new))


def insert_after(reference: Paragraph, text: str, *, alignment=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    if reference._p.pPr is not None:
        new_p.append(deepcopy(reference._p.pPr))
    reference._p.addnext(new_p)
    paragraph = Paragraph(new_p, reference._parent)
    paragraph.add_run(text)
    if alignment is not None:
        paragraph.alignment = alignment
    return paragraph


def find_paragraph(document: Document, needle: str, *, startswith=False) -> Paragraph:
    matches = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if (text.startswith(needle) if startswith else needle in text):
            matches.append(paragraph)
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph for {needle!r}, found {len(matches)}")
    return matches[0]


def append_citations(paragraph: Paragraph, citations: str) -> None:
    if citations in paragraph.text:
        return
    run = paragraph.add_run(citations)
    run.font.superscript = True


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)


def set_repeat_table_header(row) -> None:
    """Repeat the first row when a table continues on the next page."""
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:tblHeader"))
    if existing is None:
        existing = OxmlElement("w:tblHeader")
        tr_pr.append(existing)
    existing.set(qn("w:val"), "true")


def add_update_fields(document: Document) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def main() -> None:
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)
    document = Document(INPUT)
    changes: list[str] = []

    # 1. Keep the abstract consistent with the calculated pump/system intersection.
    zh_abstract = find_paragraph(document, "泵与最不利管路曲线交点约为202")
    replace_text(zh_abstract, "泵与最不利管路曲线交点约为202 m³/h、36.3 m",
                 "泵与最不利管路曲线自然交点约为220 m³/h、37.3 m，设计流量按200 m³/h通过变频控制")
    en_abstract = find_paragraph(document, "intersect at approximately 202")
    replace_text(
        en_abstract,
        "intersect at approximately 202 m³/h and 36.3 m",
        "naturally intersect at approximately 220 m³/h and 37.3 m; the operating flow is controlled at 200 m³/h by variable-frequency regulation",
    )
    changes.append("中英文摘要泵工作点统一为自然交点220 m³/h、37.3 m，并区分200 m³/h控制流量")

    # 2. Update the list of governing standards with current mandatory general codes.
    standards = find_paragraph(document, "油库等级、库址、总平面、防火间距")
    replace_text(
        standards,
        "消防给水执行GB 50974—2014",
        "消防设计同时执行GB 55036—2022《消防设施通用规范》和GB 55037—2022《建筑防火通用规范》，消防给水执行GB 50974—2014",
    )
    changes.append("设计依据补充GB 55036—2022和GB 55037—2022")

    # 3. Replace the incomplete tank-top-space statement and insert a traceable high-level check.
    level_summary = find_paragraph(document, "两类储罐的计算液位均低于18 m罐壁高度")
    set_paragraph_text(
        level_summary,
        "名义容量对应液位仅用于储罐几何核对，最高运行液位还应按SH/T 3007—2014第4.1.8条计算设计储存高液位，并按第5.4.2、5.4.3条及GB 50074—2014第15.1.1～15.1.4条设置高液位报警、高高液位报警及进罐阀联锁。最大进罐流量取水运进库300 m³/h，按15 min继续进液量75 m³计算。",
    )
    p = insert_after(
        level_summary,
        "20 000 m³汽油内浮顶罐：罐截面积1 256.64 m²，15 min进液折算高度为75/1 256.64=0.060 m；浮盘底面设计最高位置取17.40 m，安全裕量取0.30 m，设计储存高液位为17.40-0.060-0.30=17.04 m。名义容量液位15.915 m低于17.04 m；高液位报警取15.92 m，高高液位联锁取16.40 m。",
    )
    p = insert_after(
        p,
        "15 000 m³汽油内浮顶罐：罐截面积907.92 m²，15 min进液折算高度为75/907.92=0.083 m；浮盘底面设计最高位置取17.40 m，安全裕量取0.30 m，设计储存高液位为17.40-0.083-0.30=17.02 m。名义容量液位16.521 m低于17.02 m；高液位报警取16.52 m，高高液位联锁取16.90 m。",
    )
    p = insert_after(
        p,
        "15 000 m³柴油固定顶罐：泡沫产生器下沿至罐壁顶端的布置高度取0.50 m，15 min进液折算高度为0.083 m，安全裕量取0.30 m，设计储存高液位为18.00-0.50-0.083-0.30=17.12 m。名义容量液位16.521 m低于17.12 m；高液位报警取16.52 m，高高液位联锁取16.90 m。三类储罐均有不少于15 min最大流量进液的联锁处置余量。",
    )
    changes.append("按SH/T 3007—2014补算三类储罐设计储存高液位及高/高高液位设定")

    # 4. Apply the same liquid-level conclusion to Appendix B.
    appendix_replacements = [
        (
            "年周转质量G=400,000",
            "单罐D=40 m、H=18 m，名义液位15.915 m，满足罐壁顶空间要求。",
            "单罐D=40 m、H=18 m，名义液位15.915 m，低于按SH/T 3007—2014计算的设计储存高液位17.04 m；高液位报警取15.92 m，高高液位联锁取16.40 m。",
        ),
        (
            "年周转质量G=250,000",
            "单罐D=34 m、H=18 m，名义液位16.521 m，满足罐壁顶空间要求。",
            "单罐D=34 m、H=18 m，名义液位16.521 m，低于按SH/T 3007—2014计算的设计储存高液位17.02 m；高液位报警取16.52 m，高高液位联锁取16.90 m。",
        ),
        (
            "年周转质量G=300,000",
            "单罐D=34 m、H=18 m，名义液位16.521 m，满足罐壁顶空间要求。",
            "单罐D=34 m、H=18 m，名义液位16.521 m，低于按SH/T 3007—2014计算的设计储存高液位17.12 m；高液位报警取16.52 m，高高液位联锁取16.90 m。",
        ),
    ]
    for needle, old, new in appendix_replacements:
        paragraph = find_paragraph(document, needle)
        replace_text(paragraph, old, new)

    # 5. Correct piping and hydraulic consistency issues.
    suction_text = find_paragraph(document, "吸入管采用DN250，正常流量200 m³/h时流速约1.13")
    replace_text(suction_text, "1.13 m/s", "1.07 m/s")
    combined_flow = find_paragraph(document, "两种汽油同时装车时总流量240")
    replace_text(combined_flow, "2.10 m/s", "2.02 m/s")
    fire_ring = find_paragraph(document, "按式20～式23进行流速和阻力计算")
    replace_text(fire_ring, "式20～式23", "式20～式24")
    changes.append("修正DN250吸入流速、汽油合流管流速及消防环网公式引用范围")

    # 6. Correct pump inlet valve size.
    valve_table = document.tables[15]
    for row in valve_table.rows[1:]:
        if "泵入口" in row.cells[0].text:
            set_cell_text(row.cells[2], "DN250")
    changes.append("表4-4泵入口阀由DN200修正为DN250")

    # 7. Make the 300 m³/h ship-pump interface power boundary explicit in Appendix C.
    appendix_power = {
        "步骤6：按式27、泵效率ηp=0.70计算轴功率P=26.94 kW。": (
            "步骤6：按式27、泵效率ηp=0.70计算轴功率P=26.94 kW。该功率为水运进库300 m³/h接口工况的船泵轴功率估算，由船泵提供；不作为库内P-101或P-201输油泵的选型负荷。库内输油泵按表4-3的200 m³/h、36 m、30 kW校核。"
        ),
        "步骤6：按式27、泵效率ηp=0.70计算轴功率P=31.76 kW。": (
            "步骤6：按式27、泵效率ηp=0.70计算轴功率P=31.76 kW。该功率为柴油水运进库300 m³/h接口工况的船泵轴功率估算，由船泵提供；不作为库内P-301输油泵的选型负荷。库内P-301按表4-3的200 m³/h、36 m、30 kW校核，正文工作点轴功率为24.09 kW。若改由岸上接卸泵承担300 m³/h进库工况，应另行选泵配套电机。"
        ),
    }
    for prefix, replacement in appendix_power.items():
        matches = [p for p in document.paragraphs if p.text.startswith(prefix)]
        expected = 2 if "26.94" in prefix else 1
        if len(matches) != expected:
            raise ValueError(f"Expected {expected} matches for {prefix!r}, found {len(matches)}")
        for paragraph in matches:
            set_paragraph_text(paragraph, replacement)
    changes.append("附录C明确31.76 kW为船泵接口工况，30 kW库内泵按24.09 kW轴功率校核")

    # 8. Replace the placeholder fire-spacing table with calculated normative distances.
    fire_intro = find_paragraph(document, "总平面防火间距按GB 50074—2014")
    set_paragraph_text(
        fire_intro,
        "总平面防火间距按GB 50074—2014第5.1.3、5.1.8、6.1.15和6.5.2条逐项控制。钢制内浮顶汽油罐及丙A类固定顶柴油罐同组罐间净距均按0.4D计算；同一地上储罐区内相邻罐组间按较大罐直径的0.8倍计算。固定顶柴油罐与汽油内浮顶罐分别设防火堤。规范最小值、总平面采用值和校核结果列于表5-2。",
    )
    append_citations(fire_intro, "[3]")

    fire_table = document.tables[18]
    rows = [
        ["校核对象", "规范依据与计算原则", "规范最小值", "本设计采用值及结论"],
        ["G92同组T101—T102", "GB 50074—2014表6.1.15：钢制内浮顶罐0.4D", "0.4×40=16.0 m", "18.0 m，满足"],
        ["G95同组T201—T202", "GB 50074—2014表6.1.15：钢制内浮顶罐0.4D", "0.4×34=13.6 m", "15.0 m，满足"],
        ["GD同组T301—T302", "GB 50074—2014表6.1.15：丙A类固定顶罐0.4D", "0.4×34=13.6 m", "15.0 m，满足"],
        ["G92与G95相邻罐组", "GB 50074—2014第5.1.8条：相邻罐组0.8D", "0.8×40=32.0 m", "36.0 m，满足"],
        ["G95与GD相邻罐组", "GB 50074—2014第5.1.8条：相邻罐组0.8D", "0.8×34=27.2 m", "30.0 m，满足"],
        ["罐壁至防火堤内堤脚", "GB 50074—2014第6.5.2条：不小于0.5H", "0.5×18=9.0 m", "G92不小于22 m，G95/GD不小于19.2 m，满足"],
    ]
    # Existing table has seven rows. Populate all rows deterministically.
    if len(fire_table.rows) != len(rows):
        raise ValueError(f"Unexpected Table 5-2 row count: {len(fire_table.rows)}")
    for row_obj, row_values in zip(fire_table.rows, rows):
        for cell, value in zip(row_obj.cells, row_values):
            set_cell_text(cell, value)
    set_repeat_table_header(fire_table.rows[0])

    fire_note = find_paragraph(document, "说明：表中具体数值必须在两张图纸形成实际坐标")
    set_paragraph_text(
        fire_note,
        "表5-2所列采用值作为总平面布置尺寸。二级石油库罐区和装卸区消防车道按GB 50074—2014第5.2.8、5.2.9条采用6.0 m车道宽度，路面宽度6.0 m，净空高度不小于5.0 m，内缘转弯半径不小于12 m；储罐区形成环形消防车道。总平面图应逐项标注上述净距、道路宽度和转弯半径。",
    )
    road_text = find_paragraph(document, "罐区四周设置环形消防道路")
    set_paragraph_text(
        road_text,
        "罐区四周设置环形消防道路，车道宽度和路面宽度均取6.0 m，净空高度不小于5.0 m，内缘转弯半径不小于12 m。码头管廊与库区管廊沿道路一侧敷设，穿越道路处采用管涵或高架并保证消防净空。管理区与生产区分设出入口，装车车辆采用单向循环，避免与消防道路和人员通道交叉。",
    )
    changes.append("表5-2改为按GB 50074计算的罐间距、罐组间距、防火堤距离和道路控制值")

    # 9. Correct foam-system table values to match the calculation immediately above it.
    foam_table = document.tables[22]
    target_rows = [r for r in foam_table.rows if "15 000 m³汽油罐密封圈" in r.cells[0].text]
    if len(target_rows) != 1:
        raise ValueError("Could not uniquely locate the 15 000 m³ gasoline foam row")
    set_cell_text(target_rows[0].cells[1], "58.75")
    set_cell_text(target_rows[0].cells[3], "12.24/40.00")
    changes.append("表6-1中15 000 m³汽油罐保护面积和理论泡沫流量修正为58.75 m²、12.24 L/s")

    # 10. Put high-high setpoints into the alarm/interlock table.
    alarm_table = document.tables[24]
    for row in alarm_table.rows[1:]:
        if row.cells[0].text.strip() == "储罐高高液位":
            set_cell_text(row.cells[0], "储罐高高液位（20 000 m³汽油罐16.40 m；15 000 m³汽油/柴油罐16.90 m）")
    changes.append("表7-1补入高高液位联锁设定值")

    # 11. Reuse every recent reference in the substantive design section that it supports.
    citation_targets = [
        ("依据GB 50074—2014第6.1.4条", "[9]"),
        ("进库方式和出库方式保持计划表比例不变", "[2]"),
        ("总平面按约420 m×320 m", "[15]"),
        ("主要危险包括汽油蒸气形成爆炸性混合物", "[16]"),
        ("作业许可覆盖动火、受限空间", "[6]"),
        ("储罐基础、罐壁、锚固和接管按场地抗震参数设计", "[10]"),
        ("储罐底板采用可检测的防渗构造", "[11]"),
        ("按GB 50074—2014第12.2.7条", "[7]"),
        ("运行人员每班检查储罐液位趋势", "[8]"),
    ]
    for needle, citation in citation_targets:
        paragraph = find_paragraph(document, needle)
        append_citations(paragraph, citation)
    changes.append("近五年文献[2][3][6]—[11][15][16]补充到物流、储罐、总图、HSE、消防和运行章节的实际应用处")

    add_update_fields(document)
    document.save(OUTPUT)

    # Structural audit of citation order/count. Standards are intentionally excluded from the bibliography.
    reopened = Document(OUTPUT)
    reference_start = next(i for i, p in enumerate(reopened.paragraphs) if p.text.strip() == "参考文献")
    body = "\n".join(p.text for p in reopened.paragraphs[:reference_start])
    references = [p.text.strip() for p in reopened.paragraphs[reference_start + 1:] if re.match(r"^\[\d+\]", p.text.strip())]
    citation_hits = {}
    first_positions = {}
    for n in range(1, len(references) + 1):
        marker = f"[{n}]"
        citation_hits[n] = body.count(marker)
        first_positions[n] = body.find(marker)

    years = []
    foreign = []
    for ref in references:
        m = re.search(r"\b(20\d{2})\b", ref)
        years.append(int(m.group(1)) if m else None)
        if re.match(r"^\[\d+\]\s+[A-Z][A-Z\s]+,", ref):
            foreign.append(ref)

    report = {
        "input": str(INPUT),
        "output": str(OUTPUT),
        "changes": changes,
        "paragraphs": len(reopened.paragraphs),
        "tables": len(reopened.tables),
        "references_total": len(references),
        "references_2022_2026": sum(1 for y in years if y is not None and 2022 <= y <= 2026),
        "foreign_references": len(foreign),
        "citation_hits": citation_hits,
        "uncited_reference_numbers": [n for n, count in citation_hits.items() if count == 0],
        "citation_first_appearance_in_numeric_order": all(
            first_positions[n] >= 0 and first_positions[n] < first_positions[n + 1]
            for n in range(1, len(references))
        ),
    }
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
