from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX_PATH = Path(
    r"D:\毕业论文\04_最终成品\01_毕业设计说明书"
    r"\张淑鑫_浙江平湖二级油库工艺设计_毕业设计初稿_计算与规范完善版_2026-07-28.docx"
)
OUTPUT_PATH = DOCX_PATH.with_name(
    "张淑鑫_浙江平湖二级油库工艺设计_毕业设计初稿_气象数据修正版_2026-07-28.docx"
)


ROWS = [
    ("温度", "年平均气温", "16.7 ℃", "16.7 ℃", "荣晟纸机改造环评PDF第165页（报告书第160页）", "油品物性、设备环境条件"),
    ("温度", "最热月月平均气温", "29.0 ℃", "29.0 ℃", "荣晟纸机改造环评PDF第165页（报告书第160页）", "夏季运行温度边界"),
    ("温度", "最冷月月平均气温", "4.3 ℃", "4.3 ℃", "荣晟纸机改造环评PDF第165页（报告书第160页）", "冬季运行温度边界"),
    ("温度", "绝对最高气温", "39.9 ℃", "39.9 ℃", "荣晟纸机改造环评PDF第165页（原称“极端最高气温”）", "材料、仪表耐温"),
    ("温度", "绝对最低气温", "-7.7 ℃", "-7.7 ℃", "荣晟纸机改造环评PDF第165页（原称“极端最低气温”）", "防冻与低温适用性"),
    (
        "温度",
        "最低平均温度",
        "未单独提供",
        "不单独采用",
        "两份资料均未定义该独立指标，不与“最冷月月平均气温”混用",
        "待明确统计定义后使用",
    ),
    ("降雨量", "全年平均降雨量", "1185 mm", "1185 mm", "嘉化能源氯碱技改环评PDF第84页表6.1-1（报告书第81页）", "场地排水系统"),
    ("降雨量", "日最大降雨量", "276.4 mm", "276.4 mm", "嘉化能源氯碱技改环评PDF第84页表6.1-1（报告书第81页）", "事故水池雨水分量"),
    ("降雨量", "平均降雨天数", "136 d", "136 d", "嘉化能源氯碱技改环评PDF第83页（报告书第80页）", "排水与码头作业组织"),
    ("主导风向和风速", "主导风向", "E风（东风）", "E风（东风）", "荣晟纸机改造环评PDF第165页（报告书第160页）", "总图和管理区方位"),
    ("主导风向和风速", "年平均风速", "2.9 m/s", "2.9 m/s", "荣晟纸机改造环评PDF第165页（报告书第160页）", "运行环境与通风"),
    ("主导风向和风速", "最大风速", "20.3 m/s", "20.3 m/s", "嘉化能源氯碱技改环评PDF第84页表6.1-1（累年实测最大风速）", "历史最大风速边界"),
    ("主导风向和风速", "年最大风速", "12.0 m/s", "12.0 m/s", "荣晟纸机改造环评PDF第165页（报告书第160页）", "常年风况补充参数"),
    ("主导风向和风速", "极大风速", "31.7 m/s", "31.7 m/s", "嘉化能源氯碱技改环评PDF第84页表6.1-1（11级极大风力）", "台风停工与加固管理"),
    ("风荷载", "乍浦50年重现期基本风压", "0.45 kN/m²", "0.45 kN/m²", "《浙江省基本风压资料》PDF第43页（资料第37页）", "储罐抗风稳定性"),
    ("湿度", "年平均相对湿度", "78%", "78%", "荣晟纸机改造环评PDF第165页（报告书第160页）", "沿海防腐与电气选型"),
    ("天气", "年平均雾日", "35 d", "35 d", "政府环评PDF第167页（报告书第163页）", "码头可用时间"),
    ("天气", "年平均雷暴日", "28 d", "28 d", "政府环评PDF第168页（报告书第164页）", "防雷和作业管理"),
    ("地震", "地震基本烈度", "Ⅵ度", "Ⅵ度", "政府环评PDF第177页（报告书第173页）", "储罐、管线抗震"),
]


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def main():
    document = Document(DOCX_PATH)

    # Remove a historical equation fragment accidentally embedded in the
    # cover-title paragraph while retaining the visible title text “设计”.
    cover = document.paragraphs[0]
    for math_node in list(cover._p.findall(".//" + qn("m:oMath"))):
        parent = math_node.getparent()
        parent.remove(math_node)

    # An empty Heading 1 immediately before “1 设计总论” inherits
    # pageBreakBefore and produces a completely blank page after the TOC.
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() != "1 设计总论" or index == 0:
            continue
        previous = document.paragraphs[index - 1]
        if not previous.text.strip() and previous.style.name == "Heading 1":
            previous._p.getparent().remove(previous._p)
        break

    paragraph_replacements = {
        "区域属亚热带季风气候。独山港相邻工程采用的平湖站1971—2014年统计值为：年平均气温16.3 ℃、极端最高气温39.1 ℃、极端最低气温-9.3 ℃、年平均降水量1269.7 mm、年平均相对湿度80%、年平均雷暴日28 d。杭州湾北岸为强潮海区，平均潮差约4.82 m，历史最高潮位5.69 m，夏秋季台风可能引起增水。上述条件要求场地竖向设计兼顾防洪排涝，储罐及管架按沿海风环境进行抗风校核，码头软管和装卸臂设置紧急脱离及快速切断。[1]":
        "区域属亚热带季风气候。平湖市气象站近20年统计值为：年平均气温16.7 ℃、最热月平均气温29.0 ℃、最冷月平均气温4.3 ℃、极端最高气温39.9 ℃、极端最低气温-7.7 ℃、年平均相对湿度78%。嘉兴港区平湖气象特征值表给出的多年年平均降雨量为1185 mm，年均降雨日数为136 d。杭州湾北岸为强潮海区，夏秋季台风可能引起增水。上述条件要求场地竖向设计兼顾防洪排涝，储罐及管架按沿海风环境进行抗风校核，码头软管和装卸臂设置紧急脱离及快速切断。[1]",
        "自然条件采用拟建场址相邻的嘉兴港独山港区B区21、22号多用途泊位工程环评长期统计资料。该报告由浙江省政府信息公开平台发布，气象统计站为平湖站（30°37′N、121°05′E），统计期为1971—2014年，能够代表独山港陆域的温度、降水、湿度、雾、雷暴和风环境。表1-1同时列出原始PDF页码，便于后续复核。":
        "自然条件优先采用浙江省政府信息公开平台发布的两份平湖及嘉兴港区环境影响报告书。温度、主导风向、年平均风速和年最大风速采用平湖市气象站近20年统计资料；平均降雨量、降雨日数、日最大降雨量及累年实测最大风速采用《平湖气象特征值表》。雾日、雷暴日和地震烈度继续采用独山港区相邻工程公开资料。表1-1逐项列出原始PDF页码，避免混淆“最大风速”和“极大风速”。",
        "表中最大一日降水量276.4 mm直接用于第6章事故水池校核；50年基本风压0.45 kN/m²用于第2章抗风稳定性参数；极端风速31.7 m/s用于台风停工和加固管理边界，二者含义不同，不能相互替代。环评原文另有一处OCR显示“342 m/s”，与同一报告风速表及物理常识矛盾，本设计不采用该错误识别值。":
        "表中日最大降雨量276.4 mm直接用于第6章事故水池校核；50年基本风压0.45 kN/m²用于第2章抗风稳定性参数。累年实测最大风速20.3 m/s作为“最大风速”填写；31.7 m/s是表6.1-1所列11级“极大风速”，仅作为台风停工与加固管理边界；另一份近20年资料所列年最大风速为12.0 m/s。三种指标名称和统计口径不同，不相互替代。",
    }
    for paragraph in document.paragraphs:
        replacement = paragraph_replacements.get(paragraph.text)
        if replacement is not None:
            paragraph.text = replacement

    target = None
    for table in document.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header[:3] == ["类别", "参数", "原始统计值"]:
            target = table
            break
    if target is None:
        raise RuntimeError("未找到自然条件表")

    header_tr = target.rows[0]._tr
    row_template = deepcopy(target.rows[1]._tr)
    for tr in list(target._tbl.tr_lst)[1:]:
        target._tbl.remove(tr)

    for values in ROWS:
        new_tr = deepcopy(row_template)
        target._tbl.append(new_tr)
        row = target.rows[-1]
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)

    for cell in target.rows[0].cells:
        original = cell.text.strip()
        set_cell_text(cell, original, bold=True)

    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if values[:2] == ["自然条件", "平均气温/降水"]:
                set_cell_text(row.cells[2], "16.7 ℃/1185 mm")

    # Keep the header repeated if Word splits the expanded table over pages.
    tr_pr = header_tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)

    document.save(OUTPUT_PATH)
    print(f"updated: {OUTPUT_PATH}")
    print(f"natural-condition rows: {len(ROWS)}")


if __name__ == "__main__":
    main()
