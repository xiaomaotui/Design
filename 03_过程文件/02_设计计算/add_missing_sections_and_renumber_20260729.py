from __future__ import annotations

import copy
import math
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree


XSL_PATH = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}


def find_starts(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(text):
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag == qn("w:pPr"):
            continue
        paragraph._p.remove(child)


SUBSCRIPT_RE = re.compile(r"_\(([^)]+)\)|_([A-Za-z0-9Σ]+)")


def add_text_with_subscripts(paragraph: Paragraph, text: str) -> None:
    pos = 0
    for match in SUBSCRIPT_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        run = paragraph.add_run(match.group(1) or match.group(2))
        run.font.subscript = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_text(paragraph: Paragraph, text: str) -> Paragraph:
    clear_paragraph(paragraph)
    add_text_with_subscripts(paragraph, text)
    return paragraph


def insert_after(paragraph: Paragraph, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    result.style = style or paragraph.style
    add_text_with_subscripts(result, text)
    return result


def insert_before(paragraph: Paragraph, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    result = Paragraph(new_p, paragraph._parent)
    result.style = style or paragraph.style
    add_text_with_subscripts(result, text)
    return result


def append_before(anchor: Paragraph, items: list[tuple[str, object]]) -> Paragraph:
    last = None
    for kind, payload in items:
        if kind == "h2":
            last = insert_before(anchor, str(payload), find_starts(anchor._parent, "3.2 码头通过能力").style)
        elif kind == "h3":
            last = insert_before(anchor, str(payload), find_starts(anchor._parent, "3.2.2 泊位通过能力").style)
        elif kind == "p":
            last = insert_before(anchor, str(payload), find_starts(anchor._parent, "以92号汽油公路发油").style)
        else:
            raise ValueError(kind)
    return last or anchor


def latex_to_omml(latex: str, transform: etree.XSLT):
    mathml = etree.fromstring(latex_to_mathml(latex).encode("utf-8"))
    return transform(mathml).getroot()


def add_run_with_font(p, text: str, east_asia: str = "宋体"):
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), east_asia)
    fonts.set(qn("w:hAnsi"), east_asia)
    fonts.set(qn("w:eastAsia"), east_asia)
    rpr.append(fonts)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)


def add_tab(p):
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Cambria Math")
    fonts.set(qn("w:hAnsi"), "Cambria Math")
    fonts.set(qn("w:eastAsia"), "Cambria Math")
    rpr.append(fonts)
    r.append(rpr)
    tab = OxmlElement("w:tab")
    r.append(tab)
    p.append(r)


def insert_formal_equation_after(
    paragraph: Paragraph,
    latex: str,
    number: int,
    template_p,
    transform: etree.XSLT,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    # Use exactly the existing equation paragraph properties.
    ppr = template_p.find(qn("w:pPr"))
    if ppr is not None:
        new_p.append(copy.deepcopy(ppr))
    add_tab(new_p)
    new_p.append(copy.deepcopy(latex_to_omml(latex, transform)))
    add_tab(new_p)
    add_run_with_font(new_p, "（式")
    add_run_with_font(new_p, str(number), "Cambria Math")
    add_run_with_font(new_p, "）")
    return Paragraph(new_p, paragraph._parent)


def insert_display_math_after(
    paragraph: Paragraph, latex: str, transform: etree.XSLT
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    result.style = paragraph.style
    result.alignment = WD_ALIGN_PARAGRAPH.CENTER
    result.paragraph_format.first_line_indent = 0
    result.paragraph_format.left_indent = 0
    result.paragraph_format.right_indent = 0
    new_p.append(copy.deepcopy(latex_to_omml(latex, transform)))
    return result


def insert_formal_before(
    anchor: Paragraph,
    latex: str,
    number: int,
    template_p,
    transform: etree.XSLT,
) -> Paragraph:
    dummy = insert_before(anchor, "", anchor.style)
    eq = insert_formal_equation_after(dummy, latex, number, template_p, transform)
    dummy._element.getparent().remove(dummy._element)
    return eq


def replace_formula_references(doc: Document) -> None:
    # Existing equations are shifted by equations inserted in their actual sections.
    mapping = {i: i for i in range(1, 17)}
    mapping.update({i: i + 3 for i in range(17, 27)})
    mapping[27] = 30
    mapping[28] = 32
    mapping.update({i: i + 4 for i in range(29, 39)})
    mapping[39] = 45

    # Replace references contained in one text node. Formula labels split across
    # three runs are handled separately below.
    for node in doc.element.xpath(".//w:t"):
        if not node.text:
            continue

        def repl(match):
            old = int(match.group(1))
            return "式" + str(mapping.get(old, old))

        node.text = re.sub(r"式(\d+)", repl, node.text)

    # Existing equation paragraphs place “（式”, the number, and “）” in
    # separate runs. Update the middle run without rebuilding the paragraph.
    for element in doc.element.xpath(".//w:p|.//w:tbl"):
        texts = element.xpath(".//w:t")
        for i, node in enumerate(texts[:-2]):
            if node.text == "（式" and (texts[i + 1].text or "").isdigit() and texts[i + 2].text == "）":
                old = int(texts[i + 1].text)
                texts[i + 1].text = str(mapping.get(old, old))


def add_heading_after(paragraph: Paragraph, text: str, heading_style) -> Paragraph:
    return insert_after(paragraph, text, heading_style)


def main(src: Path, dst: Path) -> None:
    doc = Document(src)
    transform = etree.XSLT(etree.parse(str(XSL_PATH)))
    h2_style = find_starts(doc, "3.2 码头通过能力").style
    h3_style = find_starts(doc, "3.2.2 泊位通过能力").style
    body_style = find_starts(doc, "以92号汽油公路发油").style
    eq_template = find_starts(doc, "（式17）")._p
    # Shift only the source document's existing equation numbers first. All new
    # equations and references inserted below are therefore written directly
    # with their final numbers and cannot be mistaken for old references.
    replace_formula_references(doc)

    # 2.2.3 储罐附件与通气能力要求
    anchor = find_starts(doc, "2.3 罐壁厚度与抗风圈校核")
    insert_before(anchor, "2.2.3 储罐附件与通气能力要求", h3_style)
    insert_before(
        anchor,
        "汽油内浮顶罐设置量油孔、液位计、高高液位独立报警、边缘密封、浮盘防旋转装置、导静电装置、罐壁通气孔及紧急排水设施；柴油固定顶罐设置呼吸阀、阻火器、量油孔、液位计和高高液位报警。附件布置应满足检修可达性，并避免盘梯、抗风圈和消防管线互相遮挡。",
        body_style,
    )
    insert_before(
        anchor,
        "呼吸阀和通气孔的设计通气量应同时考虑最大进油排气量、最大出油吸气量、昼夜温差引起的热呼吸以及火灾工况紧急通气。现阶段以最大水运收油流量300 m³/h和最大库内外输流量200 m³/h作为正常工况下限，设备订货时按制造商依据储罐直径、罐高、介质闪点和设计压力给出的选型计算书复核；未取得制造商参数前，不在正文中虚构口径和台数。",
        body_style,
    )

    # 3.2.3 码头装卸臂数量与能力
    anchor = find_starts(doc, "3.3 工艺流程方案")
    insert_before(anchor, "3.2.3 码头装卸臂数量与能力校核", h3_style)
    p = insert_before(
        anchor,
        "码头按同一时段作业一种油品组织，控制流量取最大水运收油流量Q_w=300 m³/h。每种油品设置一台额定流量q_a=300 m³/h的专用装卸臂，装卸臂数量按下式计算：",
        body_style,
    )
    eq = insert_formal_before(
        anchor,
        r"n_a=\left\lceil\frac{Q_w}{q_a}\right\rceil",
        17,
        eq_template,
        transform,
    )
    insert_before(anchor, "式中：", body_style)
    insert_before(anchor, "n_a——单种油品所需装卸臂数量，台；", body_style)
    insert_before(anchor, "Q_w——单泊位设计装卸流量，m³/h；", body_style)
    insert_before(anchor, "q_a——单台装卸臂额定流量，m³/h。", body_style)
    p = insert_before(anchor, "将Q_w=300 m³/h、q_a=300 m³/h代入式17：", body_style)
    p = insert_display_math_after(p, r"n_a=\left\lceil\frac{300}{300}\right\rceil=1", transform)
    insert_before(
        anchor,
        "计算得单种油品配置1台装卸臂。92号汽油、95号汽油和0号柴油分别设置专用装卸臂，共3台；作业切换通过阀组隔离，避免共用软管造成混油。单台能力等于控制流量，泊位年通过能力仍按式15校核。",
        body_style,
    )

    # 3.4 公路装车设施
    anchor = find_starts(doc, "4 管网水力计算与设备选型")
    insert_before(anchor, "3.4 公路装车设施计算", h2_style)
    insert_before(anchor, "3.4.1 汽车发油鹤位数量计算", h3_style)
    p = insert_before(
        anchor,
        "公路发油年作业天数取D=330 d/a，每日有效作业时间t_d=16 h/d，单鹤位流量q_i=60 m³/h。第i种油品所需鹤位数量按年公路发油量校核：",
        body_style,
    )
    insert_formal_before(
        anchor,
        r"n_i=\left\lceil\frac{G_{i,r}}{q_i\rho_i t_dD}\right\rceil",
        18,
        eq_template,
        transform,
    )
    for line in [
        "式中：",
        "n_i——第i种油品所需鹤位数量，个；",
        "G_(i,r)——第i种油品公路年发油质量，t/a；",
        "q_i——单鹤位装车流量，m³/h；",
        "ρ_i——第i种油品设计密度，t/m³；",
        "t_d——每日有效装车时间，h/d；",
        "D——年作业天数，d/a。",
        "92号汽油公路年发油量为260 000 t/a，95号汽油为175 000 t/a，0号柴油为180 000 t/a。分别代入式18：",
    ]:
        insert_before(anchor, line, body_style)
    p = insert_before(anchor, "", body_style)
    p = insert_display_math_after(
        p,
        r"n_{92}=\left\lceil\frac{260000}{60\times0.760\times16\times330}\right\rceil=2",
        transform,
    )
    p = insert_display_math_after(
        p,
        r"n_{95}=\left\lceil\frac{175000}{60\times0.760\times16\times330}\right\rceil=1",
        transform,
    )
    p = insert_display_math_after(
        p,
        r"n_{0}=\left\lceil\frac{180000}{60\times0.840\times16\times330}\right\rceil=1",
        transform,
    )
    insert_before(
        anchor,
        "计算最少需要4个鹤位。为使三种油品均能双鹤位装车并在单鹤位检修时维持作业，本设计每种油品设置2个鹤位，共6个鹤位。该配置满足能力校核，同时与现有120 m³/h单品种公路发油工况一致。",
        body_style,
    )
    insert_before(anchor, "3.4.2 装车台及装车棚面积", h3_style)
    p = insert_before(
        anchor,
        "六个鹤位采用3座双侧装车岛布置。初步总图按装车台长度L_s=30 m、宽度B_s=12 m控制，装车棚投影面积按下式计算：",
        body_style,
    )
    insert_formal_before(anchor, r"A_s=L_sB_s", 19, eq_template, transform)
    for line in [
        "式中：",
        "A_s——装车棚投影面积，m²；",
        "L_s——装车台控制长度，m；",
        "B_s——装车棚控制宽度，m。",
        "将L_s=30 m、B_s=12 m代入式19：",
    ]:
        insert_before(anchor, line, body_style)
    p = insert_before(anchor, "", body_style)
    p = insert_display_math_after(p, r"A_s=30\times12=360\ {\rm m^2}", transform)
    insert_before(
        anchor,
        "装车棚投影面积取360 m²。该尺寸用于初步总平面布置，能够覆盖3座双侧装车岛及鹤管作业区；车辆转弯、排队和消防通道另在装车区道路范围内布置，不计入装车棚面积。最终尺寸以总平面图中的车辆轨迹和设备厂家界面复核。",
        body_style,
    )
    insert_before(anchor, "3.4.3 装车支管与集油管校核", h3_style)
    insert_before(
        anchor,
        "单鹤位流量60 m³/h时采用DN100支管，计算流速2.21 m/s；同品种两鹤位同时作业流量120 m³/h时采用DN150集油支管，计算流速1.96 m/s。装车总管按不利同时作业组合校核：两种汽油同时装车时总流量240 m³/h，采用DN200总管，计算流速2.10 m/s；柴油单品种最大流量120 m³/h，采用DN150管道。上述管径计算均采用第4.1节式20，详细水力损失在第4.2节按作业工况计算。",
        body_style,
    )

    # 4.2 按工况组织水力计算，不移动原公式和计算内容。
    p = find_starts(doc, "4.2 沿程阻力")
    insert_after(p, "4.2.1 计算方法与参数", h3_style)
    p = find_starts(doc, "以92号汽油公路发油最远鹤位")
    insert_before(p, "4.2.2 公路装车管路水力计算", h3_style)
    p = find_starts(doc, "柴油水运进库采用DN250")
    insert_before(p, "4.2.3 水运收发油管路水力计算", h3_style)
    p = find_starts(doc, "表4-2 主要路径水力计算汇总")
    insert_before(p, "4.2.4 公共管道与倒罐工况", h3_style)
    insert_before(
        p,
        "公共管道收发油和库内倒罐均按“流量—实际内径—雷诺数—摩阻系数—沿程损失—局部损失—系统扬程”的顺序计算。公共管道外输采用200 m³/h，倒罐采用150 m³/h；各路径采用实际管长、标高和局部构件数量，计算结果汇总于表4-2，逐条分步过程见附录A。",
        body_style,
    )

    # 5.3 分别校核汽油、柴油罐组防火堤高度。
    p = find_starts(doc, "5.3 防火堤与泄漏围控")
    insert_after(p, "5.3.1 汽油罐组防火堤高度与容积", h3_style)
    p_area = find_starts(doc, "（式30）")
    insert_formal_equation_after(
        p_area,
        r"H_d=\frac{V_{\max}}{A_n}+0.20",
        31,
        eq_template,
        transform,
    )
    p = find_starts(doc, "An——防火堤内有效净面积")
    insert_before(p, "H_d——防火堤设计高度，m；", body_style)
    insert_before(p, "V_max——罐组内最大单罐公称容量，m³；", body_style)
    p = find_starts(doc, "以G92罐组为例")
    # The source already contains two unnumbered numerical equations for this
    # example. They are replaced by the expanded area-height-volume sequence
    # below, so remove only those old equation-only paragraphs.
    cursor = p._p.getnext()
    while cursor is not None:
        texts = "".join((x.text or "") for x in cursor.iter(qn("w:t"))).strip()
        if texts.startswith("计算有效容积"):
            break
        next_cursor = cursor.getnext()
        if cursor.tag == qn("w:p") and cursor.find(".//" + qn("m:oMath")) is not None:
            cursor.getparent().remove(cursor)
        cursor = next_cursor
    set_text(
        p,
        "以G92罐组为例，将L=140 m、B=90 m、n=2、D=40 m代入式30，先计算有效净面积；再将最大单罐容量V_max=20 000 m³代入式31确定所需堤高：",
    )
    p = insert_display_math_after(
        p,
        r"A_n=140\times90-\frac{2\pi\times40^2}{4}=10086.73\ {\rm m^2}",
        transform,
    )
    p = insert_display_math_after(
        p,
        r"H_d=\frac{20000}{10086.73}+0.20=2.183\ {\rm m}",
        transform,
    )
    p = insert_after(
        p,
        "防火堤设计高度向上取2.2 m。将A_n=10 086.73 m²、H_d=2.2 m代入式32校核有效容积：",
        body_style,
    )
    p = insert_display_math_after(
        p, r"V_d=10086.73\times(2.2-0.2)=20173.46\ {\rm m^3}", transform
    )
    anchor_table = find_starts(doc, "表5-3 防火堤有效容积校核")
    insert_before(anchor_table, "5.3.2 柴油罐组防火堤高度与容积", h3_style)
    p = insert_before(
        anchor_table,
        "柴油罐组内尺寸为120 m×80 m，布置2座直径34 m、容量15 000 m³的固定顶罐。按式30计算有效净面积，按式31确定所需堤高：",
        body_style,
    )
    p = insert_display_math_after(
        p,
        r"A_n=120\times80-\frac{2\pi\times34^2}{4}=7784.16\ {\rm m^2}",
        transform,
    )
    p = insert_display_math_after(
        p,
        r"H_d=\frac{15000}{7784.16}+0.20=2.127\ {\rm m}",
        transform,
    )
    p = insert_after(
        p,
        "柴油罐组防火堤设计高度同样向上取2.2 m。按式32校核有效容积：",
        body_style,
    )
    p = insert_display_math_after(
        p, r"V_d=7784.16\times(2.2-0.2)=15568.32\ {\rm m^3}", transform
    )
    insert_after(
        p,
        "计算有效容积15 568.32 m³＞最大单罐容积15 000 m³，满足围控要求。95号汽油罐组的罐径、数量和防火堤尺寸与柴油罐组相同，几何校核结果一致。",
        body_style,
    )
    p = find_starts(doc, "防火堤内地坪坡向集水井")
    insert_before(p, "5.3.3 排水与泄漏围控", h3_style)

    # 6.1.3 消防环网管径与水力校核
    anchor = find_starts(doc, "6.2 泡沫灭火系统")
    insert_before(anchor, "6.1.3 消防给水环网管径与水力校核", h3_style)
    p = insert_before(
        anchor,
        "冷却水控制流量取Q_c=130 L/s=0.130 m³/s。消防环网采用Φ323.9×8 mm钢管，计算内径d=0.308 m。按式20～式23进行流速和阻力计算，取环网不利路径长度L=600 m、清洁钢管绝对粗糙度ε=0.045 mm、局部阻力系数总和Σζ=15，20 ℃水运动黏度ν=1.0×10⁻⁶ m²/s：",
        body_style,
    )
    p = insert_display_math_after(
        p, r"v=\frac{4\times0.130}{\pi\times0.308^2}=1.745\ {\rm m/s}", transform
    )
    p = insert_display_math_after(
        p, r"Re=\frac{1.745\times0.308}{1.0\times10^{-6}}=5.37\times10^5", transform
    )
    p = insert_display_math_after(
        p,
        r"\lambda=\frac{0.25}{[\lg(\frac{0.000045}{3.7\times0.308}+\frac{5.74}{(5.37\times10^5)^{0.9}})]^2}=0.0149",
        transform,
    )
    p = insert_display_math_after(
        p,
        r"h_f=0.0149\times\frac{600}{0.308}\times\frac{1.745^2}{2\times9.81}=4.50\ {\rm m}",
        transform,
    )
    p = insert_display_math_after(
        p, r"h_m=15\times\frac{1.745^2}{2\times9.81}=2.33\ {\rm m}", transform
    )
    insert_after(
        p,
        "环网不利路径沿程与局部损失合计6.83 m。按单向供水控制时DN300环网流速1.745 m/s；环网两向供水时各方向流量降低，阻力更小，因此DN300满足本阶段冷却水输送要求。",
        body_style,
    )
    p = find_starts(doc, "消防泵扬程按式25")
    set_text(
        p,
        "消防泵扬程按式25的组成逐项校核。冷却水控制点取罐顶高差18 m、计算环网与支路损失6.83 m、喷头入口所需压力水头20 m，系统所需扬程为44.83 m；考虑支管、阀门布置调整和设备裕量，冷却水泵初选H=60 m。泡沫控制点取罐顶高差18 m、主管及支管损失20 m、比例混合装置损失10 m、PCL产生器额定进口压力0.50 MPa折合水头51 m，计算扬程99 m，泡沫水泵初选H=110 m。PCL额定进口压力取值来自产品样本，采购型号变化时应按实际额定压力复核，不把该产品参数写成国家标准定值。",
    )

    # 6.4 给排水分区和防洪排涝。
    anchor = find_starts(doc, "6.5 消防系统运行与可靠性")
    insert_before(anchor, "6.4.3 给排水系统与排水分区", h3_style)
    insert_before(
        anchor,
        "库区给排水采用清污分流、雨污分流和事故状态可切换的组织方式。生活给水由市政管网供给；消防给水由独立消防水池和消防泵组供给。设备冲洗水、罐底切水、泵棚和装车区含油地面排水进入含油污水系统，经隔油和处理后处置；生活污水单独收集，不与含油污水混排。",
        body_style,
    )
    insert_before(
        anchor,
        "罐区、防火堤和装车区的初期污染雨水经切换井进入初期雨水池，后期清净雨水经油膜检测和人工确认后进入清净雨水系统。事故状态关闭雨水总排口，消防废水和污染雨水转入事故水池。清净雨水沟、含油污水管和事故水管在总图上分别标识流向，交叉处不得采用无隔离的直连。",
        body_style,
    )
    insert_before(anchor, "6.4.4 防洪排水沟水力校核", h3_style)
    p = insert_before(
        anchor,
        "独山港区暴雨强度采用浙江省工程建设标准DB33/T 1191—2020中平湖市暴雨强度公式的工程应用结果。嘉兴港独山港区A区10号泊位码头工程环境影响报告书在PDF浏览器第113页（文档内部第107页）给出重现期P=2 a、降雨历时t=15 min时q=251.69 L/(s·hm²)。本设计总用地面积F=420 m×320 m=13.44 hm²，混合地面综合径流系数ψ取0.75作为初步设计假定，雨水设计流量按下式计算：",
        body_style,
    )
    insert_formal_before(anchor, r"Q_r=\psi Fq", 43, eq_template, transform)
    for line in [
        "式中：",
        "Q_r——雨水设计流量，L/s；",
        "ψ——综合径流系数；",
        "F——汇水面积，hm²；",
        "q——设计暴雨强度，L/(s·hm²)。",
        "将ψ=0.75、F=13.44 hm²、q=251.69 L/(s·hm²)代入式43：",
    ]:
        insert_before(anchor, line, body_style)
    p = insert_before(anchor, "", body_style)
    p = insert_display_math_after(
        p, r"Q_r=0.75\times13.44\times251.69=2537.0\ {\rm L/s}=2.537\ {\rm m^3/s}", transform
    )
    p = insert_after(
        p,
        "库区外围截洪沟初选矩形断面，底宽b=1.50 m、设计水深h=1.00 m、纵坡S=0.003，混凝土沟曼宁糙率n_M=0.015。排水沟过流能力按曼宁公式计算：",
        body_style,
    )
    insert_formal_before(
        anchor, r"Q_d=\frac{1}{n_M}AR^{2/3}S^{1/2}", 44, eq_template, transform
    )
    for line in [
        "式中：",
        "Q_d——排水沟设计过流能力，m³/s；",
        "n_M——曼宁糙率；",
        "A——过水断面面积，m²；",
        "R——水力半径，m；",
        "S——沟底纵坡。",
        "矩形断面的过水面积A=bh=1.50 m²，湿周χ=b+2h=3.50 m，水力半径R=A/χ=0.4286 m。代入式44：",
    ]:
        insert_before(anchor, line, body_style)
    p = insert_before(anchor, "", body_style)
    p = insert_display_math_after(
        p,
        r"Q_d=\frac{1}{0.015}\times1.50\times0.4286^{2/3}\times0.003^{1/2}=3.11\ {\rm m^3/s}",
        transform,
    )
    insert_before(
        anchor,
        "排水沟过流能力3.11 m³/s＞场地设计雨水流量2.537 m³/s，初选断面满足计算要求。该校核针对全场峰值清净雨水排放，不替代第6.4.2节按污染深度计算的初期雨水池容积；施工图阶段应按实际分区汇水面积、地面高程和出水口潮位分段复核。",
        body_style,
    )

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.core_properties.title = "浙江平湖二级油库工艺设计毕业设计初稿（目录补漏完善版）"
    doc.core_properties.subject = "按最新版计划表补充码头、公路装车、防火堤、消防环网、给排水与防洪计算"
    doc.save(dst)

    # Structural checks.
    check = Document(dst)
    visible = "\n".join(p.text for p in check.paragraphs)
    required = [
        "2.2.3 储罐附件与通气能力要求",
        "3.2.3 码头装卸臂数量与能力校核",
        "3.4.1 汽车发油鹤位数量计算",
        "5.3.2 柴油罐组防火堤高度与容积",
        "6.1.3 消防给水环网管径与水力校核",
        "6.4.4 防洪排水沟水力校核",
    ]
    missing = [x for x in required if x not in visible]
    if missing:
        raise RuntimeError(f"Missing inserted sections: {missing}")
    if "式45" not in visible:
        raise RuntimeError("Formula renumbering did not reach formula 45.")
    print(f"saved={dst}")
    print(f"paragraphs={len(check.paragraphs)} tables={len(check.tables)}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: script.py SOURCE.docx OUTPUT.docx")
    main(Path(sys.argv[1]), Path(sys.argv[2]))
