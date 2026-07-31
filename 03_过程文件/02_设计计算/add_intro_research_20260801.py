from __future__ import annotations

import hashlib
import re
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


ROOT = Path(r"D:\毕业论文")
SOURCE = ROOT / "04_最终成品" / "01_毕业设计说明书" / "张淑鑫_浙江平湖二级油库工艺设计_毕业设计初稿_防火堤与一致性修订版_2026-07-31.docx"
OUTPUT = ROOT / "04_最终成品" / "01_毕业设计说明书" / "张淑鑫_浙江平湖二级油库工艺设计_毕业设计初稿_绪论与研究动态补充版_2026-08-01.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NS = {"w": W_NS, "m": M_NS}


def paragraph_text(paragraph) -> str:
    return paragraph.text.strip()


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def clone_paragraph_properties(source, target) -> None:
    source_ppr = source._p.pPr
    if source_ppr is not None:
        if target._p.pPr is not None:
            target._p.remove(target._p.pPr)
        target._p.insert(0, deepcopy(source_ppr))


def set_page_break_before(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    page_break = ppr.find(qn("w:pageBreakBefore"))
    if page_break is None:
        page_break = OxmlElement("w:pageBreakBefore")
        ppr.append(page_break)


def add_cited_text(paragraph, text: str) -> None:
    parts = re.split(r"(\[\d+\])", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        if re.fullmatch(r"\[\d+\]", part):
            run.font.superscript = True


def insert_paragraph(target, text: str, template, page_break: bool = False):
    paragraph = target.insert_paragraph_before()
    clone_paragraph_properties(template, paragraph)
    if page_break:
        set_page_break_before(paragraph)
    add_cited_text(paragraph, text)
    return paragraph


def replace_in_runs(paragraph, transform) -> None:
    for run in paragraph.runs:
        if run.text:
            run.text = transform(run.text)


def all_body_paragraphs(document: Document):
    for element in document.element.body.iter():
        if element.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            yield Paragraph(element, document.element.body)


def formula_fingerprint(path: Path):
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
    root = etree.fromstring(xml)
    formulas = root.xpath(".//m:oMathPara | .//m:oMath[not(ancestor::m:oMathPara)]", namespaces=NS)
    fingerprints = []
    for formula in formulas:
        structure = []
        for node in formula.iter():
            local = etree.QName(node).localname
            text = node.text or ""
            attrs = sorted((etree.QName(key).localname, value) for key, value in node.attrib.items())
            structure.append((local, text, attrs))
        digest = hashlib.sha256(repr(structure).encode("utf-8")).hexdigest()
        fingerprints.append(digest)
    return fingerprints


def main() -> None:
    shutil.copy2(SOURCE, OUTPUT)
    before_formulas = formula_fingerprint(SOURCE)
    document = Document(OUTPUT)

    # Locate formatting exemplars and the original first design chapter.
    design_heading = next(p for p in document.paragraphs if paragraph_text(p) == "1 设计总论")
    heading1_template = design_heading
    heading2_template = next(p for p in document.paragraphs if paragraph_text(p) == "1.1 项目概况与设计范围")
    body_template = next(p for p in document.paragraphs if paragraph_text(p).startswith("拟建油库服务浙江北部"))

    # Remove the old literature-review block under design basis. Its substance is
    # rewritten in the new introduction, so no duplicate review remains.
    paragraphs = list(document.paragraphs)
    start = next(i for i, p in enumerate(paragraphs) if paragraph_text(p) == "1.2.3 近五年研究与方案借鉴")
    end = next(i for i, p in enumerate(paragraphs) if paragraph_text(p) == "1.2.4 主要设计数据")
    for paragraph in paragraphs[start:end]:
        remove_paragraph(paragraph)

    # Insert the school-template introduction and research-dynamics structure.
    blocks = [
        ("h1", "绪论"),
        (
            "body",
            "成品油库是区域能源供应体系中连接炼化生产、港口运输、公共管道和终端配送的重要节点，其工艺设计需要在既定年周转量和运输比例下，统筹储罐容量、收发油流程、管网输送能力、装卸设备、总平面布置、消防安全与环境保护。浙江平湖独山港区兼具深水岸线、临港产业和陆路集疏运条件，适合建设承担水运接卸、管道转输和公路配送任务的区域性成品油库。",
        ),
        (
            "body",
            "本设计以92号汽油、95号汽油和0号柴油为储存介质，按照二级石油库控制建设规模，在不改变既定进出库方式及比例的条件下，完成库容与储罐组合、码头泊位、收发油工艺、管网水力、泵组能力、总平面、防火堤、泡沫与冷却水系统、事故水收集以及HSE措施设计。国内外研究成果用于识别技术发展方向和选择工程方法，具体设计限值仍以中国现行工程建设标准为依据。",
        ),
        ("h1", "国内外研究动态"),
        ("h2", "国内研究动态"),
        (
            "body",
            "国内油库研究已由单项设备计算逐步转向物流组织、工艺系统与安全环保协同优化。许少新等对成品油管铁联运物流进行优化，指出运输方式、接卸能力、库存水平和终端需求之间存在耦合关系[1]。这一认识同样适用于水运、公共管道和公路并存的商业油库：库容不能脱离年周转量单独扩大，泵和管径也应按实际作业窗口及峰值工况确定。",
        ),
        (
            "body",
            "罐区事故防控研究重点由单一防火间距转向多层安全屏障。吴守志等量化分析了安全屏障对池火灾多米诺效应概率的影响，表明可靠的探测、隔离、固定消防和应急处置可以共同降低事故升级概率[2]。在污染控制方面，王雷等分析了浮顶密封、罐壁状态和附件对储罐VOCs排放量的影响[3]；刘世达等进一步总结了源头减量、密闭收集、回收处理与运行管控相结合的深度减排技术路径[4]。因此，汽油储罐型式、浮盘密封、油气回收和泄漏检测应作为同一工艺链统筹设计。",
        ),
        (
            "body",
            "数字化和全寿命周期管理是国内研究的另一重点。朱喜平构建大型石油储库全寿命周期风险管控一体化平台，将设备状态、作业风险和应急资源纳入统一管理[5]；丑冠博等基于PPRR理论建立化工罐区应急管理能力评价方法，强调预防、准备、响应和恢复四个阶段的连续性[6]。上述研究说明，液位、压力、可燃气体检测和紧急切断不应孤立设置，而应与巡检、维护、消防和事故水切换形成联动。",
        ),
        (
            "body",
            "储罐结构与完整性研究正在向精细化分析发展。邵新军等研究全接液金属浮盘的抗爆性能，为浮盘连接和密封结构的安全设计提供依据[7]；蒋新生等分析油气爆炸荷载对储罐结构的毁伤机制[8]；康泽天等对环形格栅双层底板储罐开展疲劳强度有限元分析[9]。曲建军等提出根据设备状态和风险等级安排储罐检修[10]，武刚等利用TOPSIS方法进行储罐风险分级预警[11]。这些成果推动储罐设计由满足静态强度要求，进一步扩展到结构抗爆、腐蚀泄漏监测、状态检验和风险排序。",
        ),
        (
            "body",
            "沿海高湿环境还会影响油气回收装置的运行。田素俊等研究吸附剂对含湿油气的吸附性能及热效应，指出水分会改变吸附容量和床层温升特征[12]。因此，平湖油库油气回收系统除满足处理能力外，还应设置气液分离、冷凝预处理和温度监测，以降低高湿油气对吸附单元安全性和稳定性的影响。",
        ),
        ("h2", "国外研究动态"),
        (
            "body",
            "国外油库研究较多采用定量风险评价、数值模拟和空间信息技术分析火灾爆炸及环境影响。Doregar Zavareh等对石油产品储罐火灾爆炸开展环境风险评价，将事故后果沿空气、水体和土壤向人员、生态与设施受体的传播路径纳入分析[13]。Yang等建立油罐区区域能源风险叠加与可视化模型，用空间分析识别多个风险源共同作用下的重点防控区域[14]。这类研究有助于将罐区内部的设备风险与库外敏感目标、主导风向和应急疏散条件结合起来。",
        ),
        (
            "body",
            "在污染排放研究方面，An等基于不同油库源谱观测分析VOCs排放组成及其对臭氧形成的影响，表明储存介质、储罐结构和装卸方式会改变排放特征[15]。在事故演化方面，Malik等将数值模拟与人工神经网络结合，研究不同风速和起火位置下罐区火灾多米诺传播概率[16]。国外研究的共同特点是利用现场监测、计算流体力学和数据模型提高风险识别的时空分辨率，但工程方案仍需由适用地区的法规、标准和设备条件约束。",
        ),
        ("h2", "研究趋势"),
        (
            "body",
            "综合国内外研究，成品油库技术呈现四个发展方向：一是储罐、码头、管道和公路装卸协同组织，以降低等待时间并提高周转效率；二是内浮顶、高效密封、密闭装卸、油气回收和LDAR协同控制VOCs；三是将实时监测、风险分级和预测模型用于设备完整性及事故预警；四是把防火堤、固定消防、事故水收集和应急资源作为完整事故控制链。焦浩宇等对隔堤池火条件下储罐热响应的研究进一步表明，风速、火源范围及相邻罐受热必须联合考虑[17]；董林林从工程设计角度提出罐组布置、消防水、泡沫系统、消防泵动力和事故排水应整体配置[18]。",
        ),
        (
            "body",
            "针对浙江平湖二级成品油库，本设计吸收上述研究中的系统化方法：以年周转量和运输比例确定库容及作业能力，以储罐型式和密封控制汽油蒸发损耗，以最不利路径完成水力计算和泵选型，以风环境及池火热响应完善罐组防护，以分区围控、固定消防和事故水收集限制事故扩大，并通过液位、压力、气体检测和紧急切断提高运行可靠性。研究成果用于完善方案逻辑，不替代现行国家标准的强制性要求。",
        ),
    ]

    for kind, text in blocks:
        if kind == "h1":
            paragraph = insert_paragraph(design_heading, text, heading1_template, page_break=(text == "绪论"))
        elif kind == "h2":
            paragraph = insert_paragraph(design_heading, text, heading2_template)
        else:
            paragraph = insert_paragraph(design_heading, text, body_template)

    # The introduction follows the unnumbered form used in the school example;
    # the engineering chapters therefore retain their original numbering.
    for paragraph in document.paragraphs:
        text = paragraph_text(paragraph)
        if text == "1.2.4 主要设计数据":
            replace_in_runs(paragraph, lambda value: value.replace("1.2.4", "1.2.3"))

    # Map old citation numbers to the order of first appearance in the new text.
    citation_map = {
        1: 19,
        2: 1,
        3: 2,
        4: 3,
        5: 4,
        6: 6,
        7: 17,
        8: 5,
        9: 7,
        10: 8,
        11: 9,
        12: 10,
        13: 11,
        14: 12,
        15: 14,
        16: 13,
        17: 20,
        18: 21,
        19: 22,
    }

    def remap_citations(value: str) -> str:
        return re.sub(r"\[(\d+)\]", lambda m: f"[{citation_map.get(int(m.group(1)), int(m.group(1)))}]", value)

    reference_heading = next(p for p in document.paragraphs if paragraph_text(p) == "参考文献")
    reached_design_body = False
    for paragraph in all_body_paragraphs(document):
        if paragraph._p is design_heading._p:
            reached_design_body = True
        if not reached_design_body:
            continue
        # Bibliography entries are rebuilt below, so leave their numbers alone here.
        if paragraph._p is reference_heading._p:
            continue
        replace_in_runs(paragraph, remap_citations)

    # Rebuild bibliography in exact first-citation order.
    references = [
        "许少新, 涂仁福, 徐宁, 等. 成品油管铁联运物流优化[J]. 油气储运, 2022, 41(7): 859-868. DOI:10.6047/j.issn.1000-8241.2022.07.015.",
        "吴守志, 侯磊, 伍星光, 刘芳媛. 安全屏障对储油罐区池火灾多米诺效应概率的影响[J]. 油气储运, 2022, 41(2): 165-176. DOI:10.6047/j.issn.1000-8241.2022.02.006.",
        "王雷, 申满对, 刘奎. 外浮顶储罐VOCs排放量影响因素分析与探讨[J]. 炼油技术与工程, 2022, 52(10): 55-58.",
        "刘世达, 侯栓弟, 刘忠生, 等. 国内石化有机液体储罐VOCs深度减排管控技术进展[J]. 炼油技术与工程, 2022, 52(4): 11-18.",
        "朱喜平. 大型石油储库全寿命周期风险管控一体化平台研发[J]. 油气储运, 2023, 42(10): 1175-1183. DOI:10.6047/j.issn.1000-8241.2023.10.011.",
        "丑冠博, 刘杰, 多依丽, 等. 基于PPRR理论的化工罐区定量应急管理能力研究[J]. 辽宁石油化工大学学报, 2024, 44(4): 51-59. DOI:10.12422/j.issn.1672-6952.2024.04.007.",
        "邵新军, 周一卉, 黄兆锋, 等. 全接液金属浮盘抗爆特性实验与数值模拟[J]. 油气储运, 2024, 43(2): 200-211. DOI:10.6047/j.issn.1000-8241.2024.02.009.",
        "蒋新生, 秦希卓, 储汇, 等. 油气爆炸荷载对储罐结构的毁伤机制及评估[J]. 油气储运, 2024, 43(12): 1365-1377. DOI:10.6047/j.issn.1000-8241.2024.12.005.",
        "康泽天, 姚冰, 党文义, 等. 环形格栅双层底板立式储罐疲劳强度有限元分析[J]. 油气储运, 2022, 41(8): 939-945. DOI:10.6047/j.issn.1000-8241.2022.08.009.",
        "曲建军, 纪瑞军. 石油库储罐检修策略[J]. 油气储运, 2023, 42(11): 1307-1312. DOI:10.6047/j.issn.1000-8241.2023.11.011.",
        "武刚, 张庶鑫, 罗金恒, 等. 基于TOPSIS的原油储罐风险分级预警[J]. 油气储运, 2024, 43(6): 641-648. DOI:10.6047/j.issn.1000-8241.2024.06.005.",
        "田素俊, 黄维秋, 鄢永兵, 等. 吸附剂对含湿油气的吸附性能及热效应[J]. 油气储运, 2022, 41(8): 962-971. DOI:10.6047/j.issn.1000-8241.2022.08.012.",
        "DOREGAR ZAVAREH R, DANA T, ROAYAEI E, et al. The environmental risk assessment of fire and explosion in storage tanks of petroleum products[J]. Sustainability, 2022, 14(17): 10747. DOI:10.3390/su141710747.",
        "YANG Y, ZHANG X, XIE S, et al. Design and visual implementation of a regional energy risk superposition model for oil tank farms[J]. Energies, 2024, 17(22): 5775. DOI:10.3390/en17225775.",
        "AN W, TONG J, ZHANG L, et al. Characterization of VOC emissions based on oil depots source profiles observations and influence of ozone numerical simulation[J]. Atmosphere, 2025, 16(10): 1192. DOI:10.3390/atmos16101192.",
        "MALIK A A, NASIF M S, ARSHAD U, et al. Predictive modelling of wind-influenced dynamic fire spread probability in tank farm due to domino effect by integrating numerical simulation with ANN[J]. Fire, 2023, 6(3): 85. DOI:10.3390/fire6030085.",
        "焦浩宇, 任婧杰, 赵彦修, 毕明树. 隔堤池火条件下储罐热响应的数值模拟[J]. 化工进展, 2025, 44(12): 7349-7358. DOI:10.16085/j.issn.1000-6613.2024-2040.",
        "董林林. 大型油库消防系统设计[J]. 油气储运, 2011, 30(11): 864-866.",
        "浙江荣晟环保纸业股份有限公司, 嘉兴市环境科学研究所有限公司. 纸机绿色节能提效升级改造项目环境影响报告书（公示稿）[R]. 嘉兴, 2024.",
        "刘德俊, 杨帆, 于洋, 等. 油库技术与管理[M]. 2版. 北京: 中国石化出版社, 2021.",
        "邢科伟, 马秀让, 刘占卿. 油库加油站设计数据图表手册[M]. 北京: 中国石化出版社, 2015.",
        "KSB SE & Co. KGaA. MegaCPK centrifugal pumps with shaft seal: characteristic curves and technical data[EB/OL]. Frankenthal: KSB.",
    ]

    current_paragraphs = list(document.paragraphs)
    ref_index = next(i for i, p in enumerate(current_paragraphs) if p._p is reference_heading._p)
    thanks_index = next(i for i, p in enumerate(current_paragraphs) if paragraph_text(p) == "致谢")
    old_ref_paragraphs = current_paragraphs[ref_index + 1 : thanks_index]
    ref_template = old_ref_paragraphs[0]
    for paragraph in old_ref_paragraphs:
        remove_paragraph(paragraph)
    thanks_heading = next(p for p in document.paragraphs if paragraph_text(p) == "致谢")
    for number, entry in enumerate(references, start=1):
        paragraph = thanks_heading.insert_paragraph_before()
        clone_paragraph_properties(ref_template, paragraph)
        paragraph.add_run(f"[{number}] {entry}")

    document.core_properties.title = "浙江平湖二级油库工艺设计"
    document.core_properties.subject = "毕业设计初稿（绪论与国内外研究动态补充版）"
    document.save(OUTPUT)

    after_formulas = formula_fingerprint(OUTPUT)
    if before_formulas != after_formulas:
        raise RuntimeError(
            f"Formula preservation check failed: before={len(before_formulas)}, after={len(after_formulas)}"
        )
    print(f"OUTPUT={OUTPUT}")
    print(f"FORMULAS={len(before_formulas)} IDENTICAL")


if __name__ == "__main__":
    main()
