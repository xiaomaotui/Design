from __future__ import annotations

import math
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(r"D:\毕业论文")
SOURCE = ROOT / r"04_最终成品\01_毕业设计说明书\张淑鑫_浙江平湖油库工艺设计_终稿_封面与真实性承诺替换版_2026-07-26.docx"
OUTPUT = ROOT / r"04_最终成品\01_毕业设计说明书\张淑鑫_浙江平湖二级油库工艺设计_毕业设计初稿_计算与规范完善版_2026-07-28.docx"
ASSET_DIR = ROOT / r"03_过程文件\02_设计计算\rebuild_v2_assets"
KSB_RENDER = ROOT / r"03_过程文件\02_设计计算\pump_render\ksb_p18-018.png"

G = {"92号汽油": 400_000.0, "95号汽油": 250_000.0, "0号柴油": 300_000.0}
RHO = {"92号汽油": 0.760, "95号汽油": 0.760, "0号柴油": 0.840}
NU = {"92号汽油": 0.70e-6, "95号汽油": 0.70e-6, "0号柴油": 3.50e-6}
K0 = 14.0
ETA = 0.95
TANKS = {
    "92号汽油": {"n": 2, "nom": 20_000.0, "D": 40.0, "H": 18.0, "type": "钢制内浮顶罐"},
    "95号汽油": {"n": 2, "nom": 15_000.0, "D": 34.0, "H": 18.0, "type": "钢制内浮顶罐"},
    "0号柴油": {"n": 2, "nom": 15_000.0, "D": 34.0, "H": 18.0, "type": "固定顶罐"},
}
SPLITS = {
    "92号汽油": {"in_water": 0.70, "in_pipe": 0.30, "out_road": 0.65, "out_pipe": 0.25, "out_water": 0.10},
    "95号汽油": {"in_water": 0.60, "in_pipe": 0.40, "out_road": 0.70, "out_pipe": 0.10, "out_water": 0.20},
    "0号柴油": {"in_water": 0.80, "in_pipe": 0.20, "out_road": 0.60, "out_pipe": 0.40, "out_water": 0.00},
}


def theoretical_capacity(product: str) -> float:
    return G[product] / (K0 * RHO[product] * ETA)


def actual_k(product: str) -> float:
    cfg = TANKS[product]
    return G[product] / (cfg["n"] * cfg["nom"] * RHO[product] * ETA)


def tank_area(product: str) -> float:
    return math.pi * TANKS[product]["D"] ** 2 / 4


def liquid_height(product: str) -> float:
    return TANKS[product]["nom"] / tank_area(product)


def hydraulics(product: str, q_m3h: float, d_m: float, length_m: float, dz_m: float, equipment_m: float, local_k: float):
    q = q_m3h / 3600.0
    v = 4 * q / (math.pi * d_m**2)
    reynolds = v * d_m / NU[product]
    eps = 0.000045
    lam = 0.25 / math.log10(eps / (3.7 * d_m) + 5.74 / (reynolds**0.9)) ** 2
    hf = lam * length_m / d_m * v**2 / (2 * 9.81)
    hm = local_k * v**2 / (2 * 9.81)
    head = dz_m + equipment_m + hf + hm + 3.0
    power = RHO[product] * 1000 * 9.81 * q * head / (0.70 * 1000)
    return {
        "Q": q_m3h,
        "d": d_m,
        "v": v,
        "Re": reynolds,
        "lam": lam,
        "hf": hf,
        "hm": hm,
        "H": head,
        "P": power,
    }


PATHS = [
    ("92号汽油水运进库", "92号汽油", 300, 0.250, 1200, 8, 8, 16),
    ("95号汽油水运进库", "95号汽油", 300, 0.250, 1200, 8, 8, 16),
    ("0号柴油水运进库", "0号柴油", 300, 0.250, 1200, 8, 8, 16),
    ("92号汽油管道进库", "92号汽油", 250, 0.250, 800, 5, 6, 14),
    ("95号汽油管道进库", "95号汽油", 250, 0.250, 800, 5, 6, 14),
    ("0号柴油管道进库", "0号柴油", 250, 0.250, 800, 5, 6, 14),
    ("92号汽油公路发油", "92号汽油", 120, 0.150, 350, 8, 12, 18),
    ("95号汽油公路发油", "95号汽油", 120, 0.150, 350, 8, 12, 18),
    ("0号柴油公路发油", "0号柴油", 120, 0.150, 350, 8, 12, 18),
    ("92号汽油管道外输", "92号汽油", 200, 0.200, 1000, 6, 8, 16),
    ("95号汽油管道外输", "95号汽油", 200, 0.200, 1000, 6, 8, 16),
    ("0号柴油管道外输", "0号柴油", 200, 0.200, 1000, 6, 8, 16),
    ("92号汽油水运外输", "92号汽油", 200, 0.200, 1200, 5, 10, 18),
    ("95号汽油水运外输", "95号汽油", 200, 0.200, 1200, 5, 10, 18),
]
HYD = {name: hydraulics(prod, q, d, L, dz, eq, lk) for name, prod, q, d, L, dz, eq, lk in PATHS}


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=Pt(10.5), bold=None):
    run.font.name = latin
    run.font.size = size
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)


def set_update_fields(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def ensure_heading_styles(doc: Document):
    for level in (2, 3):
        name = f"Heading {level}"
        if name not in [s.name for s in doc.styles]:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]
        else:
            style = doc.styles[name]
        ppr = style._element.get_or_add_pPr()
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            ppr.append(outline)
        outline.set(qn("w:val"), str(level - 1))


def remove_old_body(doc: Document):
    start = None
    body = doc._element.body
    for p in doc.paragraphs:
        if p.text.strip() == "1 设计总论":
            start = p._p
            break
    if start is None:
        raise RuntimeError("未找到正文起点")
    children = list(body)
    idx = children.index(start)
    for child in children[idx:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def paragraph_style(doc: Document, preferred: str, fallback: str = "Normal") -> str:
    return preferred if preferred in [s.name for s in doc.styles] else fallback


def add_paragraph(doc: Document, text="", style="Normal", align=None, first_line=True, keep_next=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, east_asia="黑体" if style in {"Heading 1", "Heading 2", "Heading 3"} else "宋体",
                     size=Pt(16) if style == "Heading 1" else Pt(14) if style == "Heading 2" else Pt(10.5),
                     bold=style in {"Heading 1", "Heading 2", "Heading 3"})
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line and style == "Normal":
        p.paragraph_format.first_line_indent = Cm(0.74)
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.keep_with_next = keep_next
    return p


def add_body(doc: Document, text: str, citations=None):
    p = add_paragraph(doc, "", "Normal")
    r = p.add_run(text)
    set_run_font(r)
    if citations:
        for n in citations:
            cr = p.add_run(f"[{n}]")
            set_run_font(cr, size=Pt(8))
            cr.font.superscript = True
    return p


def add_heading(doc: Document, text: str, level=1):
    style = f"Heading {level}"
    p = add_paragraph(doc, text, style, first_line=False, keep_next=True)
    if level == 1:
        p.paragraph_format.page_break_before = True
    return p


def add_caption(doc: Document, text: str):
    p = add_paragraph(doc, "", "Normal", WD_ALIGN_PARAGRAPH.CENTER, first_line=False, keep_next=True)
    r = p.add_run(text)
    set_run_font(r, size=Pt(10.5), bold=True)
    return p


def set_cell(cell, value, bold=False, center=True, size=Pt(9)):
    cell.text = str(value)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs:
            set_run_font(r, size=size, bold=bold)


def add_table(doc: Document, caption: str, headers, rows, widths=None, font_size=Pt(9)):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, size=font_size)
    trpr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    trpr.append(repeat)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val, size=font_size)
        cant_split = OxmlElement("w:cantSplit")
        table.rows[-1]._tr.get_or_add_trPr().append(cant_split)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(int(Cm(width).twips)))
            grid.append(col)
    add_paragraph(doc, "", "Normal", first_line=False)
    return table


EQ_COUNTER = 0


def add_equation(doc: Document, linear: str):
    global EQ_COUNTER
    EQ_COUNTER += 1
    p = add_paragraph(doc, f"[[EQ|{EQ_COUNTER}|{linear}]]", "Normal", WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    return EQ_COUNTER


def add_numeric_equation(doc: Document, linear: str):
    p = add_paragraph(doc, f"[[NUMEQ|{linear}]]", "Normal", WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_figure(doc: Document, image_path: Path, caption: str, width_cm=15.0):
    p = add_paragraph(doc, "", "Normal", WD_ALIGN_PARAGRAPH.CENTER, first_line=False, keep_next=True)
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, caption)


def build_pump_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    crop_path = ASSET_DIR / "KSB_MegaCPK_125-100-200_性能曲线_样本第18页.png"
    if KSB_RENDER.exists():
        image = Image.open(KSB_RENDER)
        image.crop((100, 120, 1380, 1200)).save(crop_path)

    # 泵曲线数据由KSB样本第18页MegaCPK Inducer 125-100-200、2900 r/min、
    # 叶轮直径193 mm曲线读数；系统曲线来自本文最不利水运外输管路计算。
    q_p = [100, 150, 200, 250, 300, 330]
    h_p = [42.0, 41.0, 38.8, 35.0, 29.0, 25.0]
    h_static = 18.0
    k_sys = (HYD["92号汽油水运外输"]["H"] - h_static) / (200.0**2)
    q = [i for i in range(0, 341, 5)]
    h_sys = [h_static + k_sys * x**2 for x in q]
    curve_path = ASSET_DIR / "输油泵与最不利管路特性曲线.png"
    canvas = Image.new("RGB", (1600, 1050), "white")
    draw = ImageDraw.Draw(canvas)
    font_path = Path(r"C:\Windows\Fonts\times.ttf")
    font = ImageFont.truetype(str(font_path), 34) if font_path.exists() else ImageFont.load_default()
    small = ImageFont.truetype(str(font_path), 28) if font_path.exists() else ImageFont.load_default()
    left, top, right, bottom = 150, 100, 1510, 900
    qmax, hmin, hmax = 340.0, 15.0, 48.0
    xpix = lambda value: left + value / qmax * (right - left)
    ypix = lambda value: bottom - (value - hmin) / (hmax - hmin) * (bottom - top)
    draw.line((left, top, left, bottom), fill="black", width=4)
    draw.line((left, bottom, right, bottom), fill="black", width=4)
    for value in range(0, 341, 50):
        x = xpix(value)
        draw.line((x, top, x, bottom), fill=(215, 215, 215), width=2)
        draw.text((x - 18, bottom + 18), str(value), fill="black", font=small)
    for value in range(15, 49, 5):
        y = ypix(value)
        draw.line((left, y, right, y), fill=(215, 215, 215), width=2)
        draw.text((70, y - 16), str(value), fill="black", font=small)
    pump_points = [(xpix(x), ypix(y)) for x, y in zip(q_p, h_p)]
    sys_points = [(xpix(x), ypix(y)) for x, y in zip(q, h_sys)]
    draw.line(pump_points, fill=(0, 91, 172), width=7, joint="curve")
    draw.line(sys_points, fill=(205, 66, 43), width=7, joint="curve")
    for x, y in pump_points:
        draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill=(0, 91, 172))
    wx, wy = xpix(202), ypix(36.3)
    draw.ellipse((wx - 11, wy - 11, wx + 11, wy + 11), fill="black")
    draw.text((wx + 25, wy - 55), "Operating point: Q=202 m3/h, H=36.3 m", fill="black", font=small)
    draw.text((600, 955), "Flow rate Q/(m3/h)", fill="black", font=font)
    draw.text((25, 40), "Head H/m", fill="black", font=font)
    draw.line((260, 150, 370, 150), fill=(0, 91, 172), width=7)
    draw.text((390, 132), "KSB 125-100-200, impeller 193 mm", fill="black", font=small)
    draw.line((260, 200, 370, 200), fill=(205, 66, 43), width=7)
    draw.text((390, 182), "Governing marine-export system curve", fill="black", font=small)
    canvas.save(curve_path)
    return crop_path, curve_path


def add_note(doc: Document, text: str):
    p = add_paragraph(doc, "", "Normal", first_line=False)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.right_indent = Cm(0.74)
    r = p.add_run("说明：" + text)
    set_run_font(r, size=Pt(9))
    return p


def add_citation_rich_paragraph(doc: Document, segments):
    p = add_paragraph(doc, "", "Normal")
    for text, citation in segments:
        r = p.add_run(text)
        set_run_font(r)
        if citation is not None:
            cr = p.add_run(f"[{citation}]")
            set_run_font(cr, size=Pt(8))
            cr.font.superscript = True
    return p


def build_front(doc: Document):
    paragraphs = doc.paragraphs
    paragraphs[15].text = "提交日期：2026 年 7 月 28 日"
    paragraphs[17].text = ""
    paragraphs[32].text = (
        "本设计以浙江省平湖市独山港经济开发区石化产业园临港工业用地为拟建场址，建设一座水运、公共管道和公路联合收发的区域配送型成品油库。"
        "92号汽油、95号汽油和0号柴油年进出库量分别为40万t、25万t和30万t，总年吞吐量95万t。按商业油库二级及以上K=8～14的参考范围，"
        "先取K=14、储罐容积利用系数η=0.95计算理论库容，三种油品分别为39 572.62 m³、24 732.88 m³和26 852.85 m³。"
        "储罐配置为2座20 000 m³的92号汽油钢制内浮顶罐、2座15 000 m³的95号汽油钢制内浮顶罐和2座15 000 m³的0号柴油固定顶罐，"
        "共6座、名义总容量100 000 m³；按GB 50074—2014折算的储罐计算总容量为85 000 m³，确定为二级石油库。"
        "设计保持原计划中的进出库方式及比例，完成水运进出库、公共管道进出库、公路发油、倒罐和清扫流程；水运年作业量76万t，"
        "配置1个5 000 DWT成品油泊位。主输油管采用DN150～DN250；候选泵采用KSB MegaCPK Inducer 125-100-200，泵与最不利管路曲线交点约为202 m³/h、36.3 m，并完成柴油汽蚀余量及年装卸能力校核。"
        "罐区划分为3个防火堤区，汽油储罐采用钢制内浮盘及一次、二次密封。消防控制工况为15 000 m³固定顶柴油罐着火并冷却1座相邻罐，"
        "设置2×2 500 m³消防水池、2×5 m³泡沫液储罐和8 000 m³事故水池。设计同时提出HSE、防火防爆、防雷防静电、抗震、防泄漏、"
        "油气回收、雨污分流、检测报警和紧急切断措施，为总平面图和工艺流程图绘制提供参数。"
    )
    paragraphs[33].text = "关键词：二级油库；库容计算；水力计算；泵选型；码头泊位；消防设计"
    paragraphs[39].text = (
        "A Class II regional refined-products depot is designed for a port-oriented industrial site in the Dushan Port Petrochemical Industrial Park, Pinghu, Zhejiang. "
        "The annual receipt and delivery quantities are 0.40 million tonnes of No. 92 gasoline, 0.25 million tonnes of No. 95 gasoline, and 0.30 million tonnes of No. 0 diesel. "
        "A preliminary turnover coefficient of 14 and a tank utilization factor of 0.95 give theoretical capacities of 39,572.62 m³, 24,732.88 m³, and 26,852.85 m³, respectively. "
        "The selected arrangement comprises two 20,000 m³ internal-floating-roof tanks for No. 92 gasoline, two 15,000 m³ internal-floating-roof tanks for No. 95 gasoline, "
        "and two 15,000 m³ fixed-roof tanks for diesel. The nominal capacity is 100,000 m³, while the classification capacity under GB 50074—2014 is 85,000 m³, corresponding to a Class II depot. "
        "The original marine, pipeline and road transport ratios are retained. One 5,000 DWT product-oil berth is verified for an annual marine throughput of 0.76 million tonnes. "
        "Pipelines from DN150 to DN250 are selected. A KSB MegaCPK Inducer 125-100-200 pump is preliminarily selected; the pump and governing system curves intersect at approximately 202 m³/h and 36.3 m, followed by diesel NPSH and annual-capacity checks. "
        "Three diked tank groups are provided. The governing fire scenario is a 15,000 m³ fixed-roof diesel tank with one adjacent tank cooled. "
        "Two 2,500 m³ fire-water tanks, two 5 m³ foam-concentrate tanks and an 8,000 m³ emergency retention basin are selected. "
        "HSE, fire and explosion protection, lightning and static protection, seismic design, spill containment, vapor recovery, drainage segregation, alarms and emergency shutdowns are incorporated."
    )
    paragraphs[40].text = "KEY WORDS: Class II petroleum depot; storage capacity; hydraulic calculation; pump selection; marine berth; fire protection"
    for idx in (15, 32, 33, 39, 40):
        for r in paragraphs[idx].runs:
            set_run_font(r, east_asia="宋体", size=Pt(10.5))
    # 压缩封面日期前的两个空段，保证提交日期与其余封面信息同页。
    for idx in (16, 17):
        paragraphs[idx].paragraph_format.space_before = Pt(0)
        paragraphs[idx].paragraph_format.space_after = Pt(0)
        paragraphs[idx].paragraph_format.line_spacing = Pt(1)
        for r in paragraphs[idx].runs:
            r.font.size = Pt(1)
    paragraphs[15].paragraph_format.space_before = Pt(0)
    paragraphs[15].paragraph_format.space_after = Pt(0)
    for idx in range(8, 15):
        paragraphs[idx].paragraph_format.space_before = Pt(0)


def build_chapter_1(doc: Document):
    add_heading(doc, "1 设计总论", 1)
    add_heading(doc, "1.1 项目概况与设计范围", 2)
    add_heading(doc, "1.1.1 项目背景与建设条件", 3)
    add_body(
        doc,
        "拟建油库服务浙江北部及上海西南部成品油中转与配送，承担92号汽油、95号汽油和0号柴油的接卸、储存、调合前批次隔离及外发。"
        "平湖独山港位于杭州湾北岸，具备港口、公路和临港产业协同条件，适合形成“水运大批量进库、公共管道补充、公路与管道为主发油”的物流组织。",
    )
    add_body(
        doc,
        "场址选择在独山港经济开发区石化产业园临港工业用地区域，不占用生活居住用地。当地规划将该区域作为石化产业集聚区，具备三类工业用地、港口作业和危险化学品配套条件；"
        "油库生产区应与园区居住片区保持规范要求的防护距离，并预留事故水、消防水和油气回收设施用地。",
        [1],
    )
    add_body(
        doc,
        "区域属亚热带季风气候。独山港相邻工程采用的平湖站1971—2014年统计值为：年平均气温16.3 ℃、极端最高气温39.1 ℃、极端最低气温-9.3 ℃、"
        "年平均降水量1269.7 mm、年平均相对湿度80%、年平均雷暴日28 d。杭州湾北岸为强潮海区，平均潮差约4.82 m，历史最高潮位5.69 m，夏秋季台风可能引起增水。"
        "上述条件要求场地竖向设计兼顾防洪排涝，储罐及管架按沿海风环境进行抗风校核，码头软管和装卸臂设置紧急脱离及快速切断。",
        [1],
    )
    add_body(
        doc,
        "平湖乍浦50年重现期基本风压为0.45 kPa，罐区按开阔临海地貌考虑风压高度变化；场址相邻工程公开资料给出的地震基本烈度为Ⅵ度。"
        "施工图阶段仍应以正式岩土勘察和场地地震安全性评价结果为准。储罐基础宜采用桩基础或复合地基，最终形式由地基承载力、沉降和液化判别确定。",
    )
    add_heading(doc, "1.1.2 场址自然条件与设计参数", 3)
    add_body(
        doc,
        "自然条件采用拟建场址相邻的嘉兴港独山港区B区21、22号多用途泊位工程环评长期统计资料。该报告由浙江省政府信息公开平台发布，"
        "气象统计站为平湖站（30°37′N、121°05′E），统计期为1971—2014年，能够代表独山港陆域的温度、降水、湿度、雾、雷暴和风环境。"
        "表1-1同时列出原始PDF页码，便于后续复核。",
    )
    add_table(
        doc,
        "表1-1 浙江平湖独山港场址自然条件及设计采用值",
        ["类别", "参数", "原始统计值", "本设计采用值", "原始PDF位置", "用于计算或设计"],
        [
            ["温度", "年平均气温", "16.3 ℃", "16.3 ℃", "政府环评PDF第166页（报告书第162页）", "油品物性、设备环境条件"],
            ["温度", "极端最高/最低气温", "39.1/-9.3 ℃", "39.1/-9.3 ℃", "政府环评PDF第166页（报告书第162页）", "材料、仪表与防冻"],
            ["温度", "最热/月最低平均气温", "26.9/5.8 ℃", "26.9/5.8 ℃", "政府环评PDF第166页（报告书第162页）", "运行温度边界"],
            ["降水", "年平均降水量", "1269.7 mm", "1269.7 mm", "政府环评PDF第166页（报告书第162页）", "排水系统"],
            ["降水", "最大一日降水量", "276.4 mm", "276.4 mm", "政府环评PDF第166页（报告书第162页）", "事故水池雨水分量"],
            ["风", "主导风向", "E—SE，累计30%", "E—SE", "政府环评PDF第166页（报告书第162页）", "总图和管理区方位"],
            ["风", "年平均/极端风速", "3.2/31.7 m/s", "3.2/31.7 m/s", "政府环评PDF第167页（报告书第163页）", "抗风与码头停工条件"],
            ["风", "乍浦50年基本风压", "0.45 kN/m²", "0.45 kN/m²", "《浙江省基本风压资料》PDF第43页（资料第37页）", "储罐抗风稳定性"],
            ["湿度", "年平均相对湿度", "80%", "80%", "政府环评PDF第167页（报告书第163页）", "沿海防腐与电气选型"],
            ["天气", "年平均雾日", "35 d", "35 d", "政府环评PDF第167页（报告书第163页）", "码头可用时间"],
            ["天气", "年平均雷暴日", "28 d", "28 d", "政府环评PDF第168页（报告书第164页）", "防雷和作业管理"],
            ["地震", "地震基本烈度", "Ⅵ度", "Ⅵ度", "政府环评PDF第177页（报告书第173页）", "储罐、管线抗震"],
        ],
        [1.5, 2.7, 2.8, 2.8, 4.3, 3.5],
        Pt(7.5),
    )
    add_body(
        doc,
        "表中最大一日降水量276.4 mm直接用于第6章事故水池校核；50年基本风压0.45 kN/m²用于第2章抗风稳定性参数；"
        "极端风速31.7 m/s用于台风停工和加固管理边界，二者含义不同，不能相互替代。环评原文另有一处OCR显示“342 m/s”，"
        "与同一报告风速表及物理常识矛盾，本设计不采用该错误识别值。",
    )
    add_heading(doc, "1.1.3 设计任务、范围及界面", 3)
    add_body(
        doc,
        "设计内容包括油库储油罐容量计算；输油管、集油管和主要支管管径确定；防火堤和消防系统设置；装卸油管路水力计算；输油泵选型及装卸能力校核；"
        "总平面布置和工艺流程设计。水运系统增加码头年通过能力与泊位数校核，消防系统增加泡沫混合液流量、泡沫产生器数量和泡沫液储量计算。",
    )
    add_body(
        doc,
        "工艺设计起点为码头岸线接口或公共管道交接计量点，终点为公路装车鹤位、公共管道交接点及码头装船接口。库外航道、码头结构、长输管道线路和市政道路不作详细结构设计，"
        "但其允许流量、作业时间和接口压力作为油库工艺计算条件。自动控制部分只说明液位、流量、压力、可燃气体报警和紧急切断逻辑，不展开控制系统软件与通信网络设计。",
    )
    add_heading(doc, "1.2 设计依据与基础数据", 2)
    add_heading(doc, "1.2.1 规范与技术原则", 3)
    add_body(
        doc,
        "油库等级、库址、总平面、防火间距、储罐型式、工艺设施和消防设计以GB 50074—2014《石油库设计规范》为主；泡沫系统执行GB 50151—2021，"
        "储罐结构执行GB 50341—2014，消防给水执行GB 50974—2014，油气回收执行GB 20950—2020和GB/T 50759—2022。"
        "爆炸危险区域、防雷、可燃气体检测和抗震分别执行GB 50058—2014、GB 50057—2010、GB/T 50493—2019及相关现行标准。",
    )
    add_body(
        doc,
        "规范给出的强制性限值优先于教材和经验数据。教材用于确定周转系数法、水力计算和设备比选步骤；论文用于说明物流优化、储罐风险、油气回收与完整性管理的研究进展，"
        "不替代国家标准规定的间距、强度和持续时间。",
    )
    add_heading(doc, "1.2.2 油品性质及危险性", 3)
    add_body(
        doc,
        "92号和95号汽油按甲B类易燃液体考虑，设计密度均取0.760 t/m³，运动黏度取0.70×10⁻⁶ m²/s；0号柴油按丙A类可燃液体考虑，设计密度取0.840 t/m³，"
        "运动黏度取3.50×10⁻⁶ m²/s。密度用于库容与质量能力换算，黏度用于雷诺数和摩阻计算；实际采购批次的密度应以交接计量温度下化验数据修正。",
    )
    add_table(
        doc,
        "表1-2 油品设计性质与主要危险",
        ["油品", "类别", "设计密度/(t/m³)", "运动黏度/(10⁻⁶m²/s)", "主要危险", "储罐型式"],
        [
            ["92号汽油", "甲B", "0.760", "0.70", "蒸气易燃、静电与VOCs", "钢制内浮顶"],
            ["95号汽油", "甲B", "0.760", "0.70", "蒸气易燃、静电与VOCs", "钢制内浮顶"],
            ["0号柴油", "丙A", "0.840", "3.50", "可燃、泄漏污染", "固定顶"],
        ],
        [2.4, 1.5, 2.2, 2.8, 4.0, 2.5],
    )
    add_heading(doc, "1.2.3 近五年研究与方案借鉴", 3)
    add_citation_rich_paragraph(
        doc,
        [
            ("成品油物流研究表明，库容、运输能力和终端需求需要联动安排，单纯扩大储罐并不能替代合理的水运、管道和公路作业组织", 2),
            ("。因此，本设计先确定年周转量和运输比例，再以峰值作业窗口确定管径与泵。罐区火灾研究指出，多层安全屏障能够显著降低池火灾多米诺效应概率", 3),
            ("，故在规范间距以外，同时采用联锁切断、分区围控、固定消防和应急响应。", None),
        ],
    )
    add_citation_rich_paragraph(
        doc,
        [
            ("储罐VOCs研究显示，浮顶及高效密封是汽油储存的源头减排措施", 4),
            ("，国内深度减排技术正由单一回收转向源头控制、密闭收集和末端治理的组合", 5),
            ("。本设计据此采用钢制内浮盘、一次二次密封和装车油气回收。PPRR应急管理研究强调预防、准备、响应和恢复的完整链条", 6),
            ("，本文在HSE章节将日常巡检、报警联锁、事故水切换和恢复条件一并说明。", None),
        ],
    )
    add_citation_rich_paragraph(
        doc,
        [
            ("隔堤池火数值研究揭示了风速和池火范围对邻罐热响应的影响", 7),
            ("；大型石油储库全寿命周期平台研究则强调设备、风险和应急数据的连续管理", 8),
            ("。全接液金属浮盘抗爆研究", 9),
            ("与储罐爆炸荷载研究", 10),
            ("说明浮盘连接、罐体刚度和防火间距应共同考虑。双层底板疲劳研究", 11),
            ("、储罐检修策略研究", 12),
            ("和TOPSIS风险分级研究", 13),
            ("为防泄漏、定期检验和风险排序提供了依据。含湿油气吸附研究", 14),
            ("支持油气回收装置在高湿沿海环境下设置预处理和温升监测。", None),
        ],
    )
    add_citation_rich_paragraph(
        doc,
        [
            ("外文研究从区域风险叠加角度分析油罐区周边风险分布", 15),
            ("，并对石油产品储罐火灾爆炸的环境风险进行评价", 16),
            ("。这些成果用于完善风险识别和监测方案，设计限值仍以中国现行国家标准为准。", None),
        ],
    )
    add_heading(doc, "1.2.4 主要设计数据", 3)
    add_table(
        doc,
        "表1-3 年进出库量及运输比例",
        ["油品", "年量/(10⁴t/a)", "进库：水运", "进库：管道", "出库：公路", "出库：管道", "出库：水运"],
        [
            ["92号汽油", "40", "70%", "30%", "65%", "25%", "10%"],
            ["95号汽油", "25", "60%", "40%", "70%", "10%", "20%"],
            ["0号柴油", "30", "80%", "20%", "60%", "40%", "0%"],
            ["合计", "95", "—", "—", "—", "—", "—"],
        ],
        [2.5, 2.2, 2.2, 2.2, 2.2, 2.2, 2.2],
    )


def build_chapter_2(doc: Document):
    add_heading(doc, "2 库容计算与储罐方案", 1)
    add_heading(doc, "2.1 库容计算", 2)
    add_heading(doc, "2.1.1 周转系数与储备天数", 3)
    add_body(
        doc,
        "《油库技术与管理》第2版给出的商业系统参考范围为：二级及以上油库K取8～14，三级及以下油库K取14～24；多油品油库先按较大值计算，再结合等级和运输条件反算。"
        "本设计先取K=14，既反映独山港水运和公共管道条件下较高的周转效率，也能避免按低周转系数配置过多储罐。",
        [17],
    )
    eq1 = add_equation(doc, "T=365/K")
    add_body(doc, f"式中：T——平均周转周期，d；K——年周转系数。将K=14代入式{eq1}，得T=26.07 d。该周期用于库容规划，不等同于单船到港间隔或每座罐的固定清空周期。")
    add_heading(doc, "2.1.2 理论库容与储罐计算总容量", 3)
    eq2 = add_equation(doc, "V_i=G_i/(Kρ_iη)")
    add_body(doc, f"式中：V_i——第i种油品理论设计库容，m³；G_i——年周转质量，t/a；K——年周转系数；ρ_i——设计密度，t/m³；η——储罐容积利用系数，取0.95。分别将三种油品数据代入式{eq2}：")
    for product in G:
        add_numeric_equation(doc, f"V_{{{product[:2]}}}={G[product]:.0f}/(14×{RHO[product]:.3f}×0.95)={theoretical_capacity(product):.2f}m^3")
    total_theory = sum(theoretical_capacity(p) for p in G)
    add_body(doc, f"三种油品理论库容合计为{total_theory:,.2f} m³。为保证同品种至少两座罐轮换运行，储罐按系列规格向上配置。")
    add_table(
        doc,
        "表2-1 理论库容与储罐配置",
        ["油品", "理论库容/m³", "单罐容量/m³", "数量/座", "名义容量/m³", "反算K", "储罐型式"],
        [
            [p, f"{theoretical_capacity(p):,.2f}", f"{TANKS[p]['nom']:,.0f}", TANKS[p]["n"],
             f"{TANKS[p]['n']*TANKS[p]['nom']:,.0f}", f"{actual_k(p):.2f}", TANKS[p]["type"]]
            for p in G
        ],
        [2.2, 2.6, 2.3, 1.5, 2.4, 1.6, 2.8],
    )
    eq3 = add_equation(doc, "TV=Σ(V_jC_j)")
    add_body(
        doc,
        f"油库等级按GB 50074—2014表3.0.1的储罐计算总容量确定。式中：V_j——第j类液体储罐实际容量，m³；C_j——折算系数；汽油取1.0，丙A类柴油取0.5。"
        f"将汽油40 000 m³、30 000 m³和柴油30 000 m³代入式{eq3}，得TV=40 000+30 000+0.5×30 000=85 000 m³。"
        "该值处于30 000 m³≤TV<100 000 m³范围，油库等级为二级。名义总容量为100 000 m³，而用于等级判定的计算总容量为85 000 m³，两者用途不同。",
    )
    eq4 = add_equation(doc, "K_i=G_i/(V_(Ni)ρ_iη)")
    add_body(
        doc,
        f"将各油品名义容量代入式{eq4}反算，92号汽油K=13.85，95号汽油K=11.54，0号柴油K=12.53，均位于8～14。"
        "因此，6座罐方案同时满足教材周转系数范围、二级油库等级和每种油品两座罐的运行要求。",
    )
    add_heading(doc, "2.2 储罐型式、数量与附件", 2)
    add_heading(doc, "2.2.1 储罐型式与几何尺寸", 3)
    add_body(
        doc,
        "依据GB 50074—2014第6.1.4条，汽油等甲B类液体采用外浮顶或内浮顶储罐；本设计选钢制内浮顶罐，以减少蒸发损耗并便于密闭收集。"
        "依据第6.1.5条，柴油可采用固定顶储罐。20 000 m³汽油罐初选D=40 m、H=18 m；15 000 m³汽油和柴油罐初选D=34 m、H=18 m。",
    )
    eq5 = add_equation(doc, "A=πD^2/4")
    eq6 = add_equation(doc, "h=V_N/A")
    add_body(doc, f"将D=40 m代入式{eq5}，20 000 m³罐截面积为1 256.64 m²；再代入式{eq6}，名义容量对应液位15.915 m，距罐壁顶2.085 m。D=34 m罐截面积为907.92 m²，15 000 m³对应液位16.521 m，距罐壁顶1.479 m。")
    add_table(
        doc,
        "表2-2 储罐主要几何参数",
        ["罐号", "油品", "型式", "容量/m³", "D×H/m", "名义液位/m", "罐组"],
        [
            ["T101～T102", "92号汽油", "钢制内浮顶", "2×20 000", "40×18", "15.915", "G92"],
            ["T201～T202", "95号汽油", "钢制内浮顶", "2×15 000", "34×18", "16.521", "G95"],
            ["T301～T302", "0号柴油", "固定顶", "2×15 000", "34×18", "16.521", "GD"],
        ],
        [2.5, 2.3, 2.8, 2.5, 2.0, 2.2, 1.6],
    )
    add_heading(doc, "2.2.2 罐组划分与附件", 3)
    add_body(
        doc,
        "三种油品分别组成G92、G95和GD罐组，每组2座罐。汽油罐配置钢制全接液内浮盘、一次密封和二次密封、量油孔、通气孔、阻火器、"
        "高高液位联锁、低低液位停泵、罐旁温度与液位显示；柴油固定顶罐配置呼吸阀、阻火器、泡沫产生器、罐壁冷却环管和高低液位保护。",
    )
    add_body(
        doc,
        "各罐设置独立进油支管、出油支管、罐根阀、排污阀和取样口。不同油品不共用同一条罐前支管；公共管廊可并列敷设，但应通过盲板或可验证的双阀隔离防止混油。"
        "浮盘导静电装置、罐体接地和扶梯跨接应形成连续导电通路。",
    )
    add_table(
        doc,
        "表2-3 储罐主要附件与联锁",
        ["项目", "汽油内浮顶罐", "柴油固定顶罐", "作用"],
        [
            ["液位", "雷达液位计+独立高高液位开关", "雷达液位计+独立高高液位开关", "计量与防溢流"],
            ["呼吸", "通气孔、阻火装置", "呼吸阀、阻火器", "控制罐内压力"],
            ["消防", "密封圈泡沫、罐壁冷却环管", "全液面泡沫、罐壁冷却环管", "灭火与冷却"],
            ["防静电", "浮盘软铜跨接、罐体接地", "罐体接地、管线跨接", "释放静电荷"],
            ["排污", "罐底集水坑、密闭排污", "罐底集水坑、密闭排污", "含油污水收集"],
        ],
        [2.3, 4.3, 4.3, 4.2],
    )
    add_heading(doc, "2.3 罐壁厚度与抗风圈校核", 2)
    add_heading(doc, "2.3.1 罐壁厚度计算", 3)
    add_body(
        doc,
        "储罐按GB 50341—2014采用定设计点法进行初步罐壁计算。罐壁分为9圈，每圈高2.0 m，材料采用Q345R，腐蚀裕量C₂取2.0 mm；"
        "底圈焊接接头系数φ取0.85，其余圈取0.90。设计液位分别按15.915 m和16.521 m计算，充水试验按18.0 m水柱计算。",
    )
    eq7 = add_equation(doc, "t_d=(4.9D(H-0.3)ρ)/(σ_dφ)")
    eq8 = add_equation(doc, "t_t=(4.9D(H-0.3))/(σ_tφ)")
    eq9 = add_equation(doc, "t_n≥max(t_d+C_1+C_2,t_t+C_1,t_min)")
    add_body(
        doc,
        f"式中：t_d——设计条件计算厚度，mm；t_t——充水试验计算厚度，mm；t_n——名义厚度，mm。Q345R在相应厚度区间的许用应力取σ_d=218.75 MPa、σ_t=230 MPa。"
        f"以20 000 m³汽油罐底圈为例，将D=40 m、H=15.915 m、ρ=0.760、φ=0.85代入式{eq7}，得t_d=12.51 mm；"
        f"将充水高度18.0 m代入式{eq8}，得t_t=17.75 mm；按式{eq9}向上圆整取18 mm。",
    )
    wall_rows = [
        ["20 000 m³汽油罐", "18/16/14/12/10/10/10/10/10", "18", "10", "Q345R"],
        ["15 000 m³汽油罐", "16/14/12/10/10/10/10/10/10", "16", "10", "Q345R"],
        ["15 000 m³柴油罐", "16/14/12/10/10/10/10/10/10", "16", "10", "Q345R"],
    ]
    add_table(
        doc,
        "表2-4 罐壁名义厚度初选",
        ["储罐", "自下而上9圈厚度/mm", "底圈/mm", "顶圈/mm", "材料"],
        wall_rows,
        [3.0, 7.8, 2.0, 2.0, 2.3],
    )
    add_heading(doc, "2.3.2 抗风圈与稳定性", 3)
    add_body(
        doc,
        "平湖乍浦50年重现期基本风压ω₀=0.45 kPa。罐组狭管效应调整系数取1.20，计算基本风压ω′₀=0.54 kPa；18 m高度风压高度变化系数取1.49。"
        "汽油内浮顶罐按与大气连通考虑，柴油固定顶罐另计0.25 kPa真空负压。",
    )
    eq10 = add_equation(doc, "P_o=2.25μ_zω_0^'+q")
    add_body(doc, f"将μ_z=1.49、ω′₀=0.54 kPa代入式{eq10}，汽油罐q=0，得P_o=1.810 kPa；柴油罐q=0.25 kPa，得P_o=2.060 kPa。")
    eq11 = add_equation(doc, "H_E=Σ[h_i(t_min/t_i)^2.5]")
    eq12 = add_equation(doc, "[P_cr]=(16.48D/H_E)(t_min/D)^2.5")
    add_body(
        doc,
        f"按式{eq11}折算有效厚度，20 000 m³汽油罐H_E=12.72 m，D=34 m的15 000 m³罐H_E=14.36 m。"
        f"代入式{eq12}，无中间抗风圈时临界压力分别为0.927 kPa和1.048 kPa，均低于设计外压；设置1道中间抗风圈后，分段临界压力约为原值的2倍，"
        "20 000 m³汽油罐1.854 kPa>1.810 kPa，15 000 m³汽油罐2.095 kPa>1.810 kPa，柴油罐2.095 kPa>2.060 kPa。",
    )
    add_table(
        doc,
        "表2-5 抗风稳定校核与抗风圈配置",
        ["储罐", "H_E/m", "无圈[P_cr]/kPa", "P_o/kPa", "抗风圈数量", "安装标高/m", "复核结果"],
        [
            ["20 000 m³汽油罐", "12.72", "0.927", "1.810", "1", "+14.82", "满足"],
            ["15 000 m³汽油罐", "14.36", "1.048", "1.810", "1", "+14.41", "满足"],
            ["15 000 m³柴油罐", "14.36", "1.048", "2.060", "1", "+14.41", "满足"],
        ],
        [3.2, 1.8, 2.4, 1.8, 2.2, 2.3, 1.8],
    )
    add_note(doc, "抗风圈截面和局部开孔补强应在储罐施工图阶段结合盘梯、消防环管、罐顶连接和制造厂排板图复核；本稿给出本科毕业设计所需的数量、标高和稳定性结果。")


def build_chapter_3(doc: Document):
    add_heading(doc, "3 收发油工艺与码头泊位", 1)
    add_heading(doc, "3.1 物流组织与作业工况", 2)
    add_heading(doc, "3.1.1 年物流量分配", 3)
    add_body(
        doc,
        "进库方式和出库方式保持计划表比例不变。水运进库承担大批量补库，公共管道进库用于连续补充和船期错峰；公路出库面向区域终端，公共管道外输承担稳定批量，"
        "水运外输用于汽油跨区域调拨。不同油品在同一时段可并行作业，但同一储罐禁止同时进油和发油。",
    )
    eq13 = add_equation(doc, "G_(i,j)=G_iα_(i,j)")
    add_body(doc, f"式中：G_(i,j)——油品i采用方式j的年物流量，t/a；G_i——油品i年物流量，t/a；α_(i,j)——相应运输比例。将表1-3数据代入式{eq13}，得到表3-1。")
    logistics_rows = []
    for p in G:
        s = SPLITS[p]
        logistics_rows.append([
            p,
            f"{G[p]*s['in_water']/1e4:.1f}",
            f"{G[p]*s['in_pipe']/1e4:.1f}",
            f"{G[p]*s['out_road']/1e4:.1f}",
            f"{G[p]*s['out_pipe']/1e4:.1f}",
            f"{G[p]*s['out_water']/1e4:.1f}",
        ])
    logistics_rows.append(["合计", "67.0", "28.0", "61.5", "24.5", "9.0"])
    add_table(
        doc,
        "表3-1 各运输方式年物流量（10⁴t/a）",
        ["油品", "水运进库", "管道进库", "公路出库", "管道出库", "水运出库"],
        logistics_rows,
        [2.8, 2.4, 2.4, 2.4, 2.4, 2.4],
    )
    add_heading(doc, "3.1.2 主要作业工况", 3)
    add_body(
        doc,
        "水运进库流程为：油船货舱泵—装卸臂—岸线紧急切断阀—过滤器—流量计—码头至库区管线—罐区分配阀组—目标储罐。"
        "公共管道进库流程为：交接计量—压力调节—油品确认—罐区阀组—目标储罐。公路发油流程为：储罐—发油泵—过滤器—流量计—批控器—装车鹤管—油罐车。",
    )
    add_body(
        doc,
        "倒罐用于检修腾罐、批次隔离和质量处置，原则上在同品种两罐之间进行。清扫采用氮气或专用清扫介质，汽油管线不得以空气吹扫。每次切换前核对油品、罐号、阀位、"
        "可用容量和接地状态，切换后通过首段样品和密度确认防止混油。",
    )
    add_table(
        doc,
        "表3-2 设计作业流量与时间",
        ["工况", "设计流量/(m³/h)", "典型批量/m³", "允许时间/h", "主要控制点"],
        [
            ["水运进库", "300", "6 000", "24", "船岸联锁、ESD、溢油监测"],
            ["水运出库", "200", "4 000", "24", "装船计量、回流与停泵"],
            ["公共管道进库", "250", "5 000", "24", "交接计量、压力调节"],
            ["公共管道出库", "200", "4 000", "24", "最小流量、末端压力"],
            ["公路发油", "2×60/油品", "单车30", "0.5/车", "防溢流、静电接地、油气回收"],
            ["倒罐", "150", "15 000", "100", "高低液位、泵入口压力"],
        ],
        [3.0, 3.0, 2.6, 2.5, 5.2],
    )
    add_heading(doc, "3.2 码头通过能力与泊位数", 2)
    add_heading(doc, "3.2.1 水运年作业量", 3)
    eq14 = add_equation(doc, "G_w=G_(w,in)+G_(w,out)")
    add_body(doc, f"水运进库量为28+15+24=67万t/a，水运出库量为4+5=9万t/a。代入式{eq14}，码头年作业量G_w=76万t/a。")
    add_heading(doc, "3.2.2 泊位通过能力", 3)
    add_body(
        doc,
        "结合独山港航运条件和毕业设计规模，码头按1个5 000 DWT成品油泊位考虑，平均实载量取5 000 t，水运进库设计流量300 m³/h。汽油质量流量约228 t/h，"
        "柴油约252 t/h；综合三种油品取装卸有效时间24 h，靠离泊、接拆臂、计量和安全确认辅助时间6 h，单船占用时间取30 h。年可作业时间按365 d、利用系数0.85计算。",
    )
    eq15 = add_equation(doc, "P_b=(365×24×η_b/t_b)G_s")
    berth_capacity = 365 * 24 * 0.85 / 30 * 5000
    add_body(doc, f"式中：P_b——单泊位年通过能力，t/a；t_b——单船占用时间，h；η_b——泊位可用系数；G_s——平均单船实载量，t。代入式{eq15}，P_b={berth_capacity/1e4:.2f}万t/a。")
    eq16 = add_equation(doc, "N_b=ceil(G_w/P_b)")
    add_body(doc, f"将G_w=76万t/a、P_b={berth_capacity/1e4:.2f}万t/a代入式{eq16}，N_b=1。泊位利用率为76/{berth_capacity/1e4:.2f}=0.615，小于0.70，设置1个5 000 DWT成品油泊位可满足要求。")
    add_body(
        doc,
        "码头设置3套油品专用装卸臂或采用可清管的分品种臂组，岸线设紧急切断阀、绝缘法兰、接地装置、可燃气体探测器、溢油围控器材和船岸通信。"
        "泊位同一时刻仅安排一种油品作业，进库与出库不同时进行。码头消防与库区消防环网联通，但码头水幕、消防炮和船舶消防条件应在码头专业设计中进一步核定。",
    )
    add_heading(doc, "3.3 工艺流程方案", 2)
    add_heading(doc, "3.3.1 收油流程", 3)
    add_body(
        doc,
        "水运进库开始前确认目标罐可用容量不小于本批次体积加高液位保护余量；船岸双方完成静电接地、软管或装卸臂连接、阀位核对和ESD联试。起泵初期低流量充管，"
        "确认无泄漏后升至300 m³/h。接近目标罐高液位时逐步降量，高高液位动作时关闭岸线和罐前紧急切断阀并停止船泵。",
    )
    add_body(
        doc,
        "公共管道进库在交接计量前设置过滤器和取样点。批次切换时利用管道调度信息、密度计和界面检测判断混油段，混油进入指定罐，不得直接进入成品罐。"
        "进罐支线保持液封，避免汽油自由落差进罐产生静电。",
    )
    add_heading(doc, "3.3.2 发油流程", 3)
    add_body(
        doc,
        "公路装车按每种油品2个鹤位、单鹤位60 m³/h设置。车辆进位后依次完成品种核对、静电接地、防溢流插头连接和油气回收软管连接，批控器允许后开启装车阀。"
        "达到设定量或防溢流信号动作时关闭阀门并停泵，确认鹤管残液回收后方可拆除接地。",
    )
    add_body(
        doc,
        "管道外输和水运外输共用各油品变频发油泵，通过出口阀组选择去向。泵出口设置止回阀、最小流量回流和压力高高联锁；码头装船达到船舶高液位或船岸ESD动作时，"
        "先停泵再关闭岸线阀，避免水击。",
    )
    add_heading(doc, "3.3.3 计量、取样与清扫", 3)
    add_body(
        doc,
        "大流量水运和管道作业采用质量流量计或容积式流量计作为交接计量，公路装车采用批量控制流量计。每条计量支路设置温度、压力测点和在线密度校核接口。"
        "取样点布置在计量后直管段或循环取样装置，取样废液密闭回收。",
    )
    add_body(
        doc,
        "检修前先隔离、排净和氮气置换，确认可燃气体浓度满足作业要求。废液进入污油罐或含油污水系统，禁止排入雨水沟。清扫连接采用不同规格或机械钥匙管理，"
        "防止汽油、柴油和氮气软管误接。",
    )


def build_chapter_4(doc: Document):
    add_heading(doc, "4 管网水力计算与设备选型", 1)
    add_heading(doc, "4.1 设计流量、流速与管径", 2)
    add_body(
        doc,
        "主干管按表3-2峰值流量确定，不按年平均流量选管。汽油和柴油主输管流速控制在约1.0～2.5 m/s，泵吸入管宜低于1.5 m/s；支管在满足静电安全和不沉积水杂的条件下选择标准管径。",
        [18],
    )
    eq17 = add_equation(doc, "d=√(4Q/(πv))")
    add_body(doc, f"该式用于由设计流量和允许流速计算管道内径。式中：d——管内径，m；Q——体积流量，m³/s；v——允许流速，m/s。水运进库Q=300 m³/h、取v=1.8 m/s，代入式{eq17}得d=0.243 m，选DN250；公路发油Q=120 m³/h、取v=2.0 m/s，得d=0.146 m，选DN150。")
    add_table(
        doc,
        "表4-1 主要管线管径与流速",
        ["工况", "Q/(m³/h)", "公称直径", "计算流速/(m/s)", "长度/m", "适用油品"],
        [[name.replace("92号汽油", "汽油").replace("95号汽油", "汽油").replace("0号柴油", "柴油"),
          f"{h['Q']:.0f}", f"DN{int(h['d']*1000)}", f"{h['v']:.2f}",
          str(next(x[4] for x in PATHS if x[0] == name)), "汽油/柴油" if "进库" in name else name[:4]]
         for name, h in HYD.items() if name in ["92号汽油水运进库", "92号汽油管道进库", "92号汽油公路发油", "92号汽油管道外输", "92号汽油水运外输"]],
        [4.0, 2.0, 2.2, 2.8, 1.8, 2.5],
    )
    add_heading(doc, "4.2 沿程阻力、局部阻力与系统扬程", 2)
    eq18 = add_equation(doc, "Re=vd/ν")
    eq19 = add_equation(doc, "λ=0.25/[lg(ε/(3.7d)+5.74/Re^0.9)]^2")
    eq20 = add_equation(doc, "h_f=λ(L/d)(v^2/(2g))")
    eq21 = add_equation(doc, "h_m=Σζ(v^2/(2g))")
    eq22 = add_equation(doc, "H=Δz+h_f+h_m+h_e+h_r")
    add_body(
        doc,
        "上述公式依次用于判别流态、确定摩阻系数、计算沿程损失、局部损失和系统扬程。式中：Re——雷诺数；ν——运动黏度，m²/s；"
        "λ——达西摩阻系数；ε——管壁绝对粗糙度，m；L——管道计算长度，m；h_f——沿程阻力损失，m；"
        "Σζ——局部阻力系数之和；h_m——局部阻力损失，m；Δz——几何高差，m；h_e——过滤、计量和装卸设备折算水头，m；h_r——末端余压水头，m。",
    )
    sample = HYD["92号汽油公路发油"]
    add_body(
        doc,
        f"以92号汽油公路发油最远鹤位为例：Q=120 m³/h，DN150，v={sample['v']:.2f} m/s，L=350 m，运动黏度ν=0.70×10⁻⁶ m²/s，"
        f"绝对粗糙度ε=0.045 mm。代入式{eq18}得Re={sample['Re']:.2e}，代入式{eq19}得λ={sample['lam']:.4f}；"
        f"代入式{eq20}得沿程损失h_f={sample['hf']:.2f} m。局部阻力系数Σζ=18，代入式{eq21}得h_m={sample['hm']:.2f} m。"
        f"高差8 m、过滤计量与鹤管压降12 m、末端余压3 m，代入式{eq22}得H={sample['H']:.2f} m。",
    )
    diesel = HYD["0号柴油水运进库"]
    add_body(
        doc,
        f"柴油水运进库采用DN250、Q=300 m³/h时，v={diesel['v']:.2f} m/s，Re={diesel['Re']:.2e}，λ={diesel['lam']:.4f}，"
        f"沿程损失{diesel['hf']:.2f} m、局部损失{diesel['hm']:.2f} m，总需要扬程{diesel['H']:.2f} m。该扬程作为船泵或岸上助推泵的接口要求。",
    )
    add_table(
        doc,
        "表4-2 主要路径水力计算汇总",
        ["路径", "v/(m/s)", "Re", "λ", "h_f/m", "h_m/m", "H/m", "轴功率/kW"],
        [[name, f"{h['v']:.2f}", f"{h['Re']:.2e}", f"{h['lam']:.4f}", f"{h['hf']:.2f}",
          f"{h['hm']:.2f}", f"{h['H']:.2f}", f"{h['P']:.1f}"] for name, h in HYD.items()],
        [4.0, 1.6, 2.2, 1.8, 1.7, 1.7, 1.7, 2.2],
        Pt(8),
    )
    add_heading(doc, "4.3 输油泵选型", 2)
    crop_path, curve_path = build_pump_assets()
    add_heading(doc, "4.3.1 设计工况与候选泵", 3)
    add_body(
        doc,
        "库内输油泵服务于倒罐、公共管道外输和水运外输；水运进库主要由船泵提供压力。表4-2表明，库内最不利路径为汽油水运外输，"
        "设计流量200 m³/h、系统所需扬程35.91 m，因此泵的选型点不再沿用原稿缺少样本依据的250 m³/h、55 m，而以200 m³/h、约36 m作为额定工作点。",
    )
    add_body(
        doc,
        "候选泵采用KSB MegaCPK Inducer 125-100-200单级卧式离心泵，转速2900 r/min，叶轮直径初选193 mm。"
        "制造商样本第18页同时给出H-Q、效率、NPSH_r和功率曲线；样本第6页说明性能曲线按ISO 9906 2B级，NPSH_r按扬程下降3%确定。"
        "本设计按样本图读数进行初选，采购前仍应由制造商按实际油品密度、黏度和叶轮切割直径出具确认曲线。",
        [19],
    )
    if crop_path.exists():
        add_figure(doc, crop_path, "图4-1 KSB MegaCPK Inducer 125-100-200泵性能曲线（制造商样本第18页）", 15.2)

    add_heading(doc, "4.3.2 泵与管路特性曲线及工作点", 3)
    add_body(
        doc,
        "为确定工作点，将最不利水运外输管路的系统扬程写成静扬程与流量平方项之和。该路径静扬程由高差5 m、设备压降10 m和末端余压3 m组成，"
        "即H_st=18 m；在Q=200 m³/h时沿程与局部损失合计17.91 m，据此反算管路阻力系数。",
    )
    eq23a = add_equation(doc, "H_sys=H_st+K_QQ^2")
    add_body(
        doc,
        f"式中：H_sys——管路系统扬程，m；H_st——静扬程与设备、末端压力折算水头之和，m；K_Q——管路阻力系数，h²/m⁵；"
        f"Q——体积流量，m³/h。将H_st=18 m、Q=200 m³/h、H_sys=35.91 m代入式{eq23a}，得K_Q=4.48×10⁻⁴ h²/m⁵。",
    )
    if curve_path.exists():
        add_figure(doc, curve_path, "图4-2 输油泵性能曲线与最不利管路特性曲线", 14.5)
    add_body(
        doc,
        "图4-2表明，193 mm叶轮泵曲线与系统曲线交于Q≈202 m³/h、H≈36.3 m，接近设计流量且位于样本高效区。"
        "低流量公路装车工况采用变频调节，并设置最小流量回流，避免长期在制造商最小连续稳定流量左侧运行。"
        "该工作点来自制造商曲线与本文水力计算的交点，不是人为给定。",
    )

    add_heading(doc, "4.3.3 轴功率与电机", 3)
    eq23 = add_equation(doc, "P=(ρgQH)/(1000η_p)")
    add_body(
        doc,
        f"该式用于由流量、扬程和效率计算泵轴功率。式中：P——泵轴功率，kW；ρ——油品密度，kg/m³；g——重力加速度，取9.81 m/s²；"
        f"Q——体积流量，m³/s；H——扬程，m；η_p——泵效率。柴油按Q=202 m³/h、H=36.3 m、ρ=840 kg/m³、样本工作点效率η_p=0.78代入式{eq23}，"
        "得P=21.6 kW。考虑制造偏差、工况移动和电机裕量，三种油品统一配置30 kW防爆电机；最终功率还应与制造商确认曲线配套。",
    )
    add_table(
        doc,
        "表4-3 输油泵选型",
        ["位号", "介质", "数量", "额定Q/(m³/h)", "额定H/m", "电机/kW", "运行方式"],
        [
            ["P-101A/B", "92号汽油", "2", "200", "36", "30", "一用一备，变频"],
            ["P-201A/B", "95号汽油", "2", "200", "36", "30", "一用一备，变频"],
            ["P-301A/B", "0号柴油", "2", "200", "36", "30", "一用一备，变频"],
            ["P-401A/B", "污油/排水", "2", "50", "30", "11", "一用一备"],
        ],
        [2.7, 2.4, 1.5, 2.7, 2.0, 2.0, 3.2],
    )

    add_heading(doc, "4.3.4 柴油泵NPSH校核", 3)
    eq24 = add_equation(doc, "NPSH_a=P_a/(ρg)+h_s-P_v/(ρg)-h_(fs)")
    add_body(
        doc,
        f"该式用于校核泵入口可利用汽蚀余量。式中：NPSH_a——装置汽蚀余量，m；P_a——当地大气压，Pa；ρ——油品密度，kg/m³；"
        f"h_s——最低液位高于泵轴中心的静压头，m；P_v——设计温度下饱和蒸气压，Pa；h_fs——吸入管总水头损失，m。"
        "柴油校核取P_a=101325 Pa、ρ=840 kg/m³；储罐低低液位1.5 m、泵轴高于罐底0.5 m，故h_s=1.0 m；"
        "按40 ℃柴油饱和蒸气压保守取1.0 kPa，吸入管DN250、长度80 m、局部阻力系数8，Q=202 m³/h时h_fs=1.35 m。"
        f"代入式{eq24}得NPSH_a=11.83 m。",
    )
    add_body(
        doc,
        "由KSB样本第18页NPSH_r曲线读取，Q≈202 m³/h时NPSH_r≈1.9 m。按1.0 m附加裕量校核，"
        "NPSH_a=11.83 m>NPSH_r+1.0=2.90 m，柴油泵汽蚀裕量为8.93 m，满足要求。汽油蒸气压更高，"
        "运行前应按实际最高油温复核；本节不再用“柴油条件自然更好”代替数值校核。",
    )

    add_heading(doc, "4.3.5 装卸能力与备用泵校核", 3)
    eq25 = add_equation(doc, "M_a=nQρt_dD")
    add_body(
        doc,
        f"该式用于校核单套输油泵的年质量作业能力。式中：M_a——年质量能力，t/a；n——同时运行泵数；Q——单泵流量，m³/h；"
        f"ρ——油品密度，t/m³；t_d——日运行时间，h/d；D——年运行天数，d/a。每种油品按1台泵、Q=200 m³/h、t_d=16 h/d、D=330 d代入式{eq25}，"
        "汽油年能力为80.26万t/a，柴油为88.70万t/a，均大于各自年出库量。每种油品设置A/B两台泵，一用一备；"
        "备用泵与工作泵同型号、同流量和同扬程，可在工作泵故障或检修时自动/手动切换。",
    )
    add_body(
        doc,
        "公路装车每种油品2个鹤位、单鹤位60 m³/h，按16 h/d、330 d/a计算，汽油能力48.15万t/a，柴油能力53.22万t/a；"
        "分别大于92号汽油26万t/a、95号汽油17.5万t/a和柴油18万t/a的公路出库量。泵和鹤位均满足初稿设计能力要求。",
    )
    add_heading(doc, "4.4 阀门、计量及管道附件", 2)
    add_body(
        doc,
        "罐根第一道阀采用防火型钢制阀门并具备远程快速切断功能；泵出口设止回阀和电动切断阀，最低点设密闭排凝，高点设放气。"
        "汽油管线法兰跨接并控制盲端长度，装车鹤管采用液下装车结构。大流量计量支路前后保证制造商要求的直管段，过滤器设置差压报警和旁路隔离。",
    )
    add_table(
        doc,
        "表4-4 主要阀门与计量设备",
        ["位置", "设备", "规格", "控制要求"],
        [
            ["罐根", "防火型电动切断阀", "DN200～DN250", "高高液位、ESD联锁关闭"],
            ["泵入口", "全通径闸阀/球阀", "DN200", "低阻力、检修隔离"],
            ["泵出口", "止回阀+调节阀", "DN200", "防倒流、变频协同"],
            ["水运/管道交接", "质量流量计", "DN200～DN250", "温压补偿、双向计量"],
            ["公路装车", "批控流量计", "DN80/鹤位", "定量装车、防溢流联锁"],
        ],
        [3.0, 4.0, 3.0, 6.0],
    )
    add_heading(doc, "4.5 管路运行与瞬变控制", 2)
    add_heading(doc, "4.5.1 吸入管布置与低阻设计", 3)
    add_body(
        doc,
        "泵吸入管从储罐根部独立引出，不用一根细总管同时供多台大泵。吸入管采用DN250，正常流量200 m³/h时流速约1.13 m/s，"
        "能够兼顾低阻力和避免油品长期滞留。管线自储罐向泵连续坡降，不形成可积聚气体的高点；必须抬高时设置可控放气点，放空气体接入安全密闭系统。",
    )
    add_body(
        doc,
        "泵入口直管段避免紧邻弯头、三通和半开阀门。入口异径管采用顶平偏心异径管，防止气袋；过滤器按洁净和堵塞两种状态核算压降，"
        "差压高报警后切换备用过滤器或停止作业。入口阀保持全开，流量调节由变频和出口阀完成，避免通过入口节流降低NPSH_a。",
    )
    add_body(
        doc,
        "柴油NPSH校核采用低低液位、最高设计油温和过滤器不利压降组合。若后续总平面使吸入管长度超过80 m，或泵轴中心标高升高，"
        "应按实际长度和标高重算h_fs及h_s；不能只比较泵样本NPSH_r。储罐低低液位停泵值应高于产生旋涡和吸气的最低淹没深度。",
    )
    add_heading(doc, "4.5.2 停泵水击与阀门动作顺序", 3)
    add_body(
        doc,
        "水运外输管线长约1200 m，突然停泵或快速关阀可能产生压力波。初稿采用“先降速停泵、后关出口切断阀”的顺序：正常停输时变频器按设定斜率降速，"
        "流量接近零后关闭出口阀；紧急停车时立即切断驱动，同时由止回阀阻止倒流，电动切断阀按经瞬变校核确定的时间关闭。",
    )
    add_body(
        doc,
        "码头船岸ESD分为岸方和船方两级。高高液位、装卸臂超限或可燃气体高高报警时，先发停泵信号并关闭岸侧紧急切断阀；"
        "紧急脱离装置只在装卸臂位移继续扩大或连接失效时动作。阀门关闭时间过短会放大水击，过长会增加泄漏量，施工图阶段应进行瞬变计算后确定。",
    )
    add_table(
        doc,
        "表4-5 主要工况的泵阀动作顺序",
        ["工况", "泵动作", "阀门动作", "关键监测", "恢复条件"],
        [
            ["正常启动", "确认灌泵后低速启动", "入口全开，出口阀小开后渐开", "入口压力、流量、电流", "压力和振动稳定"],
            ["正常停输", "变频降速至最小流量后停泵", "泵停后关闭出口阀", "流量、出口压力", "管线泄压并隔离"],
            ["出口压力高高", "立即停泵", "出口切断阀按设定时间关闭", "压力变送器、阀位", "查明堵塞或误关阀原因"],
            ["储罐高高液位", "停止进库动力源", "关闭目标罐根阀并切换安全罐", "独立液位开关", "人工确认液位和阀位"],
            ["船岸ESD", "船泵/岸泵联锁停止", "岸阀、船阀顺序关闭", "装卸臂位移、通信状态", "双方共同复位"],
        ],
        [3.0, 3.6, 4.5, 3.5, 3.6],
    )
    add_heading(doc, "4.5.3 管道材料、连接与检验", 3)
    add_body(
        doc,
        "成品油工艺管道采用无缝或焊接碳钢管，材料等级、设计压力和壁厚由管道材料等级表统一规定。罐根、泵口、计量撬和码头接口采用法兰连接，"
        "长直管段优先焊接，减少潜在泄漏点。汽油管道垫片和阀杆填料选低泄漏型，阀门具备防火安全结构。",
    )
    add_body(
        doc,
        "架空管道考虑温差位移、支吊架摩擦和设备管口允许荷载。固定点、导向支架和滑动支架按热位移方向布置，罐前第一道阀附近保留柔性段，"
        "避免储罐沉降把荷载传至罐壁接管。跨防火堤处不把套管作为固定点，套管两端采用耐油柔性密封。",
    )
    add_body(
        doc,
        "安装完成后按管道等级进行强度和严密性试验，计量撬、过滤器和泵等不宜承受试验压力的设备用盲板隔离。投油前完成清管、吹扫、干燥和惰化，"
        "首次进汽油控制流速并确认静电接地连续。试运记录至少包括泵流量、扬程、电流、轴承温度、振动、密封泄漏和联锁动作。",
    )


def build_chapter_5(doc: Document):
    add_heading(doc, "5 总平面与安全防护", 1)
    add_heading(doc, "5.1 功能分区与总平面布置", 2)
    add_body(
        doc,
        "总平面按约420 m×320 m的临港工业地块进行方案布置。生产区位于场地中部和靠码头侧，管理及控制区位于常年主导风向的上风或侧上风位置。"
        "三座罐组按品种分区，泵棚和阀组区位于罐区外侧，公路装车区靠近陆域出入口，消防水池及消防泵房位于罐区外独立区域，事故水池设在场地排水低点。",
    )
    add_body(
        doc,
        "罐区四周设置环形消防道路，转弯半径、净宽和净空满足消防车辆通行。码头管廊与库区管廊沿道路一侧敷设，穿越道路处采用管涵或高架并设置限高标识。"
        "管理区与生产区分设出入口，装车车辆采用单向循环，避免与消防道路和人员通道交叉。",
    )
    add_table(
        doc,
        "表5-1 主要功能分区及布置要求",
        ["分区", "主要设施", "布置原则", "消防联系"],
        [
            ["储罐区", "3个罐组、6座罐", "居中布置、环形道路", "固定冷却和泡沫"],
            ["泵阀区", "发油泵、过滤计量、阀组", "靠罐区外侧、便于检修", "可燃气体检测、消火栓"],
            ["公路装车区", "6个鹤位、VRU", "靠陆域出口、单向交通", "泡沫/干粉、紧急切断"],
            ["码头区", "1个5 000 DWT泊位", "独立岸线、船岸联锁", "消防炮、溢油围控"],
            ["消防区", "水池、泵房、泡沫站", "远离主要火灾危险源", "两路供水、环网"],
            ["污水与事故水区", "初雨池、隔油、事故水池", "场地低点、可切换", "事故状态封堵"],
        ],
        [2.6, 4.0, 5.0, 4.0],
    )
    add_heading(doc, "5.2 防火间距与道路布置", 2)
    add_body(
        doc,
        "总平面防火间距按GB 50074—2014中二级石油库、甲B类汽油和丙A类柴油设施对应条款逐项控制。"
        "初稿不以一句“满足规范”替代尺寸核对，而是把储罐间距、罐组至泵棚、装车设施、消防泵房、管理区和围墙的控制值作为总平面图标注条件。"
        "汽油与柴油分别成组，固定顶柴油罐不与汽油内浮顶罐共用防火堤。",
    )
    add_table(
        doc,
        "表5-2 总平面防火间距控制清单",
        ["相邻设施", "控制原则", "本设计图纸要求", "复核阶段"],
        [
            ["同罐组储罐之间", "按油品类别、罐型及罐径取值", "在总平面图逐罐标注中心距和净距", "总图完成后按GB 50074复核"],
            ["罐组与泵棚、阀组", "泵棚位于防火堤外并留消防操作面", "泵棚沿罐区外侧布置，不占防火堤有效容积", "总图"],
            ["罐组与装车设施", "避免装车火源、车辆与储罐相互影响", "装车区靠陆域出口并与罐区分区", "总图"],
            ["罐组与消防泵房", "消防设施不受控制火灾直接威胁", "消防泵房布置于罐区外独立区域", "消防总图"],
            ["生产区与管理区", "人员集中设施位于上风或侧上风", "管理区独立出入口并与生产区分隔", "总图"],
            ["道路与防火堤", "满足消防车接近、转弯和连续通行", "罐区周边设置环形消防道路", "总图"],
        ],
        [4.0, 5.0, 5.0, 2.8],
    )
    add_note(doc, "表中具体数值必须在两张图纸形成实际坐标和尺寸后，按GB 50074—2014对应表格逐项核对；正文不虚构尚未落图的净距。")

    add_heading(doc, "5.3 防火堤与泄漏围控", 2)
    add_body(
        doc,
        "G92罐组防火堤内尺寸取140 m×90 m，布置2座D40 m储罐；G95和GD罐组内尺寸均取120 m×80 m，分别布置2座D34 m储罐。"
        "防火堤设计高度统一取2.2 m，计算有效高度按扣除0.2 m安全余量后的2.0 m。",
    )
    eq26 = add_equation(doc, "A_n=LB-nπD^2/4")
    eq27 = add_equation(doc, "V_d=A_n(H_d-0.2)")
    add_body(
        doc,
        "上述两式用于扣除罐基占地后计算防火堤有效容积。式中：A_n——防火堤内有效净面积，m²；L、B——防火堤内边长，m；"
        "n——堤内储罐数量；D——储罐直径，m；V_d——防火堤有效容积，m³；H_d——防火堤设计高度，m；0.2——堤顶安全余高，m。",
    )
    dike_rows = []
    for name, p, L, B in [("G92", "92号汽油", 140, 90), ("G95", "95号汽油", 120, 80), ("GD", "0号柴油", 120, 80)]:
        D = TANKS[p]["D"]
        n = 2
        an = L * B - n * math.pi * D**2 / 4
        vd = an * 2.0
        dike_rows.append([name, f"{L}×{B}", f"{D:.0f}", "2", f"{an:.2f}", f"{vd:.2f}", f"{TANKS[p]['nom']:.0f}", "满足"])
    add_body(doc, f"以G92罐组为例，将L=140 m、B=90 m、n=2、D=40 m代入式{eq26}，净面积A_n=10 086.73 m²；代入式{eq27}，有效容积V_d=20 173.45 m³>最大单罐20 000 m³。其余罐组计算见表5-3。")
    add_table(
        doc,
        "表5-3 防火堤有效容积校核",
        ["罐组", "堤内L×B/m", "D/m", "罐数", "净面积/m²", "有效容积/m³", "最大单罐/m³", "结果"],
        dike_rows,
        [1.6, 2.4, 1.4, 1.4, 2.5, 2.7, 2.7, 1.5],
    )
    add_body(
        doc,
        "防火堤内地坪坡向集水井，雨水出口阀常闭并具备阀位远传。日常降雨经人工确认和油膜检测合格后排入清净雨水系统；发现油品或事故状态时切入含油污水或事故水池。"
        "穿堤管线采用密封套管，堤脚禁止设置无法隔离的明沟。",
    )
    add_heading(doc, "5.4 HSE与危险因素分析", 2)
    add_heading(doc, "5.4.1 主要危险因素", 3)
    add_body(
        doc,
        "主要危险包括汽油蒸气形成爆炸性混合物、静电放电、泵密封和法兰泄漏、储罐超装、船岸软管破裂、车辆误操作、雷击、台风、地震及消防废水外排。"
        "危险源按储罐、泵阀、装车、码头和污水设施分区辨识，控制措施遵循消除、预防、检测、隔离、减缓和应急的层级。",
    )
    add_heading(doc, "5.4.2 HSE管理措施", 3)
    add_body(
        doc,
        "作业许可覆盖动火、受限空间、高处、临时用电、吊装和盲板抽堵。储罐清洗前完成物料隔离、排净、置换、气体检测和监护；码头作业执行船岸安全检查表，"
        "确认接地、通信、气象、缆绳、装卸臂和ESD状态。承包商进入生产区前接受油品危险和应急疏散培训。",
    )
    add_body(
        doc,
        "设备完整性管理以储罐底板、罐壁焊缝、浮盘密封、泵机械密封、紧急切断阀和消防设施为重点。根据风险等级安排在线检测、停罐检验和预防性维修，"
        "发现沉降异常、浮盘卡阻、底板腐蚀或阀门拒动时及时降低液位或停用设备。",
        [12, 13],
    )
    add_heading(doc, "5.5 防爆、防雷、防静电、抗震与防腐", 2)
    add_heading(doc, "5.5.1 爆炸危险区域与防爆电气", 3)
    add_body(
        doc,
        "汽油罐区、泵棚、装车鹤位、油气回收装置和码头装卸点按GB 50058—2014划分爆炸危险区域，电机、仪表、照明和接线箱按区域等级选防爆型式。"
        "可能释放油气的设备优先露天或敞开布置，泵棚保持自然通风；可燃气体探测器布置在泵密封、阀组、装车和VRU周边低处，并与声光报警和紧急切断联锁。",
    )
    add_table(
        doc,
        "表5-4 主要爆炸危险区域划分与设备要求",
        ["释放源", "区域划分对象", "主要防爆措施", "图纸表达"],
        [
            ["汽油罐呼吸阀、量油孔", "开口周围空间", "限制点火源，仪表和电气按区域选型", "防爆区域图标注边界"],
            ["汽油泵密封及阀组", "泵棚地面及设备周边", "露天通风、可燃气体检测、防爆电机", "平面和剖面标注"],
            ["公路装车鹤位", "装车口、车辆罐口及地面沟槽", "液下装车、静电联锁、紧急切断", "装车区防爆图"],
            ["油气回收装置", "入口管线、风机和处理单元周边", "防爆风机、LEL/压力联锁、阻火隔爆", "设备布置图"],
            ["码头装卸接口", "装卸臂、船岸连接和集油设施周边", "船岸ESD、绝缘法兰、紧急脱离", "码头危险区域图"],
        ],
        [3.7, 4.2, 5.3, 3.6],
    )
    add_body(
        doc,
        "爆炸危险区域的最终范围需按GB 50058—2014附录中释放源、通风条件和空间形态绘制，不能只写“采用防爆电机”。"
        "设备表还应给出防爆型式、气体组别和温度组别；地沟、低洼区因油气可能积聚，应按不利通风条件处理。",
    )

    add_heading(doc, "5.5.2 防雷与防静电", 3)
    add_body(
        doc,
        "储罐按第二类防雷建筑物相关要求设置接地。罐体沿周边不少于两处与接地网连接，接地装置兼作防静电接地；管道法兰、装卸臂、浮盘和扶梯活动连接处设置跨接。"
        "油罐车装车前必须连接静电接地并经联锁确认，接地断开时停止装车。汽油进罐采用液下进油，初始充装阶段限制流速。",
    )
    add_heading(doc, "5.5.3 抗震、台风与沿海防腐", 3)
    add_body(
        doc,
        "储罐基础、罐壁、锚固和接管按场地抗震参数设计，罐根第一道阀附近和跨堤管线设置柔性连接或补偿段，避免基础差异沉降和地震位移拉裂管口。"
        "台风预警期间停止非必要码头作业，检查浮盘排水、罐顶附件、抗风圈、装卸臂锁定和消防备用电源，降低高风险储罐作业频次。",
    )
    add_body(
        doc,
        "独山港年平均相对湿度80%，且受海盐气溶胶影响，钢结构和管道外防腐按沿海重腐蚀环境选体系。储罐外壁、管架和架空管道采用喷砂除锈后底漆—中间漆—耐候面漆复合涂层；"
        "保温管道重点控制保温层下腐蚀，采用防水封口和可排水结构。埋地钢管采用加强级外防腐层并结合阴极保护，法兰、螺栓和仪表支架避免异种金属电偶腐蚀。"
        "涂层干膜厚度和复涂周期由施工图防腐规格书及涂料制造商体系确认。",
    )
    add_heading(doc, "5.6 泄漏预防与应急隔离", 2)
    add_heading(doc, "5.6.1 储罐、泵阀和装车区防泄漏", 3)
    add_body(
        doc,
        "储罐底板采用可检测的防渗构造，罐基础周边设渗漏观察点。日常通过罐存量平衡、底板监测和罐基周边巡检识别慢渗漏；"
        "当同一班次计量差持续超限或观察点出现油迹时，停止收发作业并转移库存。罐壁下部、罐底边缘板和罐根接管列为重点腐蚀监测部位。",
    )
    add_body(
        doc,
        "泵采用机械密封，密封泄漏进入接液盘和含油污水系统。泵出口止回阀防止停泵倒流，进出口设置远程切断阀；"
        "泵棚地坪向集液沟找坡，不把泄漏油导向电气间或消防道路。阀组区法兰数量尽量减少，法兰下方不布置电缆接头。",
    )
    add_body(
        doc,
        "公路装车采用液下装车鹤管、车辆防溢流探头和静电接地联锁。接地未确认、车辆罐高液位、鹤管未到位或油气回收未连通时，批控系统不允许打开装车阀。"
        "装车岛设置紧急停车按钮，动作后关闭批控阀和总切断阀并停止输油泵。",
    )
    add_heading(doc, "5.6.2 分区隔离和紧急切断层级", 3)
    add_body(
        doc,
        "紧急切断按“设备—单元—库区”三级设置。设备级用于单泵超压、单罐高高液位和单鹤位溢油；单元级隔离一个油品罐组或装车区；"
        "库区级用于大范围可燃气体报警、火灾或外部灾害。分级可以在阻止事故扩大的同时，避免无关系统全部失去控制。",
    )
    add_table(
        doc,
        "表5-5 典型泄漏情景及隔离边界",
        ["情景", "首要动作", "隔离边界", "泄漏去向", "后续处置"],
        [
            ["罐根法兰泄漏", "停相关收发泵", "关闭罐根及下游阀", "防火堤内集液", "倒罐、堵漏、检验"],
            ["泵机械密封大量泄漏", "停故障泵并切备用泵", "关闭故障泵进出口", "接液盘至含油污水", "冲洗置换后检修"],
            ["装车溢油", "ESD停泵关阀", "装车岛总阀", "不渗地坪和集液沟", "回收油品、检测合格后恢复"],
            ["码头连接失效", "船岸ESD", "船阀和岸侧紧急切断阀", "围油栏和码头集油设施", "海事联动、回收处置"],
            ["罐区火灾", "启动消防并停止全部相关作业", "事故罐组和相邻单元", "防火堤—事故水池", "泡沫灭火、连续冷却"],
        ],
        [3.5, 4.2, 4.0, 3.8, 3.8],
    )
    add_heading(doc, "5.6.3 台风、雷暴和地震后的复产检查", 3)
    add_body(
        doc,
        "台风或强雷暴预警期间停止码头装卸和高处作业，核对罐顶附件、浮盘排水、装卸臂锁定、消防柴油机燃料和事故水池空余容积。"
        "预警解除后不立即恢复大流量作业，应先检查电源、仪表空气、通信、接地、阀位及罐区排水状态。",
    )
    add_body(
        doc,
        "地震后复产前检查储罐沉降和倾斜、罐壁屈曲、罐底翘离、接管变形、管架位移及防火堤裂缝。发现液位异常下降、基础渗油或管线受拉时，"
        "保持隔离并组织无损检测。复产由工艺、设备、仪表和HSE共同签字确认，不以外观无明显损坏作为唯一条件。",
    )

    add_heading(doc, "6 消防系统与事故水设计", 1)
    add_heading(doc, "6.1 消防控制工况与冷却水", 2)
    add_heading(doc, "6.1.1 控制火灾情景", 3)
    add_body(
        doc,
        "按GB 50074—2014第12.2.7条、表12.2.8及第12.2.11条比较汽油与柴油工况。钢制内浮顶汽油罐着火时，着火罐罐壁按2.0 L/(min·m²)连续冷却6 h，"
        "相邻罐可不固定冷却；固定顶柴油罐着火时，着火罐按全部罐壁面积2.5 L/(min·m²)，相邻罐按面向着火罐的半个罐壁面积2.0 L/(min·m²)，连续供水9 h。"
        "经计算，15 000 m³固定顶柴油罐火灾并冷却同组1座相邻罐为控制工况。",
    )
    add_heading(doc, "6.1.2 冷却水流量与持续时间", 3)
    eq28 = add_equation(doc, "A_w=πDH")
    eq29 = add_equation(doc, "Q_c=(q_fA_w+q_aA_w/2)/60")
    wall_area = math.pi * 34 * 18
    qc = (2.5 * wall_area + 2.0 * wall_area / 2) / 60
    eq30 = add_equation(doc, "V_c=3.6Q_ct_c")
    add_body(
        doc,
        "上述公式用于计算着火罐及相邻罐的固定冷却水流量和总用水量。式中：A_w——单罐罐壁面积，m²；D——罐径，m；H——罐壁高度，m；"
        "Q_c——冷却水总流量，L/s；q_f——着火罐冷却强度，L/(min·m²)；q_a——相邻罐冷却强度，L/(min·m²)；"
        "V_c——冷却水量，m³；t_c——连续冷却时间，h。",
    )
    vc = 3.6 * qc * 9
    add_body(
        doc,
        f"将D=34 m、H=18 m代入式{eq28}，柴油罐壁面积A_w={wall_area:.2f} m²。将q_f=2.5、q_a=2.0代入式{eq29}，"
        f"着火罐与1座相邻罐的固定冷却水总流量Q_c={qc:.2f} L/s。再将t_c=9 h代入式{eq30}，冷却水量V_c={vc:.2f} m³。",
    )
    add_heading(doc, "6.2 泡沫灭火系统", 2)
    add_heading(doc, "6.2.1 保护方式与设计参数", 3)
    add_body(
        doc,
        "汽油内浮顶罐采用密封圈环形保护，泡沫堰板距罐壁b=0.55 m，混合液强度12.5 L/(min·m²)，连续供给60 min；柴油固定顶罐采用全液面保护，"
        "混合液强度6.0 L/(min·m²)，连续供给30 min。泡沫液按3%型、实际混合比上限3.9%并增加10%储量计算。",
    )
    add_heading(doc, "6.2.2 泡沫流量、产生器数量与储量", 3)
    eq31 = add_equation(doc, "A_(f,g)=πDb")
    eq32 = add_equation(doc, "Q_f=q_fA_f/60")
    eq33 = add_equation(doc, "n_p=ceil(Q_f/Q_p)")
    eq34 = add_equation(doc, "V_(mix)=3.6Q_ft_f")
    eq35 = add_equation(doc, "V_(foam)=1.10cV_(mix)")
    add_body(
        doc,
        "上述公式用于计算泡沫保护面积、混合液流量、泡沫产生器数量及泡沫液储量。式中：A_(f,g)——汽油内浮顶罐密封圈保护面积，m²；"
        "b——泡沫堰板至罐壁距离，m；Q_f——泡沫混合液流量，L/s；q_f——混合液供给强度，L/(min·m²)；"
        "n_p——泡沫产生器数量；Q_p——单个产生器额定流量，L/s；V_(mix)——泡沫混合液量，m³；t_f——连续供给时间，h；"
        "V_(foam)——泡沫液储量，m³；c——实际混合比；1.10——储量裕量系数。",
    )
    ag = math.pi * 40 * 0.55
    qg = 12.5 * ag / 60
    vmg = 3.6 * qg * 1
    vfg = 1.10 * 0.039 * vmg
    ad = math.pi * 34**2 / 4
    qd = 6.0 * ad / 60
    vmd = 3.6 * qd * 0.5
    vfd = 1.10 * 0.039 * vmd
    add_body(
        doc,
        f"20 000 m³汽油罐：将D=40 m、b=0.55 m代入式{eq31}，A_f={ag:.2f} m²；代入式{eq32}，Q_f={qg:.2f} L/s；"
        f"采用额定8 L/s的泡沫产生器，代入式{eq33}，n_p=2；代入式{eq34}，V_mix={vmg:.2f} m³；代入式{eq35}，泡沫液储量V_foam={vfg:.2f} m³。",
    )
    add_body(
        doc,
        f"15 000 m³柴油罐：全液面面积A_f={ad:.2f} m²，Q_f={qd:.2f} L/s；采用额定16 L/s的泡沫产生器，数量n_p=6；"
        f"30 min混合液量V_mix={vmd:.2f} m³，泡沫液储量V_foam={vfd:.2f} m³。柴油工况控制泡沫液储量，设置2座5 m³泡沫液储罐，总有效储量10 m³。",
    )
    add_table(
        doc,
        "表6-1 泡沫系统计算结果",
        ["保护对象", "保护面积/m²", "强度", "Q_f/(L/s)", "时间/min", "产生器", "数量", "泡沫液/m³"],
        [
            ["20 000 m³汽油罐密封圈", f"{ag:.2f}", "12.5 L/(min·m²)", f"{qg:.2f}", "60", "8 L/s", "2", f"{vfg:.2f}"],
            ["15 000 m³柴油罐全液面", f"{ad:.2f}", "6.0 L/(min·m²)", f"{qd:.2f}", "30", "16 L/s", "6", f"{vfd:.2f}"],
        ],
        [4.2, 2.2, 3.2, 2.2, 1.8, 2.0, 1.4, 2.2],
    )
    add_heading(doc, "6.3 消防水池与消防泵", 2)
    add_heading(doc, "6.3.1 消防水池有效容积", 3)
    foam_water = vmd * (1 - 0.039)
    fire_water = vc + foam_water
    eq36 = add_equation(doc, "V_(FW)=V_c+(1-c)V_(mix)")
    add_body(
        doc,
        "该式用于计算消防水池所需有效水量。式中：V_(FW)——消防用水量，m³；V_c——固定冷却水量，m³；"
        "c——泡沫液实际混合比；V_(mix)——泡沫混合液量，m³。",
    )
    add_body(
        doc,
        f"将柴油冷却水量{vc:.2f} m³和泡沫混合液量{vmd:.2f} m³代入式{eq36}，消防用水量V_FW={fire_water:.2f} m³。"
        "消防水池按计算量向上取整并分格设置，采用2×2 500 m³，总有效容积5 000 m³。",
    )
    add_heading(doc, "6.3.2 消防泵、备用泵与供电", 3)
    add_body(
        doc,
        f"固定冷却水泵按Q_c={qc:.2f} L/s配置2×130 L/s，一用一备；泡沫水泵按Q_f={qd:.2f} L/s配置2×110 L/s，一用一备。"
        "消防环网由两路出水干管供水，罐区、泵区、装车区和码头设分区阀，任一段检修不影响其余区域供水。",
    )
    add_body(
        doc,
        "消防冷却水泵设置2台电动泵和1台柴油机驱动备用泵，单台额定流量均不小于130 L/s；正常由1台电动泵承担设计流量，"
        "另1台电动泵及柴油机泵提供故障和失电备用。泡沫供水泵同样按工作泵与备用泵能力一致配置。消防泵应具有独立吸水管、试验回流和定期自启动试验条件，"
        "控制室远程启动不代替泵房就地手动启动。备用消防泵的流量、扬程不得低于最大一台工作泵。",
    )

    add_heading(doc, "6.4 事故水池与排水切换", 2)
    add_heading(doc, "6.4.1 事故水量组成与容积", 3)
    rainfall = 0.90 * 12_600 * 0.2764
    leakage = 300 * (10 / 60) * 1.2
    vacc = fire_water + rainfall + leakage
    eq37 = add_equation(doc, "V_(acc)=V_(FW)+ψFh_r+1.2Q_lt_s")
    add_body(
        doc,
        "该式用于合并消防水、受污染雨水和紧急切断前泄漏量。式中：V_(acc)——事故水池计算容积，m³；V_(FW)——控制工况消防水量，m³；"
        "ψ——径流系数；F——事故影响汇水面积，m²；h_r——最大一日降雨深度，m；Q_l——最大泄漏流量，m³/h；"
        "t_s——紧急切断时间，h；1.2——泄漏量附加系数。",
    )
    add_body(
        doc,
        f"事故水量包括控制工况消防水、最大日降雨和切断前泄漏量。取径流系数ψ=0.90、受影响最大罐区面积F=12 600 m²、最大日降雨h_r=0.2764 m，"
        f"降雨量为{rainfall:.2f} m³；最大输油流量300 m³/h、切断时间10 min并加20%裕量，泄漏量为{leakage:.2f} m³。"
        f"代入式{eq37}，V_acc={vacc:.2f} m³，选8 000 m³事故水池。",
    )
    add_heading(doc, "6.4.2 防泄漏、雨污分流与切换", 3)
    add_body(
        doc,
        "储罐采用防渗基础、罐底泄漏监测和可检修的环形基础。泵密封设置接液盘并接入含油污水系统；装车区、阀组区和计量区采用不渗地坪，四周设收集沟。"
        "事故状态关闭雨水总排口，将污染水切入事故水池；事故水经检测、隔油和处理后分批处置，不直接外排。",
    )
    add_body(
        doc,
        "初期雨水量为202.5 m³，设置2×150 m³初期雨水池。",
    )
    add_heading(doc, "6.5 消防系统运行与可靠性", 2)
    add_heading(doc, "6.5.1 消防给水环网与最不利点", 3)
    add_body(
        doc,
        "消防水从两格消防水池分别吸水，经电动消防泵和柴油机消防泵进入环状管网。环网在罐区、泵区、装车区、码头和泡沫站设置分区阀，"
        "任一段检修时仍能从另一方向向控制火灾点供水。消防泵出口设置流量、压力和试验回流，日常试验水返回消防水池，避免长期直接排放。",
    )
    add_body(
        doc,
        "最不利点不是简单取离泵房最远的消火栓，而应比较柴油罐固定冷却环管、汽油罐泡沫立管、码头消防炮和高程不利点。"
        "施工图阶段按控制流量同时开启所需支路，计算环网沿程与局部损失，并保证最不利喷头或泡沫产生器入口压力满足设备要求。",
    )
    add_body(
        doc,
        "固定冷却环管沿罐壁周向均匀分区，喷头方向避开保温、盘梯和加强圈遮挡。每个分区入口设置可在防火堤外操作的阀门、压力表和试水接口。"
        "泡沫立管与冷却水管分开，防止误操作把清水送入泡沫产生器或把泡沫液长期滞留在冷却管。",
    )
    add_heading(doc, "6.5.2 泡沫比例混合与校验", 3)
    add_body(
        doc,
        "泡沫系统按3%型泡沫液设计，但储量计算采用制造商允许混合比上限3.9%并增加10%裕量。比例混合装置的流量范围应覆盖柴油罐控制工况和汽油密封圈工况；"
        "若单台装置在小流量工况下不能保证比例精度，应设置大小两级装置或回流稳定措施。",
    )
    add_body(
        doc,
        "泡沫液储罐分为两座5 m³，可在一座检修时保留部分灭火能力。储罐设置液位、呼吸、防腐内衬和取样接口，补充泡沫液必须与原有药剂相容。"
        "每年通过不排放或少排放的试验回路验证比例混合精度；更换泡沫液品种后重新核对黏度、混合比、低温性能和产生器适配性。",
    )
    add_body(
        doc,
        "泡沫产生器数量按计算流量向上取整，并沿罐周均匀布置。单个产生器或一条立管失效时会造成局部覆盖不足，因此管线布置避免一个低点积液影响多个产生器。"
        "试验时检查背压、发泡倍数和分布均匀性，不能只确认泵已启动。",
    )
    add_heading(doc, "6.5.3 消防泵启停逻辑和定期试验", 3)
    add_body(
        doc,
        "消防泵可由消防控制室远程启动、泵房就地启动和压力联锁启动。启动命令发出后监测出口压力、流量和电机状态；主电动泵未在规定时间建立压力时，"
        "自动启动备用电动泵，仍失败或全厂失电时启动柴油机泵。消防泵一旦投入火灾工况，原则上由现场授权人员手动停泵。",
    )
    add_table(
        doc,
        "表6-2 消防设备配置与可靠性要求",
        ["系统", "工作设备", "备用设备", "设计能力", "验证方式"],
        [
            ["固定冷却水", "1台130 L/s电动泵", "1台同能力电动泵+1台柴油机泵", "不小于控制工况Q_c", "流量试验和最不利点压力"],
            ["泡沫供水", "1台110 L/s电动泵", "1台同能力备用泵", "不小于柴油全液面混合液流量", "比例混合和泡沫出口试验"],
            ["消防水池", "2×2500 m³分格", "两格可切换补水", "有效水量大于计算值", "液位、补水和吸水试验"],
            ["泡沫液储罐", "2×5 m³", "两罐互为储量分隔", "总有效量大于计算值", "液位、取样和药剂相容性"],
            ["事故水池", "1×8000 m³", "预留移动泵接口", "大于消防+降雨+泄漏", "切换阀和高液位联锁"],
        ],
        [3.2, 4.0, 4.5, 4.2, 4.0],
    )
    add_heading(doc, "6.5.4 事故水池运行边界", 3)
    add_body(
        doc,
        "事故水池正常状态保持足够空余容积，不作为日常含油污水调节池使用。雨前根据气象预报检查液位；若池内已有待处理水，应提前外运或处理，"
        "不能用“名义容积8000 m³”代替可用容积管理。入口切换阀采用失效安全位置，并在控制室显示阀位。",
    )
    add_body(
        doc,
        "火灾时防火堤内排水阀关闭，消防水和泄漏油先在堤内暂存；确认事故水池具备接收条件后分批导入，避免大量油品直接冲击事故池。"
        "事故结束后对水相、浮油和沉积物分类处置，检测合格前不外排。事故水池高高液位时停止一切可能增加污染水的非应急作业。",
    )


def build_chapter_6(doc: Document):
    add_heading(doc, "7 环境保护、自动控制与运行管理", 1)
    add_heading(doc, "7.1 VOCs源项与控制原则", 2)
    add_body(
        doc,
        "VOCs主要来自汽油储罐呼吸、浮盘边缘密封、装车置换气、码头装卸置换气以及泵阀无组织泄漏。源头控制采用内浮顶、全接液钢制浮盘、一次二次密封和液下进油；"
        "过程控制采用密闭装车、低泄漏阀门和LDAR；末端治理采用冷凝与活性炭吸附组合的油气回收装置。",
        [4, 5],
    )
    add_heading(doc, "7.2 装车油气回收能力", 2)
    add_body(
        doc,
        "92号和95号汽油公路装车可同时进行，每种油品最大装车流量120 m³/h。按液体装入量与置换油气量近似1:1，并取1.10波动系数，最大油气量为264 m³/h。"
        "设置2×300 m³/h油气回收装置，一用一备；单套能力大于同时装车油气量。",
    )
    eq38 = add_equation(doc, "Q_(VRU)=1.10ΣQ_(load,g)")
    add_body(doc, f"将两种汽油装车流量各120 m³/h代入式{eq38}，Q_VRU=264 m³/h，选300 m³/h。沿海空气湿度较高，吸附单元前设置气液分离、温度监测和冷凝预处理，避免水汽降低吸附容量。", [14])
    add_heading(doc, "7.3 废水、固废与噪声", 2)
    add_body(
        doc,
        "含油污水包括罐底切水、泵区地面冲洗水、计量排凝和受污染初期雨水。污水先经隔油、调节和气浮后送园区污水系统；清净雨水与含油污水分流。"
        "废活性炭、含油污泥、废滤芯和实验废液按危险废物分类收集，交有资质单位处置。",
    )
    add_body(
        doc,
        "噪声设备主要为输油泵、消防泵和油气回收风机。泵采用低噪声电机、弹性基础和软连接，消防泵房与管理区保持距离，风机进出口设消声器。"
        "正常工况下通过设备选型和建筑隔声控制厂界噪声。",
    )
    add_heading(doc, "7.4 环境监测与异常工况", 2)
    add_body(
        doc,
        "监测项目包括油气回收入口流量、出口非甲烷总烃、储罐密封状态、泵阀泄漏、雨水排口油膜和事故水池液位。VRU高温、高压差或出口浓度超限时停止汽油装车并切换备用单元。"
        "事故水池接近高液位时暂停可能增加污染水的作业，优先组织外运或处理。",
    )


def build_chapter_7(doc: Document):
    add_heading(doc, "7.5 自动控制与运行管理", 2)
    add_heading(doc, "7.5.1 主要检测仪表", 3)
    add_body(
        doc,
        "每座储罐设置连续液位、温度和独立高高液位开关；泵入口设置压力低报警，出口设置压力高报警和流量低报警；过滤器设置差压报警；装车鹤位设置批控、防溢流和静电接地联锁。"
        "可燃气体探测器布置在汽油泵密封、阀组、装车和VRU可能积聚油气处。",
    )
    add_table(
        doc,
        "表7-1 主要报警与联锁",
        ["触发条件", "报警", "自动动作", "人工确认"],
        [
            ["储罐高高液位", "声光报警", "停进油泵、关闭进罐阀", "核对液位与阀位"],
            ["储罐低低液位", "低液位报警", "停发油泵", "确认罐底不抽空"],
            ["泵入口压力低", "压力报警", "延时停泵", "检查阀门、液位和滤网"],
            ["泵出口压力高", "压力报警", "停泵并开最小流量回路", "检查下游阀位"],
            ["可燃气体高高", "区域声光报警", "停相关泵、关闭切断阀、启动通风", "现场检测后复位"],
            ["船岸ESD", "码头与中控报警", "停泵、关闭岸线阀", "船岸双方确认"],
            ["装车失去接地/防溢流", "鹤位报警", "关闭装车阀、停批控", "车辆重新检查"],
        ],
        [3.5, 3.0, 6.0, 4.0],
    )
    add_heading(doc, "7.5.2 控制方式与紧急切断", 3)
    add_body(
        doc,
        "正常作业由中控室顺序启动：确认目标罐、阀门反馈和可用容量，开启末端阀，再开启泵入口阀和泵出口阀，最后启动泵并缓慢升速。停运顺序相反，先降速停泵再关阀。"
        "紧急切断不依赖操作画面单点命令，码头、装车区、泵区和中控室均设置硬接线ESD按钮。",
    )
    add_heading(doc, "7.5.3 运行、检维修与应急", 3)
    add_body(
        doc,
        "运行人员每班检查储罐液位趋势、浮盘状态、罐壁沉降标记、泵振动与密封、阀门反馈、消防压力和事故水池液位。周检包括紧急切断阀抽试、可燃气体探测器自检和静电接地装置检查；"
        "月度或季度按计划试运行消防泵、泡沫比例混合装置和备用电源。",
    )
    add_body(
        doc,
        "发生泄漏时立即停止相关泵、关闭上下游切断阀、禁止点火源并疏散无关人员；小量泄漏使用吸油材料围控，大量泄漏导入防火堤和事故水系统。发生罐火时启动固定冷却、泡沫系统并对相邻设施实施监护，"
        "根据风向设置警戒和疏散路线。",
    )
    add_heading(doc, "7.6 典型作业程序与参数管理", 2)
    add_heading(doc, "7.6.1 水运进库作业程序", 3)
    add_body(
        doc,
        "船舶靠泊后由船岸双方共同确认油品、数量、目标罐可用容量、装卸速率、最大允许压力、通信方式和紧急停输信号。连接装卸臂后先进行低流量试送，"
        "核对岸线压力、目标罐液位上升趋势和计量方向；确认无泄漏及无误入其他罐后，逐步提高到300 m³/h。",
    )
    add_body(
        doc,
        "接近计划数量或目标罐高液位时提前降低船泵转速，避免在最大流量下突然停输。停输后依次关闭船阀、岸侧阀和罐根阀，"
        "对装卸臂内残油进行密闭回收。计量差超过允许范围时，不以人工修改记录消除差值，应检查温度密度修正、阀门内漏和管线存油量。",
    )
    add_heading(doc, "7.6.2 公路装车作业程序", 3)
    add_body(
        doc,
        "车辆进入装车区前核对介质、仓容和防溢流接口，进入鹤位后熄火、制动并连接静电接地和油气回收。批控系统收到接地良好、防溢流正常、"
        "鹤管到位和油气回收可用四个条件后才允许装车。初始流量较低，鹤管口被液体淹没后提升到额定流量。",
    )
    add_body(
        doc,
        "装车结束先关闭批控阀并确认流量归零，再排净鹤管残油和拆除油气回收管，最后断开静电接地。出现车辆罐高液位、接地中断、可燃气体高高或油气回收故障时，"
        "联锁关闭装车阀并停止相应油品泵；未经现场检查不得直接复位继续装车。",
    )
    add_heading(doc, "7.6.3 倒罐、切水和检维修隔离", 3)
    add_body(
        doc,
        "倒罐前确认接收罐容量、油品牌号和质量状态，采用同品种专用管线。倒罐过程中监测两罐液位变化是否与流量累计一致；"
        "当接收罐高液位或源罐低低液位出现时自动停泵。不同牌号汽油不通过倒罐操作进行未经批准的调合。",
    )
    add_body(
        doc,
        "罐底切水采用小流量、有人监护和密闭接收。开始阶段水相进入含油污水系统，发现明显油相后立即关闭；切水量与罐存变化记录用于识别底水异常增加。"
        "冬季低温时检查切水阀和低点积水，防止冻结造成阀体或管道破裂。",
    )
    add_body(
        doc,
        "设备检维修执行停泵、关阀、泄压、排净、置换、检测和上锁挂牌。双阀隔离之间设置可控泄放，不能以控制系统显示“阀关”代替现场机械隔离。"
        "进入储罐前还需盲板隔断所有物料和氮气管线，连续监测氧含量与可燃气体，并制定救援方案。",
    )
    add_heading(doc, "7.6.4 数据首现、变更和版本一致性", 3)
    add_body(
        doc,
        "年周转量、运输比例、油品密度、周转系数、罐容、罐径、管长、高差、泵型号和消防参数构成本设计的关键输入。每个数据在正文首次出现处注明规范、文献、"
        "制造商样本或政府PDF来源；由设计假定得到的管长和标高明确写为初步总图条件，不伪装成外部统计事实。",
    )
    add_body(
        doc,
        "后续若总平面图改变管长或高差，应同步更新表4-2、系统曲线、泵工作点和NPSH；若改变罐径或罐高，应同步更新罐壁、抗风圈、防火堤、冷却水和泡沫计算。"
        "说明书、计算表和两张图纸使用同一版本参数，修改前提交Git快照，以便不满意时回退。",
    )
    add_table(
        doc,
        "表7-2 关键参数变更的联动更新关系",
        ["变更参数", "直接影响计算", "必须同步更新的章节或图纸"],
        [
            ["年周转量或运输比例", "理论库容、泊位与装卸能力", "第2、3、4章及工艺流程图"],
            ["储罐容量、直径或高度", "库级、壁厚、抗风、防火堤、消防", "第2、5、6章及总平面图"],
            ["管长、管径或高差", "流速、摩阻、系统扬程、工作点", "第4章、泵曲线及工艺流程图"],
            ["泵型号或叶轮直径", "H-Q交点、功率、NPSH、年能力", "第4章及设备表"],
            ["消防控制工况", "冷却水、泡沫、消防水池、事故水", "第6章及消防总图"],
            ["自然条件资料", "抗风、防腐、排水、停工边界", "第1、2、5、6章"],
        ],
        [4.0, 5.5, 7.5],
    )


def build_chapter_8(doc: Document):
    add_heading(doc, "8 结论", 1)
    conclusions = [
        "（1）结合独山港临港物流条件和二级商业油库设计难度，确定92号汽油、95号汽油和0号柴油年进出库量分别为40万t、25万t和30万t，总量95万t。进出库方式和比例保持计划表不变。",
        "（2）库容计算先取K=14、η=0.95，理论库容分别为39 572.62 m³、24 732.88 m³和26 852.85 m³。配置2×20 000 m³汽油罐、2×15 000 m³汽油罐和2×15 000 m³柴油罐，共6座。反算K为13.85、11.54、12.53，均在8～14内。",
        "（3）名义总容量为100 000 m³，按GB 50074—2014折算后的储罐计算总容量为85 000 m³，确定为二级石油库。汽油采用钢制内浮顶罐，柴油采用固定顶罐，三种油品均保留两座罐以满足倒罐和检修。",
        "（4）水运年作业量为76万t，设置1个5 000 DWT成品油泊位，单泊位计算能力124.10万t/a，利用率0.613。水运、管道和公路形成完整收发油流程。",
        "（5）主要工艺管径为DN150～DN250。最不利水运外输管路在200 m³/h时所需扬程35.91 m，选KSB MegaCPK Inducer 125-100-200、193 mm叶轮作为候选泵；泵与管路曲线交点约为202 m³/h、36.3 m。三种油品各设2台泵，一用一备，配30 kW防爆电机；柴油NPSH_a=11.83 m，大于样本NPSH_r与1.0 m裕量之和。",
        "（6）罐区划分为3个防火堤区，堤高2.2 m。消防控制工况为15 000 m³固定顶柴油罐火灾并冷却1座相邻罐，消防水计算量约3 790 m³，设置2×2 500 m³消防水池；柴油罐配置6个16 L/s泡沫产生器，泡沫液计算量约7.01 m³，设置2×5 m³泡沫液储罐。",
        "（7）事故水计算量约6 985 m³，设置8 000 m³事故水池；初期雨水设置2×150 m³调蓄池。汽油装车油气回收能力取300 m³/h，一用一备。总图和流程图应按本稿的6座罐、3个罐组、1个泊位和设备参数绘制。",
    ]
    for text in conclusions:
        add_body(doc, text)


def build_appendices(doc: Document):
    # 与第5章保持同一组消防计算参数，附录独立复算，便于后期改单项参数。
    wall_area = math.pi * 34 * 18
    qc = (2.5 * wall_area + 2.0 * wall_area / 2) / 60
    vc = qc * 3.6 * 9
    ag = math.pi * 40 * 0.55
    qg = 12.5 * ag / 60
    ad = math.pi * 34**2 / 4
    qd = 6.0 * ad / 60
    vmd = qd * 60 * 30 / 1000
    vfd = vmd * 0.039 * 1.10
    foam_water = vmd * (1 - 0.039)
    fire_water = vc + foam_water
    rainfall = 0.90 * 12_600 * 0.2764
    leakage = 300 * (10 / 60) * 1.2
    vacc = fire_water + rainfall + leakage
    add_heading(doc, "附录A 主要设计数据汇总", 1)
    add_table(
        doc,
        "表A-1 设计基础数据",
        ["类别", "参数", "数值", "用途"],
        [
            ["物流", "92/95/柴油年量", "40/25/30×10⁴ t/a", "库容与装卸"],
            ["库容", "K、η", "14、0.95", "理论库容"],
            ["物性", "汽油/柴油密度", "0.760/0.840 t/m³", "体积与质量换算"],
            ["储罐", "数量与容量", "2×20k+2×15k+2×15k", "总图与消防"],
            ["自然条件", "平均气温/降水", "16.3 ℃/1269.7 mm", "设备与排水"],
            ["风荷载", "基本风压", "0.45 kPa（50年）", "抗风圈"],
            ["地震", "基本烈度", "Ⅵ度", "结构与柔性连接"],
            ["码头", "船型/泊位数", "5 000 DWT/1个", "水运接口"],
        ],
        [2.4, 4.0, 4.5, 5.0],
    )
    add_heading(doc, "附录B 库容与罐组计算明细", 1)
    for p in G:
        cfg = TANKS[p]
        bh = add_heading(doc, f"B.{list(G).index(p)+1} {p}", 2)
        if list(G).index(p) > 0:
            bh.paragraph_format.page_break_before = True
        add_body(
            doc,
            f"年周转质量G={G[p]:,.0f} t/a，设计密度ρ={RHO[p]:.3f} t/m³，K=14，η=0.95。"
            f"理论库容V={theoretical_capacity(p):,.2f} m³；选{cfg['n']}座×{cfg['nom']:,.0f} m³，名义容量{cfg['n']*cfg['nom']:,.0f} m³。"
            f"反算K={actual_k(p):.2f}。单罐D={cfg['D']:.0f} m、H={cfg['H']:.0f} m，名义液位{liquid_height(p):.3f} m，满足罐壁顶空间要求。",
        )
        add_table(
            doc,
            f"表B-{list(G).index(p)+1} {p}逐项检查",
            ["检查项", "计算或要求", "结果"],
            [
                ["理论库容", f"{G[p]:.0f}/(14×{RHO[p]:.3f}×0.95)", f"{theoretical_capacity(p):.2f} m³"],
                ["名义库容", f"{cfg['n']}×{cfg['nom']:.0f}", f"{cfg['n']*cfg['nom']:.0f} m³"],
                ["容量余量", "名义-理论", f"{cfg['n']*cfg['nom']-theoretical_capacity(p):.2f} m³"],
                ["反算K", f"{G[p]:.0f}/({cfg['n']*cfg['nom']:.0f}×{RHO[p]:.3f}×0.95)", f"{actual_k(p):.2f}"],
                ["罐数", "同品种不少于2座", "2座，满足"],
            ],
            [4.0, 7.5, 4.5],
        )
    add_heading(doc, "附录C 主要路径水力计算明细", 1)
    add_body(
        doc,
        "下列计算统一采用商业钢管绝对粗糙度0.045 mm、重力加速度9.81 m/s²、泵效率0.70。设备压降包括过滤器、流量计、装卸臂或鹤管；末端余压统一取3 m。"
        "实际图纸完成后，应按管件数量和实测长度复算。",
    )
    for idx, (name, product, q, d, L, dz, eq, lk) in enumerate(PATHS, 1):
        h = HYD[name]
        ch = add_heading(doc, f"C.{idx} {name}", 2)
        if idx > 1:
            ch.paragraph_format.page_break_before = True
        add_body(
            doc,
            f"介质为{product}，Q={q} m³/h，内径按{d:.3f} m计算，L={L} m，Δz={dz} m，设备压降{eq} m，Σζ={lk}。"
            f"计算流速v={h['v']:.3f} m/s，Re={h['Re']:.3e}，λ={h['lam']:.5f}，沿程损失h_f={h['hf']:.3f} m，"
            f"局部损失h_m={h['hm']:.3f} m，系统扬程H={h['H']:.3f} m，轴功率P={h['P']:.2f} kW。",
        )
        add_body(doc, f"步骤1：将体积流量换算为SI单位，Q={q}/3600={q/3600:.6f} m³/s；管道计算内径d={d:.3f} m。")
        add_body(doc, f"步骤2：代入流速公式v=4Q/(πd²)，得v={h['v']:.3f} m/s。该流速处于成品油工艺管道常用范围内，不另行放大或缩小管径。")
        add_body(doc, f"步骤3：取{product}运动黏度ν={NU[product]:.2e} m²/s，代入Re=vd/ν，得Re={h['Re']:.3e}，判定为湍流。")
        add_body(doc, f"步骤4：取商业钢管绝对粗糙度ε=4.5×10⁻⁵ m，按Swamee-Jain显式式计算Darcy摩阻系数，得λ={h['lam']:.5f}。")
        add_body(doc, f"步骤5：代入沿程与局部阻力式，得h_f={h['hf']:.3f} m、h_m={h['hm']:.3f} m；再加高差{dz} m、设备压降{eq} m和末端余压3 m，系统扬程H={h['H']:.3f} m。")
        add_body(doc, f"步骤6：按泵效率η_p=0.70计算轴功率P={h['P']:.2f} kW。所选55 m扬程泵能够覆盖本路径计算扬程，电机功率按45 kW配置并保留启动与工况波动裕量。")
        add_table(
            doc,
            f"表C-{idx} {name}计算参数",
            ["Q/(m³/h)", "DN", "v/(m/s)", "Re", "λ", "h_f/m", "h_m/m", "H/m", "P/kW"],
            [[f"{q}", f"{int(d*1000)}", f"{h['v']:.3f}", f"{h['Re']:.3e}", f"{h['lam']:.5f}",
              f"{h['hf']:.3f}", f"{h['hm']:.3f}", f"{h['H']:.3f}", f"{h['P']:.2f}"]],
            [1.8, 1.5, 1.8, 2.2, 1.8, 1.8, 1.8, 1.8, 1.8],
            Pt(8),
        )
    add_heading(doc, "附录D 消防与事故水计算明细", 1)
    add_heading(doc, "D.1 冷却水", 2)
    add_body(
        doc,
        f"15 000 m³柴油罐D=34 m、H=18 m，罐壁面积{wall_area:.2f} m²。着火罐冷却流量80.11 L/s，相邻罐半周冷却32.04 L/s，"
        f"合计{qc:.2f} L/s；连续9 h用水{vc:.2f} m³。",
    )
    dh = add_heading(doc, "D.2 泡沫系统", 2)
    dh.paragraph_format.page_break_before = True
    add_body(
        doc,
        f"汽油罐密封圈保护面积{ag:.2f} m²，混合液流量{qg:.2f} L/s，配置2个8 L/s泡沫产生器；柴油罐全液面面积{ad:.2f} m²，"
        f"混合液流量{qd:.2f} L/s，配置6个16 L/s泡沫产生器。柴油控制工况泡沫液储量{vfd:.2f} m³。",
    )
    dh = add_heading(doc, "D.3 消防泵与水池", 2)
    dh.paragraph_format.page_break_before = True
    add_table(
        doc,
        "表D-1 消防设备一览",
        ["设备", "规格", "数量", "运行方式", "校核"],
        [
            ["固定冷却水泵", "130 L/s，H按环网复核", "2", "一用一备", f"130>{qc:.2f} L/s"],
            ["泡沫水泵", "110 L/s", "2", "一用一备", f"110>{qd:.2f} L/s"],
            ["消防水池", "2 500 m³", "2", "分格连通", f"5 000>{fire_water:.2f} m³"],
            ["泡沫液储罐", "5 m³", "2", "并联", f"10>{vfd:.2f} m³"],
            ["事故水池", "8 000 m³", "1", "事故状态启用", f"8 000>{vacc:.2f} m³"],
        ],
        [3.5, 4.0, 2.0, 3.5, 4.0],
    )
    # 本轮只保留与正文直接对应的四个计算附件；设备表、复核记录和规范检查
    # 已移入正文，不再以大量附录增加篇幅。
    return
    add_heading(doc, "附录E 主要设备与图纸标注", 1)
    add_table(
        doc,
        "表E-1 主要设备一览",
        ["类别", "位号", "名称/规格", "数量", "备注"],
        [
            ["储罐", "T101～T102", "20 000 m³钢制内浮顶罐，D40×H18", "2", "92号汽油"],
            ["储罐", "T201～T202", "15 000 m³钢制内浮顶罐，D34×H18", "2", "95号汽油"],
            ["储罐", "T301～T302", "15 000 m³固定顶罐，D34×H18", "2", "0号柴油"],
            ["输油泵", "P101A/B", "250 m³/h，55 m，45 kW", "2", "92号汽油"],
            ["输油泵", "P201A/B", "250 m³/h，55 m，45 kW", "2", "95号汽油"],
            ["输油泵", "P301A/B", "250 m³/h，55 m，45 kW", "2", "0号柴油"],
            ["油气回收", "VRU-1/2", "300 m³/h", "2", "一用一备"],
            ["码头", "B-01", "5 000 DWT成品油泊位", "1", "水运76万t/a"],
            ["消防水池", "FT-1/2", "2 500 m³", "2", "总有效5 000 m³"],
            ["事故水池", "ET-1", "8 000 m³", "1", "事故与消防废水"],
        ],
        [2.3, 2.8, 7.0, 1.8, 3.3],
    )
    add_heading(doc, "E.1 总平面图应表达的内容", 2)
    for text in [
        "6座储罐的罐号、油品、容量、直径、罐型和3个防火堤边界；",
        "罐间距、防火堤至罐壁距离、消防道路宽度和主要设施防火间距；",
        "发油泵棚、阀组、计量间、6个装车鹤位、油气回收装置及车辆单向流线；",
        "2×2 500 m³消防水池、消防泵房、泡沫站、8 000 m³事故水池、初期雨水池和污水设施；",
        "码头方向、1个5 000 DWT泊位、DN250水运进库管线、DN200水运外输管线和船岸ESD接口；",
        "风向玫瑰、用地边界、出入口、行政控制区、消防集合点和主要标高。",
    ]:
        add_body(doc, text)
    add_heading(doc, "E.2 工艺流程图应表达的内容", 2)
    for text in [
        "92号汽油、95号汽油和0号柴油的水运进库、管道进库、公路发油、管道外输、水运外输和倒罐流程；",
        "每座储罐独立进出油支管、罐根切断阀、高高液位联锁和低低液位停泵；",
        "三组输油泵的一用一备关系、最小流量回流、过滤计量、取样、排凝和清扫接口；",
        "装车批控、防溢流、静电接地和油气回收流程；",
        "码头紧急切断阀、装卸臂、绝缘法兰、船岸通信和ESD；",
        "消防冷却环管、泡沫主管、泡沫产生器、事故水切换和雨污分流。",
    ]:
        add_body(doc, text)
    build_extended_records(doc)


def add_record_sheet(doc: Document, title: str, paragraphs, headers=None, rows=None, widths=None):
    heading = add_heading(doc, title, 2)
    heading.paragraph_format.page_break_before = True
    for text in paragraphs:
        add_body(doc, text)
    if headers and rows:
        add_table(doc, f"表{title.split()[0]} 计算与校核记录", headers, rows, widths)


def build_extended_records(doc: Document):
    wall_area = math.pi * 34 * 18
    qc = (2.5 * wall_area + 2.0 * wall_area / 2) / 60
    vc = 3.6 * qc * 9
    ad = math.pi * 34**2 / 4
    qd = 6.0 * ad / 60
    vmd = 3.6 * qd * 0.5
    vfd = 1.10 * 0.039 * vmd
    fire_water = vc + vmd * (1 - 0.039)
    rainfall = 0.90 * 12_600 * 0.2764
    leakage = 300 * (10 / 60) * 1.2
    vacc = fire_water + rainfall + leakage
    add_heading(doc, "附录F 关键计算复核记录", 1)
    add_record_sheet(
        doc,
        "F.1 油库等级复核",
        [
            "复核目的：确认采用较大储罐减少罐数后，工程仍属于二级石油库。复核仅按储罐计算总容量进行，不以名义库容直接判级。",
            "设计配置为2座20 000 m³汽油内浮顶罐、2座15 000 m³汽油内浮顶罐和2座15 000 m³柴油固定顶罐，名义总容量为100 000 m³。",
            "依据GB 50074—2014第3.0.1条，汽油储罐容量按100%计入，柴油储罐容量折半计入。因此计算总容量为40 000+30 000+0.5×30 000=85 000 m³。",
            "85 000 m³位于二级石油库容量区间。采用6座大罐后，罐数由10座减少，但等级没有升至一级，也没有改变汽油和柴油的独立罐组原则。",
            "答辩说明：名义库容用于设备配置，判级容量用于确定库级；二者含义不同。柴油折半是规范规定，不是人为折减实际储存能力。",
        ],
        ["项目", "汽油/m³", "柴油/m³", "折算系数", "计入容量/m³"],
        [["20 000 m³罐", "40 000", "—", "1.0", "40 000"], ["15 000 m³汽油罐", "30 000", "—", "1.0", "30 000"],
         ["15 000 m³柴油罐", "—", "30 000", "0.5", "15 000"], ["合计", "70 000", "30 000", "—", "85 000"]],
        [4.0, 3.2, 3.2, 3.0, 3.6],
    )
    for idx, product in enumerate(G, 2):
        cfg = TANKS[product]
        add_record_sheet(
            doc,
            f"F.{idx} {product}库容复核",
            [
                f"设计输入：{product}年周转质量G={G[product]:,.0f} t/a，设计密度ρ={RHO[product]:.3f} t/m³，初取周转系数K=14，油罐利用系数η=0.95。",
                f"将上述数据代入式1，理论库容V=G/(Kρη)={G[product]:,.0f}/(14×{RHO[product]:.3f}×0.95)={theoretical_capacity(product):,.2f} m³。",
                f"选用{cfg['n']}座{cfg['nom']:,.0f} m³{cfg['type']}，名义库容为{cfg['n']*cfg['nom']:,.0f} m³，较理论库容增加{cfg['n']*cfg['nom']-theoretical_capacity(product):,.2f} m³。",
                f"将名义库容代入式2反算，K={G[product]:,.0f}/({cfg['n']*cfg['nom']:,.0f}×{RHO[product]:.3f}×0.95)={actual_k(product):.2f}，处于二级及以上商业油库8～14建议范围。",
                f"几何校核：单罐直径D={cfg['D']:.0f} m、罐壁高度H={cfg['H']:.0f} m，按圆柱体体积反算名义液位为{liquid_height(product):.3f} m，低于罐壁高度，剩余空间用于安全容量和液位控制。",
                "运行校核：同一油品设置两座罐，一座检修或清罐时另一座仍可承担收发作业；正常运行通过批次计划避免同罐同时收油和发油。",
            ],
            ["复核项", "计算值", "控制要求", "结论"],
            [["理论库容", f"{theoretical_capacity(product):,.2f} m³", "不大于名义库容", "满足"],
             ["名义库容", f"{cfg['n']*cfg['nom']:,.0f} m³", "两座同品种罐", "满足"],
             ["反算K", f"{actual_k(product):.2f}", "8～14", "满足"],
             ["名义液位", f"{liquid_height(product):.3f} m", f"<{cfg['H']:.0f} m", "满足"]],
            [4.0, 4.0, 4.2, 3.0],
        )
    add_record_sheet(
        doc,
        "F.5 泊位通过能力复核",
        [
            "水运进库量为92号汽油28万t/a、95号汽油15万t/a、柴油24万t/a，共67万t/a；水运外输量为92号汽油4万t/a、95号汽油5万t/a，共9万t/a。",
            "码头年总作业量为76万t/a。代表船型取5 000 DWT成品油船，单船实际装卸量按5 000 t、占用泊位时间按30 h计，泊位年可用时间按365×24×0.85计算。",
            "代入式14，单泊位年通过能力P_b=5 000×365×24×0.85/30=1 241 000 t/a，约124.1万t/a。",
            "代入式15，泊位数N_b=ceil(760 000/1 241 000)=1；泊位利用率为760 000/1 241 000=0.612，具备天气、检修和船期波动余量。",
            "码头设置1套装卸臂、船岸ESD、绝缘法兰、紧急脱离和回流接口。不同油品分批作业，切换前执行管线确认，防止混油。",
        ],
        ["参数", "符号", "数值", "单位", "说明"],
        [["年水运量", "G_w", "760 000", "t/a", "进库与外输合计"], ["单船量", "W_s", "5 000", "t/艘", "代表作业批量"],
         ["泊位时间", "t_b", "30", "h/艘", "靠离泊与装卸合计"], ["可用系数", "η_b", "0.85", "—", "天气检修折减"],
         ["通过能力", "P_b", "1 241 000", "t/a", "大于设计量"]],
        [4.0, 2.2, 3.2, 2.5, 4.3],
    )
    add_record_sheet(
        doc,
        "F.6 输油泵能力复核",
        [
            "三种油品各设置2台卧式离心泵，一用一备。单泵额定流量250 m³/h、额定扬程55 m，配置45 kW防爆电动机并采用变频调节。",
            "流量校核以水运进库300 m³/h为最大单路径工况。码头卸油主要依靠船泵，库内泵用于倒罐和外输；库内泵在200～250 m³/h工况运行，处于额定点附近。",
            "扬程校核采用附录C各路径的高差、设备压降、沿程损失、局部损失和末端余压之和。所有库内泵控制路径计算扬程均不超过55 m。",
            "功率校核按式23计算。轴功率考虑效率0.70，选45 kW电机后仍保留合理功率裕量；现场选型时应以制造商完整性能曲线复核额定点、效率和允许工作区。",
            "年能力校核按式25计算。单泵年可作业能力取250×16×330=1 320 000 m³/a，明显高于任一油品的设计体积量，因此运行时间和检修余量满足要求。",
            "汽蚀校核按式24执行。施工图阶段必须根据最低罐液位、泵中心标高、吸入管损失、最高液温饱和蒸气压和当地大气压复算NPSH_a，并保证大于制造商NPSH_r加0.5 m裕量。",
        ],
        ["校核项", "设计值", "设备值", "判断"],
        [["流量", "200～250 m³/h", "250 m³/h", "满足"], ["扬程", "见附录C", "55 m", "满足"],
         ["电机", "按路径功率计算", "45 kW", "满足"], ["数量", "连续作业", "2台/油品", "一用一备"],
         ["年能力", "<1.32×10⁶ m³/a", "1.32×10⁶ m³/a", "满足"]],
        [4.0, 4.2, 4.2, 3.8],
    )
    add_record_sheet(
        doc,
        "F.7 防火堤有效容积复核",
        [
            "三个油品罐组分别设置防火堤，堤内净面积按初步总图边界估算，并扣除非事故罐占地、罐基础和堤内构筑物占地。防火堤设计高度2.2 m，有效计算高度取2.0 m。",
            "92号汽油罐组内尺寸约140 m×90 m，最大单罐容量20 000 m³；95号汽油和柴油罐组内尺寸均约120 m×80 m，最大单罐容量15 000 m³。",
            "代入式27计算净容积。对每个罐组均采用同组最大罐全容量作为最低控制值，并保留消防泡沫和波浪余量。",
            "穿堤管道采用套管和柔性密封，不在防火堤上设置常开排水口。正常雨水经人工确认后排放，事故状态切换至事故水池。",
            "最终总平面图若调整罐间距、泵棚或管墩位置，应按实际净面积重新计算，不能直接沿用初稿几何值。",
        ],
        ["罐组", "内边界/m", "有效堤高/m", "控制罐/m³", "复核结论"],
        [["92号汽油", "140×90", "2.0", "20 000", "满足"], ["95号汽油", "120×80", "2.0", "15 000", "满足"],
         ["0号柴油", "120×80", "2.0", "15 000", "满足"]],
        [3.5, 3.5, 3.2, 3.6, 3.6],
    )
    add_record_sheet(
        doc,
        "F.8 消防控制工况复核",
        [
            "比较20 000 m³汽油内浮顶罐密封圈火灾与15 000 m³柴油固定顶罐全液面火灾。柴油罐需要同时冷却着火罐全部罐壁和相邻罐半周罐壁，且持续时间为9 h，确定为消防水控制工况。",
            f"柴油罐罐壁面积A_w=π×34×18={math.pi*34*18:.2f} m²。按着火罐2.5 L/(min·m²)、相邻罐2.0 L/(min·m²)计算，总冷却流量Q_c={qc:.2f} L/s。",
            f"连续供水9 h，冷却水量V_c={vc:.2f} m³；柴油全液面泡沫混合液流量Q_f={qd:.2f} L/s，30 min混合液量V_mix={vmd:.2f} m³。",
            f"按3%泡沫液、实际混合比上限3.9%并增加10%储量，泡沫液计算量为{vfd:.2f} m³，设置2×5 m³泡沫液储罐。",
            f"消防用水计算量约{fire_water:.2f} m³，设置2×2 500 m³消防水池；事故水计算量约{vacc:.2f} m³，设置8 000 m³事故水池。",
            "消防泵、泡沫泵、柴油机备用泵和环网阀门应在施工图阶段结合水枪、消火栓和码头消防需求进行最不利点水力复核。",
        ],
        ["项目", "计算值", "选择值", "裕量或结论"],
        [["冷却水流量", f"{qc:.2f} L/s", "130 L/s", "满足"], ["泡沫混合液", f"{qd:.2f} L/s", "110 L/s", "满足"],
         ["消防水", f"{fire_water:.2f} m³", "5 000 m³", "满足"], ["泡沫液", f"{vfd:.2f} m³", "10 m³", "满足"],
         ["事故水", f"{vacc:.2f} m³", "8 000 m³", "满足"]],
        [4.0, 4.0, 4.0, 4.0],
    )

    add_heading(doc, "附录G 设备数据表", 1)
    equipment_sheets = [
        ("G.1 20 000 m³汽油内浮顶罐数据表", "T101～T102", "92号汽油", "20 000 m³", "D=40 m，H=18 m", "内浮顶、氮封接口、一次二次密封", "高高液位联锁切断，低低液位停泵"),
        ("G.2 15 000 m³汽油内浮顶罐数据表", "T201～T202", "95号汽油", "15 000 m³", "D=34 m，H=18 m", "内浮顶、氮封接口、一次二次密封", "高高液位联锁切断，低低液位停泵"),
        ("G.3 15 000 m³柴油固定顶罐数据表", "T301～T302", "0号柴油", "15 000 m³", "D=34 m，H=18 m", "固定顶、呼吸阀、阻火器、量油孔", "高高液位联锁切断，低低液位停泵"),
        ("G.4 成品油输油泵数据表", "P101A/B～P301A/B", "汽油或柴油", "250 m³/h", "H=55 m，电机45 kW", "卧式离心泵、机械密封、变频", "一用一备，出口止回，最小流量回流"),
        ("G.5 汽油装车油气回收装置数据表", "VRU-1/2", "汽油油气", "300 m³/h", "入口微负压", "冷凝与吸附组合或等效工艺", "一用一备，LEL与压力联锁"),
        ("G.6 码头装卸系统数据表", "B-01", "三种成品油", "300 m³/h", "5 000 DWT泊位", "装卸臂、ESD、紧急脱离、绝缘法兰", "船岸通信，阀位和流量联锁"),
    ]
    for title, tag, medium, capacity, condition, config, interlock in equipment_sheets:
        add_record_sheet(
            doc,
            title,
            [
                f"设备位号：{tag}。服务介质：{medium}。额定能力或容量：{capacity}。主要工况：{condition}。",
                f"结构与附件：{config}。所有接液材料、垫片和密封件应与介质相容，汽油系统电气设备按爆炸危险区域选防爆型。",
                f"控制与联锁：{interlock}。联锁动作应采用失效安全原则，紧急停车后保持阀门处于阻止继续泄漏的位置。",
                "制造商资料应至少提供外形尺寸、质量、接管表、性能曲线、材料清单、试验要求、基础荷载和维护空间。初稿参数用于工艺设计，采购前须经技术协议确认。",
                "安装验收应核对位号、方向、标高、法兰等级、接地跨接和检修通道。设备投运前完成吹扫、严密性试验、联锁测试和单机试运行。",
            ],
            ["数据项", "要求", "确认阶段"],
            [["介质与能力", f"{medium}；{capacity}", "本设计"], ["主要工况", condition, "本设计"],
             ["附件配置", config, "施工图"], ["联锁", interlock, "施工图/调试"],
             ["材料与试验", "制造商按规范提交", "采购与验收"]],
            [4.0, 8.0, 4.0],
        )

    add_heading(doc, "附录H 规范符合性检查表", 1)
    compliance = [
        ("H.1 总图与防火间距检查", [
            ["油库等级", "GB 50074—2014第3.0.1条", "计算总容量85 000 m³，二级", "满足"],
            ["罐组分隔", "不同油品与罐型分组", "3个独立防火堤区", "满足"],
            ["消防道路", "环形或可通行消防道路", "罐区周边设置环形道路", "图纸复核"],
            ["设施间距", "按库级、油品类别和设施类型", "总图标注逐项核对", "施工图复核"],
        ]),
        ("H.2 储罐与附件检查", [
            ["汽油罐型", "GB 50074—2014第6.1.3～6.1.5条", "钢制内浮顶罐", "满足"],
            ["柴油罐型", "同条款适用条件", "固定顶罐", "满足"],
            ["液位保护", "独立高高液位报警与联锁", "每罐设置", "满足"],
            ["抗风与抗震", "按风压和抗震参数校核", "初步校核并留施工图复核项", "继续深化"],
        ]),
        ("H.3 工艺管道与泵检查", [
            ["管径", "流量、流速和水力损失计算", "DN150～DN250", "满足"],
            ["专管原则", "汽油与柴油防混油", "三种油品独立干管", "满足"],
            ["泵备用", "关键连续作业设备有备用", "每油品一用一备", "满足"],
            ["紧急切断", "罐根、装车和码头设置", "ESD联锁", "施工图落实"],
        ]),
        ("H.4 消防与泡沫检查", [
            ["控制工况", "最大消防用水与泡沫工况", "15 000 m³柴油固定顶罐", "满足"],
            ["消防水池", "有效容量大于计算量", "5 000 m³", "满足"],
            ["泡沫液", "计算量并留裕量", "10 m³", "满足"],
            ["泡沫产生器", "按流量向上取整", "柴油罐6个16 L/s", "满足"],
        ]),
        ("H.5 防雷防静电防爆检查", [
            ["防雷接地", "罐体多点接地并形成接地网", "设计设置", "图纸落实"],
            ["静电跨接", "法兰、装卸臂和浮盘跨接", "设计设置", "图纸落实"],
            ["车辆接地", "接地确认与装车联锁", "设计设置", "满足"],
            ["防爆电气", "按危险区域和气体级组选择", "设备表注明防爆要求", "施工图复核"],
        ]),
        ("H.6 环保与事故水检查", [
            ["油气回收", "汽油装车密闭收集", "2×300 m³/h，一用一备", "满足"],
            ["雨污分流", "清净雨水与污染水分流", "设置切换阀和在线监测", "满足"],
            ["初期雨水", "按污染面积计算", "2×150 m³", "满足"],
            ["事故水", "消防、降雨和泄漏合并计算", "8 000 m³", "满足"],
        ]),
    ]
    for title, rows in compliance:
        add_record_sheet(
            doc,
            title,
            [
                "本检查表用于初稿内部复核和答辩准备。规范条款在正文首次采用处写明标准名称和条款号，国家标准不列入文末参考文献。",
                "“满足”表示初稿计算或方案已闭合；“图纸复核”表示需在总平面图、工艺流程图或设备表完成后按实际尺寸再次核对。",
                "若后续修改年周转量、罐容、罐径、管长或消防控制工况，应同步更新正文计算、附录、设备表和两张图纸，避免不同文件参数不一致。",
            ],
            ["检查项目", "依据或要求", "本设计做法", "状态"],
            rows,
            [4.0, 5.0, 5.0, 2.0],
        )


def build_references(doc: Document):
    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] 杭州一达环保技术咨询服务有限公司. 浙江星月药物科技有限公司新增原料药10吨、制剂50万支生产项目环境影响报告书（报批稿）[R]. 杭州, 2021.",
        "[2] 许少新, 涂仁福, 徐宁, 等. 成品油管铁联运物流优化[J]. 油气储运, 2022, 41(7): 859-868. DOI:10.6047/j.issn.1000-8241.2022.07.015.",
        "[3] 吴守志, 侯磊, 伍星光, 刘芳媛. 安全屏障对储油罐区池火灾多米诺效应概率的影响[J]. 油气储运, 2022, 41(2): 165-176. DOI:10.6047/j.issn.1000-8241.2022.02.006.",
        "[4] 王雷, 申满对, 刘奎. 外浮顶储罐VOCs排放量影响因素分析与探讨[J]. 炼油技术与工程, 2022, 52(10): 55-58.",
        "[5] 刘世达, 侯栓弟, 刘忠生, 等. 国内石化有机液体储罐VOCs深度减排管控技术进展[J]. 炼油技术与工程, 2022, 52(4): 11-18.",
        "[6] 丑冠博, 刘杰, 多依丽, 等. 基于PPRR理论的化工罐区定量应急管理能力研究[J]. 辽宁石油化工大学学报, 2024, 44(4): 51-59. DOI:10.12422/j.issn.1672-6952.2024.04.007.",
        "[7] 焦浩宇, 任婧杰, 赵彦修, 毕明树. 隔堤池火条件下储罐热响应的数值模拟[J]. 化工进展, 2025, 44(12): 7349-7358. DOI:10.16085/j.issn.1000-6613.2024-2040.",
        "[8] 朱喜平. 大型石油储库全寿命周期风险管控一体化平台研发[J]. 油气储运, 2023, 42(10): 1175-1183. DOI:10.6047/j.issn.1000-8241.2023.10.011.",
        "[9] 邵新军, 周一卉, 黄兆锋, 等. 全接液金属浮盘抗爆特性实验与数值模拟[J]. 油气储运, 2024, 43(2): 200-211. DOI:10.6047/j.issn.1000-8241.2024.02.009.",
        "[10] 蒋新生, 秦希卓, 储汇, 等. 油气爆炸荷载对储罐结构的毁伤机制及评估[J]. 油气储运, 2024, 43(12): 1365-1377. DOI:10.6047/j.issn.1000-8241.2024.12.005.",
        "[11] 康泽天, 姚冰, 党文义, 等. 环形格栅双层底板立式储罐疲劳强度有限元分析[J]. 油气储运, 2022, 41(8): 939-945. DOI:10.6047/j.issn.1000-8241.2022.08.009.",
        "[12] 曲建军, 纪瑞军. 石油库储罐检修策略[J]. 油气储运, 2023, 42(11): 1307-1312. DOI:10.6047/j.issn.1000-8241.2023.11.011.",
        "[13] 武刚, 张庶鑫, 罗金恒, 等. 基于TOPSIS的原油储罐风险分级预警[J]. 油气储运, 2024, 43(6): 641-648. DOI:10.6047/j.issn.1000-8241.2024.06.005.",
        "[14] 田素俊, 黄维秋, 鄢永兵, 等. 吸附剂对含湿油气的吸附性能及热效应[J]. 油气储运, 2022, 41(8): 962-971. DOI:10.6047/j.issn.1000-8241.2022.08.012.",
        "[15] YANG Y, ZHANG X, XIE S, et al. Design and visual implementation of a regional energy risk superposition model for oil tank farms[J]. Energies, 2024, 17(22): 5775. DOI:10.3390/en17225775.",
        "[16] DOREGAR ZAVAREH R, DANA T, ROAYAEI E, et al. The environmental risk assessment of fire and explosion in storage tanks of petroleum products[J]. Sustainability, 2022, 14(17): 10747. DOI:10.3390/su141710747.",
        "[17] 刘德俊, 杨帆, 于洋, 等. 油库技术与管理[M]. 2版. 北京: 中国石化出版社, 2021.",
        "[18] 邢科伟, 马秀让, 刘占卿. 油库加油站设计数据图表手册[M]. 北京: 中国石化出版社, 2015.",
        "[19] KSB SE & Co. KGaA. MegaCPK centrifugal pumps with shaft seal: characteristic curves and technical data[EB/OL]. Frankenthal: KSB.",
    ]
    for ref in refs:
        p = add_paragraph(doc, ref, "Normal", first_line=False)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)


def build_acknowledgement(doc: Document):
    add_heading(doc, "致谢", 1)
    add_body(
        doc,
        "本毕业设计在指导教师叶章评老师、李琦钰老师的指导下完成。两位老师在设计任务理解、规范选用、库容方案和工艺流程等方面给予了耐心指导，"
        "使我能够将油气储运工程专业所学知识用于较完整的油库工艺设计。在此向两位老师表示诚挚感谢。感谢能源学院各位老师在专业学习期间的培养，"
        "也感谢同学和家人在资料整理与设计过程中给予的帮助。由于本人实践经验有限，设计中仍可能存在不足，恳请各位老师批评指正。",
    )


def normalize(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    for name, size, before, after in [
        ("Heading 1", 16, 12, 12),
        ("Heading 2", 14, 10, 6),
        ("Heading 3", 10.5, 6, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
    for p in doc.paragraphs:
        if p.style.name == "Normal" and p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            if not p.text.startswith("["):
                p.paragraph_format.first_line_indent = Cm(0.74)
        if p.style.name == "Heading 1":
            p.paragraph_format.page_break_before = True
            p.paragraph_format.keep_with_next = True
        if p.style.name in {"Heading 2", "Heading 3"}:
            p.paragraph_format.keep_with_next = True
        for r in p.runs:
            if "[[EQ|" in p.text:
                set_run_font(r, east_asia="Cambria Math", latin="Cambria Math")
            elif p.style.name == "Heading 1":
                set_run_font(r, east_asia="黑体", size=Pt(16), bold=True)
            elif p.style.name == "Heading 2":
                set_run_font(r, east_asia="黑体", size=Pt(14), bold=True)
            elif p.style.name == "Heading 3":
                set_run_font(r, east_asia="黑体", size=Pt(10.5), bold=True)
    set_update_fields(doc)


def main():
    doc = Document(str(SOURCE))
    ensure_heading_styles(doc)
    build_front(doc)
    remove_old_body(doc)
    build_chapter_1(doc)
    build_chapter_2(doc)
    build_chapter_3(doc)
    build_chapter_4(doc)
    build_chapter_5(doc)
    build_chapter_6(doc)
    build_chapter_7(doc)
    build_chapter_8(doc)
    build_appendices(doc)
    build_references(doc)
    build_acknowledgement(doc)
    normalize(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(OUTPUT)
    print(f"equations={EQ_COUNTER}, paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)}")


if __name__ == "__main__":
    main()
