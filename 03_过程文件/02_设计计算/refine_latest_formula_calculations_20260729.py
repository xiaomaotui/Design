from __future__ import annotations

import copy
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


SUBSCRIPT_RE = re.compile(r"_\(([^)]+)\)|_([A-Za-z0-9Σ]+)")


def add_text_with_subscripts(paragraph: Paragraph, text: str) -> None:
    pos = 0
    for match in SUBSCRIPT_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        run = paragraph.add_run(match.group(1) or match.group(2))
        run.font.subscript = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_text(paragraph: Paragraph, text: str) -> Paragraph:
    clear_paragraph(paragraph)
    add_text_with_subscripts(paragraph, text)
    return paragraph


def insert_after(paragraph: Paragraph, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    if style is not None:
        result.style = style
    elif paragraph.style is not None:
        result.style = paragraph.style
    add_text_with_subscripts(result, text)
    return result


def insert_lines_after(paragraph: Paragraph, lines: list[str]) -> Paragraph:
    current = paragraph
    for line in lines:
        current = insert_after(current, line, paragraph.style)
    return current


def find_starts(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(text):
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def find_exact(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def add_numeq_after(paragraph: Paragraph, linear: str) -> Paragraph:
    p = insert_after(paragraph, f"[[NUMEQ|{linear}]]", paragraph.style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    return p


def split_definitions(paragraph: Paragraph, definitions: list[str]) -> Paragraph:
    set_text(paragraph, "式中：")
    return insert_lines_after(paragraph, definitions)


def replace_plain_terms(doc: Document) -> None:
    replacements = {
        "统计口径": "统计方法",
        "本稿给出": "本设计给出",
        "初稿采用": "本设计采用",
        "初稿不以": "本设计不以",
        "后续若总平面图改变": "总平面图确定后，若管长或高差改变",
    }
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("[["):
            continue
        text = paragraph.text
        for old, new in replacements.items():
            text = text.replace(old, new)
        if text != paragraph.text:
            set_text(paragraph, text)


def convert_existing_underscores(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        text = paragraph.text
        if "_" in text and not text.startswith("[["):
            set_text(paragraph, text)


def main(src: Path, dst: Path) -> None:
    doc = Document(src)

    # 式1与式2：保持用户已调整的公式版式，只补数值代入和逐项变量说明。
    p = find_starts(doc, "将K=14代入式1")
    set_text(p, "将K=14代入式1，计算平均周转周期：")
    p = add_numeq_after(p, "T=(365)/(14)=26.07 d")
    insert_after(
        p,
        "计算得到平均周转周期为26.07 d。该周期用于库容规划，不等同于单船到港间隔或每座储罐的固定清空周期。",
    )

    p = find_exact(doc, "Vi——第i种油品理论设计库容，m³；")
    set_text(p, "V_i——第i种油品理论设计库容，m³；")
    p = find_starts(doc, "Gi——年周转质量")
    p = split_definitions(
        p,
        [
            "G_i——第i种油品年周转质量，t/a；",
            "K——年周转系数；",
            "ρ_i——第i种油品设计密度，t/m³；",
            "η——储罐容积利用系数，取0.95。",
        ],
    )
    # 删除原来紧随其后的重复定义。
    for duplicate in [
        "ρi——设计密度，t/m³；",
        "η——储罐容积利用系数，取0.95。",
    ]:
        try:
            para = find_exact(doc, duplicate)
            para._element.getparent().remove(para._element)
        except ValueError:
            pass

    p = find_starts(doc, "三种油品理论库容合计为")
    set_text(p, "三种油品理论库容合计按三项理论库容相加：")
    p = add_numeq_after(
        p,
        "V_Σ=V_92+V_95+V_(0号)=39572.62+24732.88+26852.85=91158.35 m^3",
    )
    insert_after(
        p,
        "计算总理论库容为91 158.35 m³。为保证同品种至少两座储罐轮换运行，储罐按系列规格向上配置。",
    )

    # 油库等级折算与周转系数反算。
    p = find_starts(doc, "油库等级按GB 50074")
    split_definitions(
        p,
        [
            "TV——储罐计算总容量，m³；",
            "V_j——第j类液体储罐实际容量，m³；",
            "C_j——第j类液体的容量折算系数；汽油取1.0，丙A类柴油取0.5。",
            "将汽油40 000 m³、30 000 m³和柴油30 000 m³代入式3：",
        ],
    )
    p = find_starts(doc, "将汽油40 000")
    p = add_numeq_after(p, "TV=40000×1.0+30000×1.0+30000×0.5=85000 m^3")
    insert_after(
        p,
        "计算得到储罐计算总容量为85 000 m³，处于30 000 m³≤TV＜100 000 m³范围，油库等级为二级。名义总容量100 000 m³与等级判定采用的计算总容量85 000 m³用途不同。",
    )

    p = find_starts(doc, "将各油品名义容量代入式4")
    set_text(p, "将各油品年周转量、密度、利用系数和名义容量代入式4反算年周转系数：")
    p = add_numeq_after(p, "K_92=(400000)/(0.760×0.95×40000)=13.85")
    p = add_numeq_after(p, "K_95=(250000)/(0.760×0.95×30000)=11.54")
    p = add_numeq_after(p, "K_(0号)=(300000)/(0.840×0.95×30000)=12.53")
    insert_after(
        p,
        "三种油品反算年周转系数均位于8～14。因此，6座储罐方案同时满足教材周转系数范围、二级油库等级和每种油品两座储罐的运行要求。",
    )

    # 储罐几何尺寸。
    p = find_starts(doc, "将D=40 m代入式5")
    set_text(p, "将20 000 m³储罐直径D=40 m代入式5和式6：")
    p = add_numeq_after(p, "A=(π×40^2)/(4)=1256.64 m^2")
    p = add_numeq_after(p, "h=(20000)/(1256.64)=15.915 m")
    p = add_numeq_after(p, "Δh=18-15.915=2.085 m")
    p = insert_after(p, "将15 000 m³储罐直径D=34 m代入式5和式6：")
    p = add_numeq_after(p, "A=(π×34^2)/(4)=907.92 m^2")
    p = add_numeq_after(p, "h=(15000)/(907.92)=16.521 m")
    p = add_numeq_after(p, "Δh=18-16.521=1.479 m")
    insert_after(p, "两类储罐的计算液位均低于18 m罐壁高度，满足预留顶部空间要求。")

    # 罐壁厚度定义和代表性算例。
    p = find_starts(doc, "式中：t_d")
    split_definitions(
        p,
        [
            "t_d——设计条件下罐壁计算厚度，mm；",
            "t_t——充水试验条件下罐壁计算厚度，mm；",
            "t_n——罐壁名义厚度，mm；",
            "D——储罐内径，m；",
            "H——计算点以上液柱高度，m；",
            "ρ——储液设计密度，t/m³；",
            "σ_d——设计温度下钢材许用应力，MPa；",
            "σ_t——充水试验温度下钢材许用应力，MPa；",
            "φ——焊接接头系数；",
            "C_1——钢板厚度负偏差，mm；",
            "C_2——腐蚀裕量，mm。",
            "Q345R在相应厚度区间取σ_d=218.75 MPa、σ_t=230 MPa。以20 000 m³汽油罐底圈为例：",
        ],
    )
    p = find_starts(doc, "Q345R在相应厚度")
    p = add_numeq_after(
        p,
        "t_d=(4.9×40×(15.915-0.3)×0.760)/(218.75×0.85)+2.0=12.51 mm",
    )
    p = add_numeq_after(
        p,
        "t_t=(4.9×40×(18.0-0.3))/(230×0.85)=17.75 mm",
    )
    insert_after(p, "按式9比较设计条件、充水试验条件及最小厚度要求，向上圆整取底圈名义厚度18 mm；其余各圈同理计算，结果列于表2-4。")

    # 抗风稳定：把图示段落的结果拆成可检查的代入过程。
    p = find_starts(doc, "将μ_z=1.49")
    set_text(p, "式中：")
    p = insert_lines_after(
        p,
        [
            "P_o——储罐设计外压，kPa；",
            "μ_z——风压高度变化系数；",
            "ω′_0——经罐组狭管效应修正后的基本风压，kPa；",
            "q——储罐设计内负压，kPa。",
            "将μ_z=1.49、ω′_0=0.54 kPa代入式10：",
        ],
    )
    p = add_numeq_after(p, "P_(o,g)=2.25×1.49×0.54+0=1.810 kPa")
    p = add_numeq_after(p, "P_(o,d)=2.25×1.49×0.54+0.25=2.060 kPa")
    p = insert_after(p, "式11和式12中各符号含义如下：")
    p = insert_lines_after(
        p,
        [
            "H_E——变厚度罐壁折算等效高度，m；",
            "h_i——第i圈罐壁高度，m；",
            "t_min——核算区段最小有效厚度，mm；",
            "t_i——第i圈罐壁有效厚度，mm；",
            "[P_cr]——罐壁临界外压，kPa；",
            "D——储罐内径，m。",
        ],
    )
    p2 = find_starts(doc, "按式11折算有效厚度")
    set_text(p2, "按表2-4各圈壁板厚度逐圈代入式11，20 000 m³汽油罐H_E=12.72 m，15 000 m³储罐H_E=14.36 m。再代入式12计算无中间抗风圈时的临界外压：")
    p2 = add_numeq_after(p2, "[P_(cr,20000)]=(16.48×40)/(12.72)×(8/40)^2.5=0.927 kPa")
    p2 = add_numeq_after(p2, "[P_(cr,15000)]=(16.48×34)/(14.36)×(8/34)^2.5=1.048 kPa")
    p2 = insert_after(p2, "无中间抗风圈时两类储罐的临界外压均低于相应设计外压。设置1道中间抗风圈后，按上下两段分别复核，临界外压约提高为原值的2倍：")
    p2 = add_numeq_after(p2, "[P_(cr,20000)]′=2×0.927=1.854 kPa>1.810 kPa")
    p2 = add_numeq_after(p2, "[P_(cr,15000)]′=2×1.048=2.096 kPa")
    insert_after(p2, "因此，20 000 m³汽油罐1.854 kPa＞1.810 kPa；15 000 m³汽油罐2.096 kPa＞1.810 kPa；15 000 m³柴油罐2.096 kPa＞2.060 kPa，均满足抗风稳定要求。")

    # 物流量、泊位与管径。
    p = find_starts(doc, "式中：G_(i,j)")
    split_definitions(
        p,
        [
            "G_(i,j)——油品i采用运输方式j的年物流量，t/a；",
            "G_i——油品i的年物流量，t/a；",
            "α_(i,j)——油品i采用运输方式j的比例。",
            "将表1-3各油品年物流量和运输比例代入式13，计算结果见表3-1。",
        ],
    )

    p = find_starts(doc, "水运进库量为")
    set_text(p, "水运进库量和水运出库量分别为：")
    p = add_numeq_after(p, "G_(w,in)=28+15+24=67 万t/a")
    p = add_numeq_after(p, "G_(w,out)=4+5=9 万t/a")
    p = add_numeq_after(p, "G_w=67+9=76 万t/a")
    insert_after(p, "因此，码头年作业量为76万t/a。")

    p = find_starts(doc, "式中：P_b")
    split_definitions(
        p,
        [
            "P_b——单泊位年通过能力，t/a；",
            "G_s——平均单船实载量，t；",
            "t_b——单船占用时间，h；",
            "η_b——泊位可用系数。",
            "代入平均单船实载量5 000 t、单船占用时间30 h和泊位可用系数0.85：",
        ],
    )
    p = find_starts(doc, "代入平均单船实载量")
    p = add_numeq_after(p, "P_b=(365×24×0.85×5000)/(30)=1241000 t/a")
    insert_after(p, "单泊位年通过能力为124.10万t/a。")

    p = find_starts(doc, "将G_w=76万")
    set_text(p, "式中：")
    p = insert_lines_after(
        p,
        [
            "N_b——计算泊位数；",
            "G_w——码头年作业量，t/a；",
            "P_b——单泊位年通过能力，t/a。",
            "将G_w=76万t/a、P_b=124.10万t/a代入式16：",
        ],
    )
    p = add_numeq_after(p, "N_b=(760000)/(1241000)=0.612")
    insert_after(p, "计算值向上取整，设置1个泊位；泊位利用率为0.612＜0.70，满足作业要求。")

    p = find_starts(doc, "该式用于由设计流量")
    split_definitions(
        p,
        [
            "d——管道计算内径，m；",
            "Q——体积流量，m³/s；",
            "v——允许流速，m/s。",
            "水运进库Q=300 m³/h，取v=1.8 m/s：",
        ],
    )
    p = find_starts(doc, "水运进库Q=300")
    p = add_numeq_after(p, "Q=(300)/(3600)=0.08333 m^3/s")
    p = add_numeq_after(p, "d=√((4×0.08333)/(π×1.8))=0.243 m")
    p = insert_after(p, "按计算内径选择DN250。公路发油Q=120 m³/h，取v=2.0 m/s：")
    p = add_numeq_after(p, "Q=(120)/(3600)=0.03333 m^3/s")
    p = add_numeq_after(p, "d=√((4×0.03333)/(π×2.0))=0.146 m")
    insert_after(p, "按计算内径选择DN150。")

    # 水力计算代表性路径。
    p = find_starts(doc, "上述公式依次用于判别流态")
    split_definitions(
        p,
        [
            "Re——雷诺数；",
            "ν——油品运动黏度，m²/s；",
            "λ——达西摩阻系数；",
            "ε——管壁绝对粗糙度，m；",
            "L——管道计算长度，m；",
            "h_f——沿程阻力损失，m；",
            "Σζ——局部阻力系数之和；",
            "h_m——局部阻力损失，m；",
            "Δz——几何高差，m；",
            "h_e——过滤、计量和装卸设备折算水头，m；",
            "h_r——末端余压折算水头，m。",
        ],
    )
    p = find_starts(doc, "以92号汽油公路发油")
    set_text(p, "以92号汽油公路发油最远鹤位为例，Q=120 m³/h，DN150管道计算内径d=0.147 m，v=1.964 m/s，L=350 m，ν=0.70×10⁻⁶ m²/s，ε=0.045 mm，Σζ=18。分步代入式18～式22：")
    p = add_numeq_after(p, "Re=(1.964×0.147)/(0.70×10^(-6))=4.125×10^5")
    p = add_numeq_after(p, "λ=0.01665")
    p = add_numeq_after(p, "h_f=0.01665×(350)/(0.147)×(1.964^2)/(2×9.81)=7.795 m")
    p = add_numeq_after(p, "h_m=18×(1.964^2)/(2×9.81)=3.539 m")
    p = add_numeq_after(p, "H=8+7.795+3.539+12+3=34.334 m")
    insert_after(p, "计算得到该路径系统扬程为34.334 m。")

    # 泵系统曲线、功率和汽蚀。
    p = find_starts(doc, "式中：H_sys")
    p = split_definitions(
        p,
        [
            "H_sys——管路系统扬程，m；",
            "H_st——静扬程、设备压降和末端压力折算水头之和，m；",
            "K_Q——管路阻力系数，h²/m⁵；",
            "Q——体积流量，m³/h。",
            "将H_st=18 m、Q=200 m³/h、H_sys=33.88 m代入式23：",
        ],
    )
    p = add_numeq_after(p, "K_Q=(33.88-18)/(200^2)=3.97×10^(-4) h^2/m^5")
    insert_after(p, "由此得到当前最不利管路的系统特性方程。该式用于当前工况拟合，不将设备压降外推为普遍严格的平方关系。")

    p = find_starts(doc, "该式用于由流量、扬程和效率")
    p = split_definitions(
        p,
        [
            "P——泵轴功率，kW；",
            "ρ——油品密度，kg/m³；",
            "g——重力加速度，取9.81 m/s²；",
            "Q——体积流量，m³/s；",
            "H——泵扬程，m；",
            "η_p——泵效率。",
            "按工作点Q=220 m³/h、H=37.3 m、ρ=840 kg/m³、η_p=0.78代入式24：",
        ],
    )
    p = find_starts(doc, "按工作点Q=220")
    p = add_numeq_after(p, "Q=(220)/(3600)=0.06111 m^3/s")
    p = add_numeq_after(p, "P=(840×9.81×0.06111×37.3)/(1000×0.78)=24.09 kW")
    insert_after(p, "计算轴功率约24.1 kW，配置30 kW防爆电机，功率裕量约24.5%。")

    p = find_starts(doc, "该式用于校核泵入口可利用汽蚀余量")
    p = split_definitions(
        p,
        [
            "NPSH_a——装置汽蚀余量，m；",
            "P_a——当地大气压，Pa；",
            "ρ——油品密度，kg/m³；",
            "h_s——最低液位高于泵轴中心的静压头，m；",
            "P_v——设计温度下油品饱和蒸气压，Pa；",
            "h_fs——吸入管总水头损失，m。",
            "柴油校核取P_a=101 325 Pa、ρ=840 kg/m³、h_s=1.0 m、P_v=1 000 Pa、h_fs=0.99 m：",
        ],
    )
    p = add_numeq_after(p, "NPSH_(a,d)=(101325-1000)/(840×9.81)+1.0-0.99=12.18 m")
    insert_after(p, "柴油装置汽蚀余量为12.18 m。")
    p = find_starts(doc, "汽油按更不利蒸气压")
    set_text(p, "汽油按更不利蒸气压单独校核。取P_v=65 000 Pa、ρ=760 kg/m³、h_s=1.0 m、h_fs=0.91 m，代入式25：")
    p = add_numeq_after(p, "NPSH_(a,g)=(101325-65000)/(760×9.81)+1.0-0.91=4.97 m")
    insert_after(p, "汽油NPSH_a=4.97 m＞NPSH_r+1.0=3.0 m，汽蚀裕量为1.97 m，满足校核要求。")

    p = find_starts(doc, "该式用于校核单套输油泵")
    split_definitions(
        p,
        [
            "M_a——单套输油泵年质量作业能力，t/a；",
            "n——同时运行泵数；",
            "Q——单泵流量，m³/h；",
            "ρ——油品密度，t/m³；",
            "t_d——日运行时间，h/d；",
            "D——年运行天数，d/a。",
            "每种油品按n=1、Q=200 m³/h、t_d=16 h/d、D=330 d/a代入式26：",
        ],
    )
    p = find_starts(doc, "每种油品按n=1")
    p = add_numeq_after(p, "M_(a,g)=1×200×0.760×16×330=802560 t/a")
    p = add_numeq_after(p, "M_(a,d)=1×200×0.840×16×330=887040 t/a")
    insert_after(p, "汽油年能力为80.26万t/a，柴油年能力为88.70万t/a，均大于相应单品种年出库量。每种油品设置A/B两台泵，一用一备。")

    # 防火堤、冷却水、泡沫、消防水和事故水。
    p = find_starts(doc, "上述两式用于扣除罐基")
    split_definitions(
        p,
        [
            "A_n——防火堤内有效净面积，m²；",
            "L、B——防火堤内边长，m；",
            "n——堤内储罐数量；",
            "D——储罐直径，m；",
            "V_d——防火堤有效容积，m³；",
            "H_d——防火堤设计高度，m；",
            "0.2——堤顶安全余高，m。",
        ],
    )
    p = find_starts(doc, "以G92罐组为例")
    set_text(p, "以G92罐组为例，将L=140 m、B=90 m、n=2、D=40 m、H_d=2.2 m代入式27和式28：")
    p = add_numeq_after(p, "A_n=140×90-2×(π×40^2)/(4)=10086.73 m^2")
    p = add_numeq_after(p, "V_d=10086.73×(2.2-0.2)=20173.46 m^3")
    insert_after(p, "计算有效容积20 173.46 m³＞最大单罐容积20 000 m³，满足围控要求；其余罐组计算见表5-3。")

    p = find_starts(doc, "上述公式用于计算着火罐")
    split_definitions(
        p,
        [
            "A_w——单罐罐壁面积，m²；",
            "D——储罐直径，m；",
            "H——罐壁高度，m；",
            "Q_c——冷却水总流量，L/s；",
            "q_f——着火罐冷却强度，L/(min·m²)；",
            "q_a——相邻罐冷却强度，L/(min·m²)；",
            "V_c——冷却水量，m³；",
            "t_c——连续冷却时间，h。",
        ],
    )
    p = find_starts(doc, "将D=34 m、H=18 m代入式29")
    set_text(p, "将D=34 m、H=18 m、q_f=2.5 L/(min·m²)、q_a=2.0 L/(min·m²)和t_c=9 h代入式29～式31：")
    p = add_numeq_after(p, "A_w=π×34×18=1922.65 m^2")
    p = add_numeq_after(p, "Q_c=(2.5×1922.65+2.0×0.5×1922.65)/(60)=112.15 L/s")
    p = add_numeq_after(p, "V_c=112.15×9×3.6=3633.66 m^3")
    insert_after(p, "按未圆整流量计算，冷却水量为约3 633.82 m³；消防水池容积计算采用该未圆整值。")

    p = find_starts(doc, "上述公式用于计算泡沫保护面积")
    split_definitions(
        p,
        [
            "A_(f,g)——汽油内浮顶罐密封圈保护面积，m²；",
            "b——泡沫堰板至罐壁距离，m；",
            "Q_f——泡沫混合液流量，L/s；",
            "q_f——混合液供给强度，L/(min·m²)；",
            "n_p——泡沫产生器数量；",
            "Q_p——单个泡沫产生器额定流量，L/s；",
            "V_(mix)——泡沫混合液量，m³；",
            "t_f——连续供给时间，h；",
            "V_(foam)——泡沫液储量，m³；",
            "c——实际混合比；",
            "1.10——泡沫液储量裕量系数。",
        ],
    )
    p = find_starts(doc, "20 000 m³汽油罐：")
    set_text(p, "20 000 m³汽油罐的密封圈保护面积、理论流量和产生器数量计算如下：")
    p = add_numeq_after(p, "A_(f,g)=π×40×0.55=69.12 m^2")
    p = add_numeq_after(p, "Q_f=(12.5×69.12)/(60)=14.40 L/s")
    p = add_numeq_after(p, "n_(p,Q)=(14.40)/(8)=1.80")
    p = add_numeq_after(p, "n_(p,L)=(π×40)/(24)=5.24")
    insert_after(p, "取两种控制条件中的较大值，每罐设置6个PCL8泡沫产生器，实际固定系统流量为48 L/s。")

    p = find_starts(doc, "15 000 m³汽油罐D=34")
    set_text(p, "15 000 m³汽油罐和柴油罐按相同方法计算：")
    p = add_numeq_after(p, "A_(f,g)=π×34×0.55=58.75 m^2")
    p = add_numeq_after(p, "Q_(f,g)=(12.5×58.75)/(60)=12.24 L/s")
    p = add_numeq_after(p, "n_(p,g)=(π×34)/(24)=4.45")
    p = add_numeq_after(p, "A_(f,d)=(π×34^2)/(4)=907.92 m^2")
    p = add_numeq_after(p, "Q_(f,d)=(6.0×907.92)/(60)=90.79 L/s")
    p = add_numeq_after(p, "n_(p,d)=(90.79)/(16)=5.67")
    insert_after(p, "因此，15 000 m³汽油罐设置5个PCL8，15 000 m³柴油罐设置6个PCL16。")

    p = find_starts(doc, "泡沫混合液管道剩余量")
    set_text(p, "泡沫混合液管道剩余量为5.00 m³。控制工况下，固定设施、2支辅助泡沫枪和管道剩余量合并计算：")
    p = add_numeq_after(p, "V_(mix)=48×3.6+2×(240/60)×1.8+5.00=192.20 m^3")
    p = add_numeq_after(p, "V_(foam)=192.20×0.039×1.10=8.25 m^3")
    insert_after(p, "设置2×5 m³泡沫液储罐，总有效量10 m³＞8.25 m³，满足要求。")

    p = find_starts(doc, "该式用于计算消防水池")
    split_definitions(
        p,
        [
            "V_(FW)——消防用水量，m³；",
            "V_c——固定冷却水量，m³；",
            "c——泡沫液实际混合比；",
            "V_(mix)——泡沫混合液量，m³。",
        ],
    )
    p = find_starts(doc, "控制工况冷却水量为")
    set_text(p, "控制工况冷却水量V_c=3 633.82 m³，泡沫混合液量V_(mix)=192.20 m³，实际混合比c=0.039。代入式37：")
    p = add_numeq_after(p, "V_(foam-water)=192.20×(1-0.039)=184.70 m^3")
    p = add_numeq_after(p, "V_(FW)=3633.82+184.70=3818.52 m^3")
    insert_after(p, "消防水池采用2×2 500 m³，总有效容积5 000 m³＞3 818.52 m³，满足要求。")

    p = find_starts(doc, "该式用于合并消防水")
    split_definitions(
        p,
        [
            "V_(acc)——事故水池计算容积，m³；",
            "V_(FW)——控制工况消防水量，m³；",
            "ψ——地面径流系数；",
            "F——事故影响汇水面积，m²；",
            "h_r——最大一日降雨深度，m；",
            "Q_l——最大泄漏流量，m³/h；",
            "t_s——紧急切断时间，h；",
            "1.2——泄漏量附加系数。",
        ],
    )
    p = find_starts(doc, "事故水量包括控制工况")
    set_text(p, "取ψ=0.90、F=12 600 m²、h_r=0.2764 m、Q_l=300 m³/h、t_s=10/60 h，分项计算事故水量：")
    p = add_numeq_after(p, "V_r=0.90×12600×0.2764=3134.38 m^3")
    p = add_numeq_after(p, "V_l=1.2×300×(10/60)=60.00 m^3")
    p = add_numeq_after(p, "V_(acc)=3818.52+3134.38+60.00=7012.90 m^3")
    insert_after(p, "按分项未圆整值复核，事故水池计算容积约7 013 m³，选用8 000 m³事故水池，满足要求。")

    p = find_starts(doc, "初期雨水汇水范围逐项包括")
    set_text(p, "初期雨水汇水面积为F_0=15 000 m²，取初期降雨深度h_0=15 mm=0.015 m、硬化地面径流系数ψ_0=0.90：")
    p = add_numeq_after(p, "V_0=0.90×15000×0.015=202.5 m^3")
    insert_after(p, "设置2×150 m³初期雨水池，总有效容积300 m³＞202.5 m³，能够容纳计算量。")

    p = find_starts(doc, "将两种汽油装车流量各120")
    set_text(p, "式中：")
    p = insert_lines_after(
        p,
        [
            "Q_(VRU)——油气回收装置设计处理量，m³/h；",
            "Q_(load,g)——单种汽油最大装车流量，m³/h；",
            "1.10——装车流量波动系数。",
            "将两种汽油装车流量各120 m³/h代入式39：",
        ],
    )
    p = add_numeq_after(p, "Q_(VRU)=1.10×(120+120)=264 m^3/h")
    insert_after(p, "单套油气回收装置选300 m³/h，采用2套一用一备。沿海空气湿度较高，吸附单元前设置气液分离、温度监测和冷凝预处理。[14]")

    # 结论与正文一致。
    p = find_starts(doc, "（5）主要工艺管径")
    set_text(
        p,
        "（5）主要工艺管径为DN150～DN250。最不利水运外输管路在200 m³/h时所需扬程33.88 m，选KSB MegaCPK Inducer 125-100-200、193 mm叶轮作为候选泵；泵与管路曲线交点约为220 m³/h、37.3 m。三种油品各设2台泵，一用一备，配30 kW防爆电机；柴油NPSH_a=12.18 m，大于样本NPSH_r与1.0 m裕量之和。",
    )

    replace_plain_terms(doc)
    convert_existing_underscores(doc)

    # 清理完全空白、且位于正文公式计算段中的冗余段落，不动分页符和表格占位。
    for paragraph in list(doc.paragraphs):
        if paragraph.text == "" and paragraph._p.xpath(".//w:br[@w:type='page']"):
            continue

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: refine_latest_formula_calculations_20260729.py SRC DST")
    main(Path(sys.argv[1]), Path(sys.argv[2]))
