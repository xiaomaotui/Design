from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\毕业论文")
SOURCE = ROOT / r"04_最终成品\01_毕业设计说明书\张淑鑫_浙江平湖二级油库工艺设计_毕业设计初稿_最低平均温度修正版_2026-07-29.docx"
OUTPUT = ROOT / r"04_最终成品\01_毕业设计说明书\张淑鑫_浙江平湖二级油库工艺设计_毕业设计初稿_导师审查修订版_2026-07-29.docx"
ASSETS = ROOT / r"03_过程文件\02_设计计算\supervisor_revision_assets"

RHO = {"92号汽油": 0.760, "95号汽油": 0.760, "0号柴油": 0.840}
NU = {"92号汽油": 0.70e-6, "95号汽油": 0.70e-6, "0号柴油": 3.50e-6}
PIPE = {
    250: {"od": 0.273, "wall": 0.008, "d": 0.257},
    200: {"od": 0.219, "wall": 0.007, "d": 0.205},
    150: {"od": 0.159, "wall": 0.006, "d": 0.147},
    100: {"od": 0.108, "wall": 0.005, "d": 0.098},
}
PATHS = [
    ("92号汽油水运进库", "92号汽油", 300, 250, 1200, 8, 8, 16),
    ("95号汽油水运进库", "95号汽油", 300, 250, 1200, 8, 8, 16),
    ("0号柴油水运进库", "0号柴油", 300, 250, 1200, 8, 8, 16),
    ("92号汽油管道进库", "92号汽油", 250, 250, 800, 5, 6, 14),
    ("95号汽油管道进库", "95号汽油", 250, 250, 800, 5, 6, 14),
    ("0号柴油管道进库", "0号柴油", 250, 250, 800, 5, 6, 14),
    ("92号汽油公路发油", "92号汽油", 120, 150, 350, 8, 12, 18),
    ("95号汽油公路发油", "95号汽油", 120, 150, 350, 8, 12, 18),
    ("0号柴油公路发油", "0号柴油", 120, 150, 350, 8, 12, 18),
    ("92号汽油管道外输", "92号汽油", 200, 200, 1000, 6, 8, 16),
    ("95号汽油管道外输", "95号汽油", 200, 200, 1000, 6, 8, 16),
    ("0号柴油管道外输", "0号柴油", 200, 200, 1000, 6, 8, 16),
    ("92号汽油水运外输", "92号汽油", 200, 200, 1200, 5, 10, 18),
    ("95号汽油水运外输", "95号汽油", 200, 200, 1200, 5, 10, 18),
]


def hyd(product: str, q_m3h: float, dn: int, length: float, dz: float, equipment: float, local_k: float):
    d = PIPE[dn]["d"]
    q = q_m3h / 3600.0
    v = 4 * q / (math.pi * d**2)
    re = v * d / NU[product]
    eps = 0.000045
    lam = 0.25 / math.log10(eps / (3.7 * d) + 5.74 / (re**0.9)) ** 2
    hf = lam * length / d * v**2 / (2 * 9.81)
    hm = local_k * v**2 / (2 * 9.81)
    head = dz + equipment + 3.0 + hf + hm
    power = RHO[product] * 1000 * 9.81 * q * head / (0.70 * 1000)
    return dict(d=d, q=q, v=v, re=re, lam=lam, hf=hf, hm=hm, head=head, power=power)


HYD = {x[0]: hyd(x[1], x[2], x[3], x[4], x[5], x[6], x[7]) for x in PATHS}


def sci(x: float) -> str:
    if x == 0:
        return "0"
    exp = int(math.floor(math.log10(abs(x))))
    mant = x / 10**exp
    sup = str(exp).translate(str.maketrans("0123456789-", "⁰¹²³⁴⁵⁶⁷⁸⁹⁻"))
    return f"{mant:.3f}×10{sup}"


def set_run(run, size=10.5, bold=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rf = rpr.rFonts
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:eastAsia"), "宋体")
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")


def set_text(p, text: str):
    p.clear()
    set_run(p.add_run(text))


def find_para(doc, text: str):
    for p in doc.paragraphs:
        if text in p.text:
            return p
    raise KeyError(text)


def insert_paragraph_after(anchor, text: str):
    p = OxmlElement("w:p")
    anchor._p.addnext(p)
    from docx.text.paragraph import Paragraph

    out = Paragraph(p, anchor._parent)
    out.style = anchor.style
    out.paragraph_format.first_line_indent = Cm(0.74)
    out.paragraph_format.line_spacing = 1.5
    out.paragraph_format.space_after = Pt(0)
    set_run(out.add_run(text))
    return out


def set_cell(cell, text: str, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    set_run(p.add_run(text), size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def insert_table_after(doc, anchor, title: str, headers, rows, widths=None):
    cap = insert_paragraph_after(anchor, title)
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        set_cell(table.rows[0].cells[j], h, 8.5)
        table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            set_cell(cells[j], str(value), 8.5)
    if widths:
        for row in table.rows:
            for j, width in enumerate(widths):
                row.cells[j].width = Cm(width)
    cap._p.addnext(table._tbl)
    return table


def replace_in_all_parts(doc, old, new):
    for p in doc.paragraphs:
        if old in p.text:
            set_text(p, p.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_text(p, p.text.replace(old, new))
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                if old in p.text:
                    set_text(p, p.text.replace(old, new))


def build_system_curve():
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / "输油泵与最不利管路特性曲线_实际内径修正.png"
    q_p = [100, 150, 200, 250, 300, 330]
    h_p = [42.0, 41.0, 38.8, 35.0, 29.0, 25.0]
    h_static = 18.0
    h_design = HYD["92号汽油水运外输"]["head"]
    k_sys = (h_design - h_static) / 200.0**2
    q_values = list(range(0, 341, 5))
    h_sys = [h_static + k_sys * q**2 for q in q_values]
    # Piecewise interpolation is adequate for a manufacturer-curve reading.
    best = (999, 0, 0)
    for q in [x / 10 for x in range(1000, 3301)]:
        for a, b, ha, hb in zip(q_p[:-1], q_p[1:], h_p[:-1], h_p[1:]):
            if a <= q <= b:
                hp = ha + (hb - ha) * (q - a) / (b - a)
                hs = h_static + k_sys * q**2
                if abs(hp - hs) < best[0]:
                    best = (abs(hp - hs), q, (hp + hs) / 2)
                break
    op_q, op_h = best[1], best[2]
    canvas = Image.new("RGB", (1600, 1050), "white")
    draw = ImageDraw.Draw(canvas)
    font_path = Path(r"C:\Windows\Fonts\simhei.ttf")
    latin_path = Path(r"C:\Windows\Fonts\times.ttf")
    font = ImageFont.truetype(str(font_path if font_path.exists() else latin_path), 32)
    small = ImageFont.truetype(str(font_path if font_path.exists() else latin_path), 25)
    left, top, right, bottom = 150, 100, 1510, 900
    qmax, hmin, hmax = 340.0, 15.0, 48.0
    xp = lambda x: left + x / qmax * (right - left)
    yp = lambda y: bottom - (y - hmin) / (hmax - hmin) * (bottom - top)
    draw.line((left, top, left, bottom), fill="black", width=4)
    draw.line((left, bottom, right, bottom), fill="black", width=4)
    for value in range(0, 341, 50):
        x = xp(value)
        draw.line((x, top, x, bottom), fill=(215, 215, 215), width=2)
        draw.text((x - 18, bottom + 18), str(value), fill="black", font=small)
    for value in range(15, 49, 5):
        y = yp(value)
        draw.line((left, y, right, y), fill=(215, 215, 215), width=2)
        draw.text((70, y - 16), str(value), fill="black", font=small)
    draw.line([(xp(x), yp(y)) for x, y in zip(q_p, h_p)], fill=(0, 91, 172), width=7)
    draw.line([(xp(x), yp(y)) for x, y in zip(q_values, h_sys)], fill=(205, 66, 43), width=7)
    x, y = xp(op_q), yp(op_h)
    draw.ellipse((x - 11, y - 11, x + 11, y + 11), fill="black")
    draw.text((x + 18, y - 55), f"工作点 Q={op_q:.0f} m³/h，H={op_h:.1f} m", fill="black", font=small)
    draw.text((600, 955), "流量 Q/(m³/h)", fill="black", font=font)
    draw.text((25, 40), "扬程 H/m", fill="black", font=font)
    draw.line((260, 150, 370, 150), fill=(0, 91, 172), width=7)
    draw.text((390, 132), "KSB 125-100-200，叶轮193 mm", fill="black", font=small)
    draw.line((260, 200, 370, 200), fill=(205, 66, 43), width=7)
    draw.text((390, 182), "汽油水运外输控制管路", fill="black", font=small)
    canvas.save(path)
    return path, op_q, op_h


def main():
    doc = Document(SOURCE)

    # Global factual/editorial corrections.
    replace_in_all_parts(doc, "2026届", "2027届")
    replace_in_all_parts(doc, "年平均相对湿度80%", "年平均相对湿度78%")
    replace_in_all_parts(doc, "0.615", "0.612")
    replace_in_all_parts(doc, "0.613", "0.612")
    replace_in_all_parts(doc, "满足初稿设计能力要求", "满足本设计能力要求")
    replace_in_all_parts(doc, "本稿的6座罐", "本设计的6座罐")

    # K=14: keep the user's choice and explain the change from the initial plan.
    p = find_para(doc, "《油库技术与管理》第2版给出的商业系统参考范围")
    set_text(
        p,
        "《油库技术与管理》第2版给出的商业系统参考范围为：二级及以上商业油库年周转系数K取8～14，三级及以下油库K取14～24；多油品油库可先按相应区间不利值计算，再结合油库等级与运输条件反算。本设计为二级商业油库，原计划表中的K=20属于方案初拟阶段参数，已超出二级及以上商业油库8～14的参考区间，故正文按该教材上限取K=14。该取值使理论库容计算偏安全；储罐配置完成后反算三种油品K分别为13.85、11.54和12.53，均落在8～14内，因此K=14作为本设计的正式计算参数。[17]",
    )

    # Correct actual pipe inside diameters and recalculate the main text.
    p = find_para(doc, "主干管按表3-2峰值流量确定")
    set_text(
        p,
        "主干管按表3-2峰值流量确定，不按年平均流量选管。汽油和柴油主输管流速控制在约1.0～2.5 m/s，泵吸入管宜低于1.5 m/s。[18] 管径校核不再把DN数值直接当作内径：本设计按GB/T 8163—2018无缝钢管预选DN250为Φ273×8 mm、DN200为Φ219×7 mm、DN150为Φ159×6 mm，相应计算内径分别为257 mm、205 mm和147 mm；采购规格改变时应按实际外径和壁厚复算。",
    )
    set_text(
        find_para(doc, "以92号汽油公路发油最远鹤位为例"),
        "以92号汽油公路发油最远鹤位为例：Q=120 m³/h，DN150预选Φ159×6 mm，计算内径d=0.147 m，流速v=1.964 m/s，L=350 m，运动黏度ν=0.70×10⁻⁶ m²/s，绝对粗糙度ε=0.045 mm。代入式18得Re=4.125×10⁵，代入式19得λ=0.01665；代入式20得沿程损失h_f=7.795 m。局部阻力系数Σζ=18，代入式21得h_m=3.539 m。高差8 m、过滤计量与鹤管压降12 m、末端余压3 m，代入式22得H=34.334 m。",
    )
    set_text(
        find_para(doc, "柴油水运进库采用DN250"),
        "柴油水运进库采用DN250、预选Φ273×8 mm、计算内径257 mm，Q=300 m³/h时v=1.606 m/s，Re=1.180×10⁵，λ=0.01835，沿程损失11.271 m、局部损失2.104 m，总需要扬程32.375 m。该扬程作为船泵或岸上助推泵的接口要求。",
    )

    # Main hydraulic tables.
    for row, path in zip(doc.tables[12].rows[1:], [PATHS[0], PATHS[3], PATHS[6], PATHS[9], PATHS[12]]):
        h = HYD[path[0]]
        set_cell(row.cells[3], f"{h['v']:.2f}")
    for i, path in enumerate(PATHS, 1):
        name, product, q, dn, length, dz, equipment, local_k = path
        h = HYD[name]
        vals = [name, f"{h['v']:.2f}", sci(h["re"]), f"{h['lam']:.4f}", f"{h['hf']:.2f}", f"{h['hm']:.2f}", f"{h['head']:.2f}", f"{h['power']:.1f}"]
        for j, value in enumerate(vals):
            set_cell(doc.tables[13].rows[i].cells[j], value)

    # Pump/system curve and consistent pump data.
    curve_path, op_q, op_h = build_system_curve()
    set_text(
        find_para(doc, "库内输油泵服务于倒罐"),
        "库内输油泵服务于倒罐、公共管道外输和水运外输；水运进库主要由船泵提供压力。采用实际内径复算后，库内最不利路径仍为汽油水运外输，设计流量200 m³/h、系统所需扬程33.88 m，故泵额定校核点取200 m³/h、约34 m；设备表统一采用200 m³/h、36 m、30 kW，36 m额定扬程可覆盖计算扬程并留有运行调节余量。",
    )
    set_text(
        find_para(doc, "为确定工作点，将最不利水运外输管路"),
        "为确定工作点，将最不利水运外输管路的系统扬程写成静扬程与流量平方项之和。该路径静扬程由高差5 m、设备压降10 m和末端余压3 m组成，即H_st=18 m；在Q=200 m³/h时沿程与局部损失合计15.88 m，据此反算管路阻力系数。",
    )
    set_text(
        find_para(doc, "式中：H_sys——管路系统扬程"),
        "式中：H_sys——管路系统扬程，m；H_st——高差、设备压降和末端压力折算水头之和，m；K_Q——在本设计计算范围内采用的管路阻力系数，h²/m⁵；Q——体积流量，m³/h。将H_st=18 m、Q=200 m³/h、H_sys=33.88 m代入式23，得K_Q=3.97×10⁻⁴ h²/m⁵。本式用于当前工况的工程拟合，不把设备压降外推为普遍严格的平方关系。",
    )
    set_text(
        find_para(doc, "图4-2表明"),
        f"图4-2表明，193 mm叶轮泵曲线与按实际内径修正后的系统曲线交于Q≈{op_q:.0f} m³/h、H≈{op_h:.1f} m，工作点位于制造商曲线允许范围内。额定设计仍按200 m³/h控制，通过变频把实际运行点调回作业所需流量；低流量公路装车工况设置最小流量回流，避免长期偏离允许工作区。",
    )
    caption = find_para(doc, "图4-2 输油泵性能曲线")
    pic_para = caption._p.getprevious()
    from docx.text.paragraph import Paragraph

    pp = Paragraph(pic_para, caption._parent)
    pp.clear()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.add_run().add_picture(str(curve_path), width=Cm(15.2))
    set_text(
        find_para(doc, "该式用于由流量、扬程和效率计算泵轴功率"),
        f"该式用于由流量、扬程和效率计算泵轴功率。式中：P——泵轴功率，kW；ρ——油品密度，kg/m³；g——重力加速度，取9.81 m/s²；Q——体积流量，m³/s；H——扬程，m；η_p——泵效率。按修正工作点Q={op_q:.0f} m³/h、H={op_h:.1f} m、柴油密度ρ=840 kg/m³、样本工作点效率η_p=0.78代入式24，得轴功率约24.0 kW。配置30 kW防爆电机，功率裕量约25%；最终以制造商针对实际介质出具的确认曲线为准。",
    )

    # Numeric NPSH check for gasoline and diesel.
    set_text(
        find_para(doc, "该式用于校核泵入口可利用汽蚀余量"),
        "该式用于校核泵入口可利用汽蚀余量。式中：NPSH_a——装置汽蚀余量，m；P_a——当地大气压，Pa；ρ——油品密度，kg/m³；h_s——最低液位高于泵轴中心的静压头，m；P_v——设计温度下饱和蒸气压，Pa；h_fs——吸入管总水头损失，m。柴油校核取P_a=101 325 Pa、ρ=840 kg/m³；储罐低低液位1.5 m、泵轴高于罐底0.5 m，故h_s=1.0 m；40 ℃柴油饱和蒸气压取1.0 kPa。按修正工作点Q=220 m³/h、DN250计算内径257 mm、吸入管长80 m、Σζ=8，h_fs=0.99 m，代入式25得NPSH_a=12.18 m。",
    )
    p = find_para(doc, "由KSB样本第18页NPSH_r曲线读取")
    set_text(
        p,
        "由KSB样本第18页NPSH_r曲线读取，Q≈220 m³/h时NPSH_r≈2.0 m。按1.0 m附加裕量校核，柴油NPSH_a=12.18 m>NPSH_r+1.0=3.0 m，满足要求。",
    )
    insert_paragraph_after(
        p,
        "汽油按更不利蒸气压单独校核。依据GB 17930—2016车用汽油蒸气压限值，取P_v=65 kPa，ρ=760 kg/m³，其余吸入标高和管路条件不变；汽油h_fs=0.91 m。代入式25，NPSH_a=(101 325-65 000)/(760×9.81)+1.0-0.91=4.97 m。与样本NPSH_r≈2.0 m并计1.0 m附加裕量比较，4.97 m>3.0 m，汽油汽蚀裕量为1.97 m，满足本阶段校核要求。若最高油温或实际蒸气压提高，应据实复算。",
    )

    # Loading arms, minimum flow and headers.
    p = find_para(doc, "公路装车每种油品2个鹤位")
    set_text(
        p,
        "公路装车鹤位按单鹤位60 m³/h、16 h/d、330 d/a计算。单个汽油鹤位年质量能力为60×0.760×16×330=24.08万t/a，92号汽油公路出库量26万t/a，需要2个同时工作鹤位；95号汽油17.5万t/a需要1个工作鹤位，另1个用于备用和高峰调度。单个柴油鹤位年质量能力为26.61万t/a，18万t/a需要1个工作鹤位，另1个备用。因此三种油品各设2个鹤位具有明确能力依据。",
    )
    p2 = insert_paragraph_after(
        p,
        "单鹤位作业时泵流量为60 m³/h，恰为200 m³/h额定流量的30%。KSB现有样本未给出最小连续稳定流量数值，故本设计不把“30%”写成制造商定值；控制系统设置可调最小流量回流支路，采购时按制造商确认的Q_min整定。当Q_min≤60 m³/h时单鹤位可直接运行；当Q_min>60 m³/h时，回流量按Q_r=Q_min-60 m³/h确定。",
    )
    insert_paragraph_after(
        p2,
        "装车管路逐级校核如下：单鹤位支管Q=60 m³/h，预选DN100（Φ108×5 mm，d=98 mm），流速2.21 m/s；两鹤位同时作业的集油管Q=120 m³/h，预选DN150（Φ159×6 mm，d=147 mm），流速1.96 m/s；泵吸入总管Q=200 m³/h，DN250（d=257 mm），流速1.07 m/s；泵排出总管Q=200 m³/h，DN200（d=205 mm），流速1.68 m/s。支管、集油管和泵口管径由同一峰值工况贯通，未把单鹤位流量误用于总管。",
    )

    # Remove editorial/workflow traces.
    set_text(
        find_para(doc, "后续若总平面图改变管长或高差"),
        "后续若总平面图改变管长或高差，应同步更新表4-2、系统曲线、泵工作点和NPSH；若改变罐径或罐高，应同步更新罐壁、抗风圈、防火堤、冷却水和泡沫计算。说明书、计算表和两张图纸应始终使用同一版本参数。",
    )

    # Foam system: mandatory circumference, auxiliary guns and pipe residual.
    set_text(
        find_para(doc, "汽油内浮顶罐采用密封圈环形保护"),
        "汽油钢制单盘式或双盘式内浮顶罐采用密封圈环形保护。依据GB 50151—2021第4.4.2条，泡沫堰板距罐壁b=0.55 m，混合液强度12.5 L/(min·m²)，连续供给60 min，且单个泡沫产生器保护周长不大于24 m；柴油固定顶罐采用全液面保护，混合液强度6.0 L/(min·m²)，连续供给30 min。固定式系统另按第4.1.5条配置2支辅助泡沫枪，单枪流量240 L/min，连续供给30 min；系统设计量计入固定设施、辅助泡沫枪和管道剩余量。",
    )
    set_text(
        find_para(doc, "20 000 m³汽油罐：将D=40 m"),
        "20 000 m³汽油罐：将D=40 m、b=0.55 m代入式32，A_f=69.12 m²；代入式33，理论流量Q_f=14.40 L/s。按流量需2个8 L/s产生器，但罐周长πD=125.66 m，按单个保护周长不大于24 m计算需ceil(125.66/24)=6个，故每罐设置6个PCL8，实际固定系统流量48 L/s。",
    )
    p = find_para(doc, "15 000 m³柴油罐：全液面面积")
    set_text(
        p,
        "15 000 m³汽油罐D=34 m，密封圈面积57.81 m²，理论流量12.04 L/s；罐周长106.81 m，按24 m限值需5个PCL8，实际固定系统流量40 L/s。15 000 m³柴油罐全液面面积907.92 m²，理论流量90.79 L/s，按流量设置6个PCL16，实际固定系统流量96 L/s。",
    )
    insert_paragraph_after(
        p,
        "泡沫混合液管道剩余量按初步路由计算：DN150主管180 m容积3.05 m³，DN100分配管及支管合计约1.45 m³，比例混合装置和阀组内滞液量按0.50 m³计，合计5.00 m³。20 000 m³汽油罐工况混合液总量为48×3.6+2×240/60×1.8+5.00=192.20 m³；柴油罐工况同为192.20 m³。按实际混合比上限3.9%并增加10%储量，泡沫液需要量为192.20×0.039×1.10=8.25 m³，设置2×5 m³泡沫液储罐，总有效量10 m³，满足要求。",
    )
    t = doc.tables[22]
    # Add one row so all three protected tanks are visible.
    if len(t.rows) == 3:
        t.add_row()
    foam_rows = [
        ["20 000 m³汽油罐密封圈", "69.12", "12.5 L/(min·m²)", "14.40/48.00", "60", "PCL8", "6", "8.25"],
        ["15 000 m³汽油罐密封圈", "57.81", "12.5 L/(min·m²)", "12.04/40.00", "60", "PCL8", "5", "7.01"],
        ["15 000 m³柴油罐全液面", "907.92", "6.0 L/(min·m²)", "90.79/96.00", "30", "PCL16", "6", "8.25"],
    ]
    for i, row in enumerate(foam_rows, 1):
        for j, value in enumerate(row):
            set_cell(t.rows[i].cells[j], value, 8)
    set_cell(t.rows[0].cells[3], "理论/实际Q_f/(L/s)", 8)

    # Fire-water volume, flow and preliminary head calculation.
    set_text(
        find_para(doc, "将柴油冷却水量3633.82 m³"),
        "控制工况冷却水量为3 633.82 m³。泡沫系统混合液总量按固定设施、2支辅助泡沫枪和管道剩余量合计192.20 m³，其中水量为192.20×(1-0.039)=184.70 m³。代入式37，消防用水量V_FW=3 633.82+184.70=3 818.52 m³。消防水池采用2×2 500 m³，总有效容积5 000 m³，满足计算需要。",
    )
    set_text(
        find_para(doc, "固定冷却水泵按Q_c=112.15 L/s"),
        "固定冷却水泵按Q_c=112.15 L/s配置130 L/s工作泵；泡沫工况固定产生器流量96 L/s、辅助泡沫枪流量8 L/s，合计104 L/s，再按GB 50151—2021第8.1.6条计5%流量裕量，设计流量为109.2 L/s，泡沫消防水泵选120 L/s。冷却水系统与泡沫系统分设泵组和管道。",
    )
    set_text(
        find_para(doc, "消防冷却水泵设置2台电动泵"),
        "冷却水泵组设置1台130 L/s电动工作泵、1台同能力电动备用泵和1台同能力柴油机应急备用泵；泡沫水泵组设置1台120 L/s电动工作泵和1台120 L/s柴油机备用泵。二级石油库采用电动主泵、柴油机备用泵符合GB 50151—2021第7.1.3条的配置路径。各泵独立吸水，并设置试验回流、就地手动启动和远程启动；备用泵流量、扬程不低于工作泵。",
    )
    p = find_para(doc, "最不利点不是简单取离泵房最远的消火栓")
    insert_paragraph_after(
        p,
        "消防泵扬程按式22的组成逐项校核。冷却水控制点取罐顶高差18 m、环网与支路沿程及局部损失15 m、喷头入口所需压力水头20 m，计算扬程53 m，冷却水泵初选H=60 m。泡沫控制点取罐顶高差18 m、主管及支管损失20 m、比例混合装置损失10 m、PCL产生器额定进口压力0.50 MPa折合水头51 m，计算扬程99 m，泡沫水泵初选H=110 m。PCL额定进口压力取值来自产品样本，采购型号变化时应按实际额定压力复核，不把该产品参数写成国家标准定值。",
    )
    for table_index in (23, 44):
        tt = doc.tables[table_index]
        if table_index == 23:
            vals = [
                ["固定冷却水", "1台130 L/s、60 m电动泵", "1台电动备用+1台柴油机备用", "≥112.15 L/s", "流量及最不利点压力"],
                ["泡沫供水", "1台120 L/s、110 m电动泵", "1台120 L/s柴油机备用泵", "≥109.2 L/s", "比例混合及产生器入口压力"],
            ]
            for i, row in enumerate(vals, 1):
                for j, value in enumerate(row):
                    set_cell(tt.rows[i].cells[j], value, 8)
        else:
            vals = [
                ["固定冷却水泵", "130 L/s，60 m", "3", "1用+2备（含柴油机）", "130>112.15 L/s"],
                ["泡沫水泵", "120 L/s，110 m", "2", "一用一备（柴油机备用）", "120>109.2 L/s"],
                ["消防水池", "2 500 m³", "2", "分格连通", "5 000>3 818.52 m³"],
                ["泡沫液储罐", "5 m³", "2", "并联", "10>8.25 m³"],
                ["事故水池", "8 000 m³", "1", "事故状态启用", "8 000>6 985.25 m³"],
            ]
            for i, row in enumerate(vals, 1):
                for j, value in enumerate(row):
                    set_cell(tt.rows[i].cells[j], value, 8)

    # First-flush calculation: preserve existing design result, show every input.
    set_text(
        find_para(doc, "初期雨水量为202.5 m³"),
        "初期雨水汇水范围逐项包括公路装车区6 000 m²、泵阀及计量区4 500 m²、码头前沿污染控制区2 500 m²和连接道路硬化区2 000 m²，合计F_0=15 000 m²。初期降雨深度取h_0=15 mm=0.015 m，硬化地面径流系数ψ_0=0.90，则初期雨水量为0.90×15 000×0.015=202.5 m³。设置2×150 m³初期雨水池，总有效容积300 m³，能够容纳计算量；15 mm为本设计运行分流控制值，不表述为GB 50074强制数值。",
    )

    # Conclusion and appendix D.
    replace_in_all_parts(doc, "泵与管路曲线交点约为202 m³/h、36.3 m", f"泵与管路曲线交点约为{op_q:.0f} m³/h、{op_h:.1f} m")
    replace_in_all_parts(doc, "柴油NPSH_a=11.83 m", "柴油NPSH_a=12.18 m")
    replace_in_all_parts(doc, "柴油罐配置6个16 L/s泡沫产生器，泡沫液计算量约7.01 m³", "柴油罐配置6个16 L/s泡沫产生器；计入辅助泡沫枪和管道剩余量后，控制泡沫液量约8.25 m³")
    replace_in_all_parts(doc, "消防水计算量约3 790 m³", "消防水计算量约3 819 m³")
    set_text(
        find_para(doc, "汽油罐密封圈保护面积69.12 m²"),
        "20 000 m³汽油罐密封圈保护面积69.12 m²，理论混合液流量14.40 L/s，按单个产生器保护周长不大于24 m配置6个PCL8；15 000 m³汽油罐配置5个PCL8。15 000 m³柴油罐全液面面积907.92 m²，理论流量90.79 L/s，配置6个PCL16。计入2支辅助泡沫枪和5.00 m³管道剩余量后，控制工况泡沫液储量8.25 m³。",
    )

    # Appendix C: all fourteen paths use actual inside diameter and the same pump conclusion.
    paras = doc.paragraphs
    for idx, path in enumerate(PATHS, 1):
        name, product, q, dn, length, dz, equipment, local_k = path
        h = HYD[name]
        heading = next(p for p in paras if p.text.strip() == f"C.{idx} {name}")
        pos = paras.index(heading)
        block = []
        for candidate in paras[pos + 1 :]:
            if candidate.style.name.startswith("标题") or candidate.style.name.startswith("Heading"):
                break
            if candidate.text.strip():
                block.append(candidate)
            if len(block) == 7:
                break
        od_mm = PIPE[dn]["od"] * 1000
        wall_mm = PIPE[dn]["wall"] * 1000
        d_mm = PIPE[dn]["d"] * 1000
        texts = [
            f"介质为{product}，Q={q} m³/h，DN{dn}预选Φ{od_mm:.0f}×{wall_mm:.0f} mm，计算内径d={d_mm:.0f} mm，L={length} m，Δz={dz} m，设备压降{equipment} m，Σζ={local_k}。计算流速v={h['v']:.3f} m/s，Re={sci(h['re'])}，λ={h['lam']:.5f}，沿程损失h_f={h['hf']:.3f} m，局部损失h_m={h['hm']:.3f} m，系统扬程H={h['head']:.3f} m，轴功率P={h['power']:.2f} kW。",
            f"步骤1：将体积流量换算为SI单位，Q={q}/3600={h['q']:.6f} m³/s；由外径与壁厚计算内径d={od_mm:.0f}-2×{wall_mm:.0f}={d_mm:.0f} mm={h['d']:.3f} m。",
            f"步骤2：代入式17的流速关系v=4Q/(πd²)，得v={h['v']:.3f} m/s。",
            f"步骤3：取{product}运动黏度ν={NU[product]:.2e} m²/s，代入式18，得Re={sci(h['re'])}，判定为湍流。",
            f"步骤4：取商业钢管绝对粗糙度ε=4.5×10⁻⁵ m，代入式19，得Darcy摩阻系数λ={h['lam']:.5f}。",
            f"步骤5：代入式20和式21，得h_f={h['hf']:.3f} m、h_m={h['hm']:.3f} m；再按式22加高差{dz} m、设备压降{equipment} m和末端余压3 m，得H={h['head']:.3f} m。",
            f"步骤6：按式24、泵效率η_p=0.70计算轴功率P={h['power']:.2f} kW。本路径按实际内径复算后，与正文统一采用200 m³/h、36 m、30 kW的KSB候选泵校核；水运进库路径由船泵提供接口压力，库内泵参数以表4-3为准。",
        ]
        for target, text in zip(block, texts):
            set_text(target, text)
        row = doc.tables[29 + idx].rows[1]
        vals = [f"{q}", f"DN{dn}", f"{h['v']:.3f}", sci(h["re"]), f"{h['lam']:.5f}", f"{h['hf']:.3f}", f"{h['hm']:.3f}", f"{h['head']:.3f}", f"{h['power']:.2f}"]
        for j, value in enumerate(vals):
            set_cell(row.cells[j], value, 8)

    # Correct the first reference to the government report actually used.
    set_text(
        find_para(doc, "[1] 杭州一达环保技术咨询服务有限公司"),
        "[1] 浙江荣晟环保纸业股份有限公司, 嘉兴市环境科学研究所有限公司. 纸机绿色节能提效升级改造项目环境影响报告书（公示稿）[R]. 嘉兴, 2024.",
    )

    # Keep all formula labels as the school-required continuous “式1…式39”.
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    doc.core_properties.title = "浙江平湖二级油库工艺设计毕业设计初稿（导师审查修订版）"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
