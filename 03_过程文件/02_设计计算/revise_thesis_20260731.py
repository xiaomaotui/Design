from __future__ import annotations

import copy
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree


ROOT = Path(r"D:\毕业论文")
SRC = ROOT / "04_最终成品" / "01_毕业设计说明书" / "张淑鑫_浙江平湖二级油库工艺设计_毕业设计初稿_规范与一致性修订版_2026-07-30.docx"
DST = ROOT / "04_最终成品" / "01_毕业设计说明书" / "张淑鑫_浙江平湖二级油库工艺设计_毕业设计初稿_防火堤与一致性修订版_2026-07-31.docx"
XSL = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")


def clear_paragraph(p: Paragraph) -> None:
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)


SUB_RE = re.compile(r"_\{([^}]+)\}")


def add_rich_text(p: Paragraph, text: str) -> None:
    pos = 0
    for match in SUB_RE.finditer(text):
        if match.start() > pos:
            p.add_run(text[pos:match.start()])
        run = p.add_run(match.group(1))
        run.font.subscript = True
        pos = match.end()
    if pos < len(text):
        p.add_run(text[pos:])


def set_text(p: Paragraph, text: str) -> Paragraph:
    clear_paragraph(p)
    add_rich_text(p, text)
    return p


def find_one(doc: Document, needle: str, *, startswith: bool = False) -> Paragraph:
    found = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if (t.startswith(needle) if startswith else needle in t):
            found.append(p)
    if len(found) != 1:
        raise ValueError(f"Expected one paragraph for {needle!r}, got {len(found)}")
    return found[0]


def insert_after(ref: Paragraph, text: str, style=None) -> Paragraph:
    node = OxmlElement("w:p")
    ref._p.addnext(node)
    p = Paragraph(node, ref._parent)
    p.style = style or ref.style
    add_rich_text(p, text)
    return p


def set_cell(cell, text: str) -> None:
    set_text(cell.paragraphs[0], text)
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    el = tr_pr.find(qn("w:tblHeader"))
    if el is None:
        el = OxmlElement("w:tblHeader")
        tr_pr.append(el)
    el.set(qn("w:val"), "true")


def set_fixed_table_layout(table, widths_cm: list[float], font_pt: float = 8.0) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_cm[idx]
            cell.width = Cm(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(round(width * 567)))
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(font_pt)


def transform_latex(latex: str, transform: etree.XSLT):
    mathml = etree.fromstring(latex_to_mathml(latex).encode("utf-8"))
    return transform(mathml).getroot()


def add_tab(p_el) -> None:
    r = OxmlElement("w:r")
    tab = OxmlElement("w:tab")
    r.append(tab)
    p_el.append(r)


def add_label(p_el, number: int) -> None:
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "宋体")
    fonts.set(qn("w:hAnsi"), "宋体")
    fonts.set(qn("w:eastAsia"), "宋体")
    rpr.append(fonts)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = f"（式{number}）"
    r.append(t)
    p_el.append(r)


def set_formula(p: Paragraph, latex: str, number: int, transform: etree.XSLT) -> None:
    ppr = copy.deepcopy(p._p.pPr) if p._p.pPr is not None else None
    for child in list(p._p):
        p._p.remove(child)
    if ppr is not None:
        p._p.append(ppr)
    add_tab(p._p)
    p._p.append(copy.deepcopy(transform_latex(latex, transform)))
    add_tab(p._p)
    add_label(p._p, number)


def set_display_math(p: Paragraph, latex: str, transform: etree.XSLT) -> None:
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.paragraph_format.left_indent = 0
    p.paragraph_format.right_indent = 0
    p._p.append(copy.deepcopy(transform_latex(latex, transform)))


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    el = settings.find(qn("w:updateFields"))
    if el is None:
        el = OxmlElement("w:updateFields")
        settings.append(el)
    el.set(qn("w:val"), "true")


def replace_in_all(doc: Document, old: str, new: str) -> int:
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            set_text(p, p.text.replace(old, new))
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_text(p, p.text.replace(old, new))
                        count += 1
    return count


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    transform = etree.XSLT(etree.parse(str(XSL)))
    doc = Document(SRC)

    # K值仅保留教材依据、范围和本设计取值。
    set_text(
        find_one(doc, "《油库技术与管理》第2版给出的商业系统参考范围"),
        "《油库技术与管理》第2版规定：二级及以上商业油库年周转系数K取8～14。本设计取K=14。[17]",
    )

    # 名义容量液位的性质与GB 50074第6.1.10条的适用范围。
    p_15000 = find_one(doc, "15 000 m³汽油内浮顶罐：罐截面积")
    set_text(
        p_15000,
        "15 000 m³汽油内浮顶罐：罐截面积907.92 m²，15 min进液折算高度为75/907.92=0.083 m；浮盘底面设计最高位置取17.40 m，安全裕量取0.30 m，设计储存高液位为17.40-0.083-0.30=17.02 m。名义容量对应液位为16.521 m，占罐壁高度的91.78%，该比值仅用于几何核对；最高运行液位按SH/T 3007—2014第4.1.8条确定。高液位报警取16.52 m，高高液位联锁取16.90 m。GB 50074—2014第6.1.10条规定同一罐组的总容量上限，不作为储罐充装高度的判定条款。",
    )

    # 罐壁：计算厚度不计腐蚀裕量，名义厚度比较时再计C1、C2。
    wall_intro = find_one(doc, "储罐按GB 50341—2014采用定设计点法")
    set_text(
        wall_intro,
        "储罐按GB 50341—2014采用定设计点法进行初步罐壁计算。罐壁分为9圈，每圈高2.0 m，材料采用Q345R，腐蚀裕量C₂取2.0 mm；钢板订货采用厚度负偏差不大于0的控制条件，C₁取0 mm。底圈焊接接头系数φ取0.85，其余圈取0.90。设计液位分别按15.915 m和16.521 m计算，充水试验按18.0 m水柱计算。",
    )
    set_formula(doc.paragraphs[145], r"t_d=\frac{4.9D(H-0.3)\rho}{\sigma_d\varphi}", 7, transform)
    set_formula(doc.paragraphs[146], r"t_t=\frac{4.9D(H-0.3)}{\sigma_t\varphi}", 8, transform)
    set_text(find_one(doc, "C1——钢板厚度负偏差"), "C₁——钢板厚度负偏差，取0 mm；")
    set_text(find_one(doc, "C2——腐蚀裕量"), "C₂——腐蚀裕量，取2.0 mm。")
    set_text(
        find_one(doc, "Q345R在相应厚度区间取"),
        "Q345R在相应厚度区间取σ_{d}=218.75 MPa、σ_{t}=230 MPa。式7和式8仅计算承载所需厚度，不计入腐蚀裕量。以20 000 m³汽油罐底圈为例：",
    )
    set_display_math(doc.paragraphs[160], r"t_d=\frac{4.9\times40\times(15.915-0.3)\times0.760}{218.75\times0.85}=12.51\ \mathrm{mm}", transform)
    set_display_math(doc.paragraphs[161], r"t_t=\frac{4.9\times40\times(18.0-0.3)}{230\times0.85}=17.75\ \mathrm{mm}", transform)
    set_text(
        find_one(doc, "按式9比较设计条件、充水试验条件及最小厚度要求"),
        "按式9比较t_{d}+C₁+C₂、t_{t}+C₁及规范最小厚度。底圈设计条件所需名义厚度为12.51+0+2.00=14.51 mm，充水试验条件所需名义厚度为17.75+0=17.75 mm，控制值为17.75 mm，向上圆整取18 mm；其余各圈同理计算，结果列于表2-4。",
    )

    # 呼吸能力表述改为工程设计语言。
    set_text(
        find_one(doc, "呼吸阀和通气孔的设计通气量应同时考虑"),
        "呼吸阀和通气孔的设计通气量应同时考虑最大进油排气量、最大出油吸气量、昼夜温差引起的热呼吸以及火灾工况紧急通气。正常工况最低通气能力分别按最大水运收油流量300 m³/h和最大库内外输流量200 m³/h控制，呼吸阀及通气孔的额定通气量不得小于相应计算值，并应依据制造商性能曲线复核。",
    )

    # 抗风圈：按等效高度自罐顶向下累计到HE/2确定实际标高。
    set_text(
        find_one(doc, "按表2-4各圈壁板厚度逐圈代入式11"),
        "表2-4名义厚度扣除腐蚀裕量后，20 000 m³汽油罐各圈有效厚度自下而上为16、14、12、10、8、8、8、8、8 mm，代入式11得H_{E}=12.72 m；15 000 m³储罐为14、12、10、8、8、8、8、8、8 mm，得H_{E}=14.36 m。中间抗风圈按上下两段等效高度相等布置，目标等效高度分别为H_{E}/2=6.36 m和7.18 m。由罐顶向下逐圈累计折算高度，换算至罐底的实际安装标高分别为18-6.36=11.64 m和18-7.18=10.82 m。再按式12计算无中间抗风圈时的临界外压：",
    )
    set_text(
        find_one(doc, "无中间抗风圈时两类储罐的临界外压均低于"),
        "无中间抗风圈时两类储罐的临界外压均低于相应设计外压。设置1道中间抗风圈并使上下两段折算等效高度均为原等效高度的1/2后，由式12可知临界外压为原值的2倍：",
    )
    set_text(
        find_one(doc, "因此，20 000 m³汽油罐1.854 kPa"),
        "复核结果为：20 000 m³汽油罐1.854 kPa＞1.810 kPa；15 000 m³汽油罐2.096 kPa＞1.810 kPa；15 000 m³柴油罐2.096 kPa＞2.060 kPa，均满足抗风稳定要求。",
    )
    set_text(
        find_one(doc, "说明：抗风圈截面和局部开孔补强"),
        "抗风圈截面、连接焊缝及罐壁开孔处的局部补强应结合盘梯、消防环管、罐顶连接和壁板排板进行校核，抗风圈不得与消防管线和检修通道冲突。",
    )
    wind_table = doc.tables[9]
    set_cell(wind_table.rows[1].cells[5], "+11.64")
    set_cell(wind_table.rows[2].cells[5], "+10.82")
    set_cell(wind_table.rows[3].cells[5], "+10.82")

    # 防火间距：补齐与实际设计设施有关的检查对象。
    set_text(
        find_one(doc, "总平面防火间距按GB 50074—2014"),
        "总平面防火间距按GB 50074—2014表5.1.3及第5.1.8、6.1.15、6.5.2条逐项控制。钢制内浮顶汽油罐和丙A类固定顶柴油罐同组罐间净距按0.4D计算；相邻罐组间按较大罐直径的0.8倍计算；罐组与泵棚、公路装车设施、液体装卸码头、办公及控制建筑、消防泵房、明火地点和围墙的距离按表5.1.3校核。汽油罐组和柴油罐组分别设置防火堤。规范最小值、总平面采用值和校核结果列于表5-2。[3]",
    )
    spacing = doc.tables[18]
    set_cell(spacing.rows[6].cells[3], "G92为23.0 m，G95/GD为19.5 m，满足")
    extra_rows = [
        ["汽油内浮顶罐至易燃液体泵棚", "GB 50074—2014表5.1.3", "15 m", "20 m，满足"],
        ["汽油内浮顶罐至公路装车设施", "GB 50074—2014表5.1.3；装车设油气回收", "15 m", "30 m，满足"],
        ["汽油内浮顶罐至液体装卸码头", "GB 50074—2014表5.1.3", "35 m", "60 m，满足"],
        ["汽油内浮顶罐至办公及中心控制建筑", "GB 50074—2014表5.1.3", "38 m", "50 m，满足"],
        ["汽油内浮顶罐至消防泵房", "GB 50074—2014表5.1.3", "26 m", "50 m，满足"],
        ["汽油内浮顶罐至明火或散发火花地点", "GB 50074—2014表5.1.3", "26 m", "50 m，满足"],
        ["汽油内浮顶罐至库区围墙", "GB 50074—2014表5.1.3", "11 m", "30 m，满足"],
    ]
    for values in extra_rows:
        row = spacing.add_row()
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)
    repeat_header(spacing.rows[0])
    set_text(
        find_one(doc, "表5-2所列采用值作为总平面布置尺寸"),
        "表5-2覆盖本设计存在的储罐、罐组、泵棚、公路装车、码头、办公控制、消防泵房、明火地点和围墙等主要对象。二级石油库罐区和装卸区消防车道按GB 50074—2014第5.2.8、5.2.9条采用6.0 m车道宽度，路面宽度6.0 m，净空高度不小于5.0 m，内缘转弯半径不小于12 m；储罐区形成环形消防车道。总平面图应逐项标注上述净距、道路宽度和转弯半径。",
    )

    # 防火堤：统一几何尺寸，完整考虑储罐排挤、构筑物占用和安全余高。
    set_text(
        find_one(doc, "G92罐组防火堤内尺寸取140 m×90 m"),
        "G92罐组防火堤内尺寸取144 m×92 m，布置2座D=40 m储罐，罐间净距18 m；沿罐组长边方向罐壁至堤内脚为(144-40-18-40)/2=23.0 m。G95和GD罐组防火堤内尺寸均取122 m×82 m，分别布置2座D=34 m储罐，罐间净距15 m，沿长边方向罐壁至堤内脚为(122-34-15-34)/2=19.5 m。三组均大于0.5H=9.0 m。防火堤设计高度取2.2 m，堤内设计液面高度H_{j}=2.0 m，堤顶安全余高0.20 m。",
    )
    set_formula(doc.paragraphs[435], r"V_d=AH_j-(V_1+V_2+V_3+V_4)", 30, transform)
    set_formula(doc.paragraphs[436], r"H_d=H_j+0.20", 31, transform)
    set_formula(
        doc.paragraphs[437],
        r"P_w=\frac{1}{2}\gamma H_j^2,\quad M_o=\frac{P_wH_j}{3},\quad K_s=\frac{\mu\sum W}{P_w},\quad K_o=\frac{\sum M_r}{M_o}",
        32,
        transform,
    )
    # Capture all following paragraphs before inserting additional symbol definitions,
    # so their Python objects remain stable when paragraph indices shift.
    p_g92_intro = find_one(doc, "以G92罐组为例，将L=140 m")
    p_g92_v1 = doc.paragraphs[449]
    p_g92_vd = doc.paragraphs[450]
    p_g92_note = doc.paragraphs[451]
    p_g92_h = doc.paragraphs[452]
    p_g92_result = doc.paragraphs[453]
    p_532_heading = doc.paragraphs[454]
    p_532_intro = doc.paragraphs[455]
    p_532_v1 = doc.paragraphs[456]
    p_532_vd = doc.paragraphs[457]
    p_532_note = doc.paragraphs[458]
    p_532_h = doc.paragraphs[459]
    p_532_result = doc.paragraphs[460]
    p_table53_caption = doc.paragraphs[461]
    structure_heading = find_one(doc, "5.3.3 排水与泄漏围控")
    variable_lines = [
        "式中：",
        "V_{d}——防火堤有效容积，m³；",
        "A——防火堤内水平投影面积，m²；",
        "H_{j}——防火堤内设计液面高度，m；",
        "V₁——堤内储罐在H_{j}以下所占体积，m³；",
        "V₂——隔堤在H_{j}以下所占体积，m³；",
        "V₃——防火堤内坡在H_{j}以下所占体积，m³；",
        "V₄——堤内管墩、集水坑、阀门基础及台阶等构筑物所占体积，m³；",
        "H_{d}——防火堤设计高度，m；",
        "P_{w}——每延米防火堤承受的液体水平合力，kN/m；",
        "γ——泄漏液体重度，kN/m³；",
        "M_{o}——每延米防火堤倾覆力矩，kN·m/m；",
        "K_{s}、K_{o}——抗滑、抗倾覆稳定系数；",
        "μ——基础底面摩擦系数；",
        "ΣW、ΣM_{r}——每延米自重及其抗倾覆力矩。",
    ]
    existing = doc.paragraphs[438:448]
    for p, line in zip(existing, variable_lines[: len(existing)]):
        set_text(p, line)
    cursor = existing[-1]
    for line in variable_lines[len(existing):]:
        cursor = insert_after(cursor, line, existing[-1].style)

    set_text(
        p_g92_intro,
        "G92罐组按最不利单罐20 000 m³泄漏校核。采用直立式钢筋混凝土防火堤，故V₂=V₃=0；管墩、集水坑、阀门基础及台阶等占用体积V₄暂取200 m³。堤内两座储罐在2.0 m液面以下所占体积为：",
    )
    set_display_math(p_g92_v1, r"V_1=2\times\frac{\pi\times40^2}{4}\times2.0=5\,026.55\ \mathrm{m^3}", transform)
    set_display_math(p_g92_vd, r"V_d=144\times92\times2.0-5\,026.55-0-0-200=21\,269.45\ \mathrm{m^3}", transform)
    set_text(
        p_g92_note,
        "有效容积21 269.45 m³＞最大单罐容量20 000 m³。再反算满足容积要求的最低设计液面高度和防火堤高度：",
    )
    set_display_math(p_g92_h, r"H_{j,\min}=\frac{20\,000+200}{144\times92-2\pi\times40^2/4}=1.882\ \mathrm{m},\quad H_{d,\min}=2.082\ \mathrm{m}", transform)
    set_text(
        p_g92_result,
        "采用H_{j}=2.0 m、H_{d}=2.2 m，容积和安全余高均满足要求。",
    )
    set_text(p_532_heading, "5.3.2 95号汽油及柴油罐组防火堤容积")
    set_text(
        p_532_intro,
        "G95和GD罐组的储罐直径、数量及防火堤尺寸相同，均按最大单罐15 000 m³校核。两座D=34 m储罐在2.0 m液面以下所占体积为：",
    )
    set_display_math(p_532_v1, r"V_1=2\times\frac{\pi\times34^2}{4}\times2.0=3\,631.68\ \mathrm{m^3}", transform)
    set_display_math(p_532_vd, r"V_d=122\times82\times2.0-3\,631.68-0-0-200=16\,176.32\ \mathrm{m^3}", transform)
    set_text(p_532_note, "有效容积16 176.32 m³＞最大单罐容量15 000 m³。最低设计液面高度及防火堤高度为：")
    set_display_math(p_532_h, r"H_{j,\min}=\frac{15\,000+200}{122\times82-2\pi\times34^2/4}=1.856\ \mathrm{m},\quad H_{d,\min}=2.056\ \mathrm{m}", transform)
    set_text(p_532_result, "采用H_{j}=2.0 m、H_{d}=2.2 m，G95和GD罐组的有效容积及安全余高均满足要求。")
    set_text(p_table53_caption, "表5-3 防火堤有效容积校核")

    dike_table = doc.tables[19]
    dike_rows = [
        ["罐组", "堤内L×B/m", "H_{j}/m", "储罐占用/m³", "构筑物/m³", "V_{d}/m³", "V_{max}/m³", "结果"],
        ["G92", "144×92", "2.0", "5026.55", "200", "21269.45", "20000", "满足"],
        ["G95", "122×82", "2.0", "3631.68", "200", "16176.32", "15000", "满足"],
        ["GD", "122×82", "2.0", "3631.68", "200", "16176.32", "15000", "满足"],
    ]
    for row_obj, vals in zip(dike_table.rows, dike_rows):
        for cell, val in zip(row_obj.cells, vals):
            set_cell(cell, val)
    repeat_header(dike_table.rows[0])
    set_fixed_table_layout(dike_table, [1.2, 2.3, 1.1, 2.0, 1.8, 2.1, 2.1, 1.1], 8.0)

    # 新增防火堤结构、抗滑、抗倾覆、地基压力和配筋计算。
    set_text(structure_heading, "5.3.3 防火堤结构与稳定性校核")
    original_drainage_body = find_one(doc, "防火堤内地坪坡向集水井")
    normal_style = original_drainage_body.style
    h3_style = structure_heading.style
    cursor = structure_heading
    structural_paragraphs = [
        "防火堤采用C30钢筋混凝土悬臂式结构，墙高2.20 m、墙厚0.30 m，底板宽2.80 m、厚0.45 m。混凝土重度取25 kN/m³；按密度较大的0号柴油校核，液体重度γ=0.84×9.81=8.24 kN/m³。基础底面摩擦系数μ取0.50，地基承载力特征值暂按100 kPa进行初步校核，施工设计时应以岩土勘察报告复核。",
        "将γ=8.24 kN/m³、H_{j}=2.0 m代入式32，液体水平合力P_{w}=0.5×8.24×2.0²=16.48 kN/m，作用点距底板顶面H_{j}/3=0.667 m，倾覆力矩M_{o}=16.48×0.667=10.99 kN·m/m。",
        "墙身自重W₁=0.30×2.20×25=16.50 kN/m，底板自重W₂=2.80×0.45×25=31.50 kN/m，ΣW=48.00 kN/m。按底板中心距堤趾1.40 m计算，抗倾覆力矩ΣM_{r}=48.00×1.40=67.20 kN·m/m。抗滑稳定系数K_{s}=0.50×48.00/16.48=1.46＞1.30；抗倾覆稳定系数K_{o}=67.20/10.99=6.11＞1.50。",
        "基底合力偏心距e=1.40-(67.20-10.99)/48.00=0.229 m，小于B/6=2.80/6=0.467 m。基底最大、最小压力分别为q_{max}=48.00/2.80×(1+6×0.229/2.80)=25.6 kPa、q_{min}=8.7 kPa，均为压应力，且q_{max}＜100 kPa，初步满足地基承载要求。",
        "墙根控制弯矩取10.99 kN·m/m，有效高度h₀取250 mm、钢筋设计强度f_{y}=360 MPa，所需受力钢筋面积A_{s}=10.99×10⁶/(0.9×360×250)=135.7 mm²/m。按最小配筋率0.20%控制，A_{s,min}=0.002×1 000×300=600 mm²/m，液体侧竖向受力钢筋采用Φ12@180，实配628 mm²/m；水平分布钢筋采用Φ10@200。防火堤设置变形缝并采用耐油止水材料，穿堤管道周边采用不燃且耐油的柔性密封。",
        "每个罐组设置不少于2处向外开启的踏步或坡道，相邻踏步或坡道间距不大于60 m。堤内地坪坡向集水坑，正常情况下排水阀保持关闭，排水须经油膜检测和人工确认。防火堤结构设计及构造要求按GB 50351—2014执行。",
    ]
    for text in structural_paragraphs:
        cursor = insert_after(cursor, text, normal_style)
    cursor = insert_after(cursor, "5.3.4 排水与泄漏围控", h3_style)

    # 事故水量统一为同一组分项未圆整结果。
    replace_in_all(doc, "6 985 m³", "7 013 m³")
    replace_in_all(doc, "6 985.25 m³", "7 012.90 m³")
    replace_in_all(doc, "6985.25", "7012.90")
    replace_in_all(doc, "6 985.25", "7 012.90")
    set_cell(doc.tables[44].rows[5].cells[4], "8 000>7 012.90 m³")

    # 去除论文中的写作过程和版本管理措辞，保留必要的工程联动关系。
    meta_heading = find_one(doc, "7.6.4 数据首现、变更和版本一致性")
    set_text(meta_heading, "7.6.4 关键设计参数及联动校核")
    set_text(
        find_one(doc, "年周转量、运输比例、油品密度、周转系数"),
        "年周转量、运输比例、油品密度、周转系数、罐容、罐径、管长、高差、泵性能和消防参数是相互关联的设计参数。储罐尺寸确定后，应同步用于罐壁、抗风、防火堤、冷却水和泡沫系统计算；管线长度及高差确定后，应同步用于水力计算、泵工作点和NPSH校核。",
    )
    set_text(
        find_one(doc, "总平面图确定后，若管长或高差改变管长或高差"),
        "总平面布置变化引起管长或高差改变时，应重新计算表4-2所列阻力、管路特性曲线、泵工作点和NPSH；储罐直径或高度改变时，应重新校核罐壁厚度、抗风圈、防火堤、消防冷却水和泡沫灭火系统。关键参数及其联动校核关系见表7-2。",
    )

    # 单列液位判定说明，避免把几何液位比误解为运行充装率。
    diesel_level = find_one(doc, "15 000 m³柴油固定顶罐：泡沫产生器下沿")
    insert_after(
        diesel_level,
        "液位判定说明：16.521 m由15 000 m³名义容量除以罐截面积换算得到，16.521/18=91.78%仅表示名义容量液位与罐壁几何高度之比，不等同于运行充装率，也不等同于设计储存高液位。储罐运行上限由SH/T 3007—2014规定的设计储存高液位以及高液位、高高液位报警联锁共同控制。本设计15 000 m³汽油罐的设计储存高液位为17.02 m，高液位报警值为16.52 m，高高液位联锁值为16.90 m；GB 50074—2014第6.1.10条控制同一罐组的总容量，不含“罐壁高度90%”的充装限制。因此，91.78%不能作为该储罐违反GB 50074—2014第6.1.10条的依据。",
        diesel_level.style,
    )

    # 修正附录引用，并清理少量不适合正文的写作措辞。
    replace_in_all(doc, "详细路径见附录A", "详细路径见附录C")
    replace_in_all(doc, "本稿", "本设计")
    replace_in_all(doc, "版本一致性", "参数一致性")
    set_update_fields(doc)
    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DST)
    print(DST)


if __name__ == "__main__":
    main()
