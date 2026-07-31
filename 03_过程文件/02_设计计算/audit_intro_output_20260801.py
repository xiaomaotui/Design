from __future__ import annotations

import re
from pathlib import Path

from docx import Document


path = Path(r"D:\毕业论文\04_最终成品\01_毕业设计说明书\张淑鑫_浙江平湖二级油库工艺设计_毕业设计初稿_绪论与研究动态补充版_2026-08-01.docx")
document = Document(path)

headings = []
for index, paragraph in enumerate(document.paragraphs):
    text = paragraph.text.strip()
    if text in {"绪论", "国内外研究动态", "国内研究动态", "国外研究动态", "研究趋势", "1 设计总论", "1.2.3 主要设计数据"}:
        headings.append((index, paragraph.style.name, text))

reference_heading_index = next(i for i, p in enumerate(document.paragraphs) if p.text.strip() == "参考文献")
thanks_heading_index = next(i for i, p in enumerate(document.paragraphs) if p.text.strip() == "致谢")

first_occurrence = {}
all_citations = []
for index, paragraph in enumerate(document.paragraphs[:reference_heading_index]):
    for match in re.finditer(r"\[(\d+)\]", paragraph.text):
        number = int(match.group(1))
        all_citations.append(number)
        first_occurrence.setdefault(number, index)

references = []
for paragraph in document.paragraphs[reference_heading_index + 1 : thanks_heading_index]:
    match = re.match(r"\[(\d+)\]\s+(.+)", paragraph.text.strip())
    if match:
        references.append((int(match.group(1)), match.group(2)))

print("HEADINGS")
for item in headings:
    print(item)
print("FIRST_OCCURRENCE_ORDER", sorted(first_occurrence, key=first_occurrence.get))
print("CITED_UNIQUE", sorted(set(all_citations)))
print("REFERENCE_NUMBERS", [number for number, _ in references])
print("REFERENCE_COUNT", len(references))
print("UNCITED", sorted(set(number for number, _ in references) - set(all_citations)))
print("MISSING_REFERENCE", sorted(set(all_citations) - set(number for number, _ in references)))
print("OLD_REVIEW_PRESENT", any("近五年研究与方案借鉴" in p.text for p in document.paragraphs))
print("CHAPTER_HEADINGS", [p.text.strip() for p in document.paragraphs if re.match(r"^[1-8]\s", p.text.strip()) and p.style.name == "Heading 1"])
