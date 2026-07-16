from pathlib import Path
from copy import deepcopy
import json, shutil
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

ROOT = Path(r"D:\毕业论文")
docx_path = ROOT / r"04_最终成品\01_开题报告\张淑鑫_浙江平湖油库工艺设计_开题报告.docx"
backup_dir = ROOT / r"03_过程文件\01_开题阶段\用户修改稿备份"
backup_dir.mkdir(parents=True, exist_ok=True)
backup = backup_dir / "张淑鑫_浙江平湖油库工艺设计_开题报告_用户删改后备份.docx"
shutil.copy2(docx_path, backup)

doc = Document(docx_path)
cell = doc.tables[0].cell(0, 0)

def set_text_keep_format(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)

def find(text):
    return next(p for p in cell.paragraphs if p.text == text)

def insert_after(anchor, text):
    new_xml = deepcopy(anchor._p)
    anchor._p.addnext(new_xml)
    p = Paragraph(new_xml, anchor._parent)
    set_text_keep_format(p, text)
    return p

# Remove economic-analysis wording throughout the report while keeping the user's remaining prose.
replacements = {
    "安全、连续、灵活且经济的物流系统": "安全、连续、灵活的物流系统",
    "技术经济条件": "技术条件",
    "技术经济比较": "技术适用性比较",
    "经济流速": "设计流速",
    "技术、安全、环保和经济比较": "技术、安全和环保比较",
}
economic_changes = []
for p in cell.paragraphs:
    old = p.text
    new = old
    for a, b in replacements.items():
        new = new.replace(a, b)
    if new != old:
        set_text_keep_format(p, new)
        economic_changes.append([old, new])

# Add anti-wind ring content after the storage tank arrangement item.
if not any("抗风圈" in p.text for p in cell.paragraphs):
    insert_after(find("2.2.2 罐组划分及附件配置"), "2.2.3 储罐抗风圈设计与稳定性校核")

# Refine pump selection into explicit selection and loading/unloading capacity checks.
if not any("泵的装卸能力校核" in p.text for p in cell.paragraphs):
    anchor = find("4.3 输油泵、阀门及计量设备选型")
    for text in [
        "4.3.1 输油泵选型",
        "4.3.2 泵的装卸能力校核",
        "4.3.3 阀门及计量设备选型",
    ]:
        anchor = insert_after(anchor, text)

# Replace the former two broad safety entries with a complete HSE and protection outline.
set_text_keep_format(find("5.2 防火堤、消防道路与应急疏散"), "5.2 HSE管理与危险因素分析")
set_text_keep_format(find("5.3 消防冷却、泡沫灭火与事故水控制"), "5.3 防火、防爆设计")
anchor = find("5.3 防火、防爆设计")
new_safety_items = [
    "5.2.1 HSE设计原则与管理措施",
    "5.2.2 泄漏风险识别与防泄漏措施",
    "5.3.1 防火分区、防火间距与防火堤",
    "5.3.2 爆炸危险区域划分与防爆措施",
    "5.4 防雷电、防静电与防震设计",
    "5.4.1 防雷接地与防静电措施",
    "5.4.2 抗震设防与设备管线防震",
    "5.5 消防系统与事故水控制",
    "5.5.1 消防冷却与泡沫灭火",
    "5.5.2 应急疏散与事故水收集",
]
# Insert 5.2.x before 5.3, then the remaining items after 5.3.
p52 = find("5.2 HSE管理与危险因素分析")
last52 = p52
for text in new_safety_items[:2]:
    last52 = insert_after(last52, text)
anchor = find("5.3 防火、防爆设计")
for text in new_safety_items[2:]:
    anchor = insert_after(anchor, text)

doc.save(docx_path)

check = Document(docx_path)
text = "\n".join(p.text for p in check.tables[0].cell(0, 0).paragraphs)
required = ["抗风圈", "HSE", "防火", "防爆", "防雷电", "防震", "防泄漏", "输油泵选型", "泵的装卸能力校核"]
audit = {
    "required_items_present": {x: x in text for x in required},
    "economic_word_absent": "经济" not in text,
    "backup": str(backup),
    "output": str(docx_path),
    "economic_replacements": len(economic_changes),
}
(ROOT / r"03_过程文件\01_开题阶段\写作提纲补充审计.json").write_text(
    json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
print(json.dumps(audit, ensure_ascii=False))
