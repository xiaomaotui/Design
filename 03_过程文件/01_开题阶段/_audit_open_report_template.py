from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(r"D:\毕业论文")
DOCX = ROOT / "01_原始资料" / "02_学生当前稿件" / "开题报告.docx"
OUT = ROOT / "03_过程文件" / "01_开题阶段"

doc = Document(str(DOCX))

def paragraph_record(paragraph, location):
    return {
        "location": location,
        "style": paragraph.style.name if paragraph.style else "",
        "text": paragraph.text,
        "alignment": int(paragraph.alignment) if paragraph.alignment is not None else None,
        "runs": [
            {
                "text": run.text,
                "font": run.font.name,
                "size_pt": run.font.size.pt if run.font.size else None,
                "bold": run.bold,
                "italic": run.italic,
                "color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
            }
            for run in paragraph.runs
        ],
    }

paragraphs = [paragraph_record(p, f"document/paragraph[{i}]") for i, p in enumerate(doc.paragraphs)]
tables = []
for ti, table in enumerate(doc.tables):
    table_record = {"index": ti, "rows": len(table.rows), "cols": len(table.columns), "cells": []}
    seen = set()
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            key = id(cell._tc)
            if key in seen:
                continue
            seen.add(key)
            table_record["cells"].append(
                {
                    "row": ri,
                    "col": ci,
                    "paragraphs": [
                        paragraph_record(p, f"table[{ti}]/row[{ri}]/cell[{ci}]/paragraph[{pi}]")
                        for pi, p in enumerate(cell.paragraphs)
                    ],
                }
            )
    tables.append(table_record)

parts = []
with zipfile.ZipFile(DOCX) as zf:
    for info in sorted(zf.infolist(), key=lambda x: x.filename):
        data = zf.read(info.filename)
        parts.append(
            {
                "path": info.filename,
                "bytes": len(data),
                "sha256": hashlib.sha256(data).hexdigest(),
            }
        )

payload = {
    "reference": str(DOCX),
    "sha256": hashlib.sha256(DOCX.read_bytes()).hexdigest(),
    "paragraphs": paragraphs,
    "tables": tables,
    "package_parts": parts,
}
(OUT / "开题报告模板_结构证据.json").write_text(
    json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
)
print(json.dumps({"paragraph_count": len(paragraphs), "table_count": len(tables), "tables": [(t['rows'], t['cols']) for t in tables]}, ensure_ascii=False))
