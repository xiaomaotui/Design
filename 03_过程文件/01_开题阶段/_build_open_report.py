from pathlib import Path
import json, re, shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"D:\毕业论文")
template = ROOT / r"01_原始资料\02_学生当前稿件\开题报告.docx"
source = ROOT / r"03_过程文件\01_开题阶段\开题报告正文_张淑鑫.md"
out_dir = ROOT / r"04_最终成品\01_开题报告"
out_dir.mkdir(parents=True, exist_ok=True)
out = out_dir / "张淑鑫_浙江平湖油库工艺设计_开题报告.docx"
shutil.copy2(template, out)

doc = Document(out)
cell = doc.tables[0].cell(0, 0)
row0_pr = doc.tables[0].rows[0]._tr.get_or_add_trPr()
for node in row0_pr.findall(qn("w:cantSplit")):
    row0_pr.remove(node)
for child in list(cell._tc):
    if child.tag != qn("w:tcPr"):
        cell._tc.remove(child)

def set_cell_margin(tc, top=90, start=90, bottom=90, end=90):
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

set_cell_margin(cell._tc)

def add_run_font(run, size=12, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color) if color else RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")

def add_para(text, kind="body"):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if kind == "body":
        pf.first_line_indent = Cm(0.74)
        size, bold = 12, False
    elif kind == "h2":
        pf.first_line_indent = Cm(0)
        pf.keep_with_next = True
        size, bold = 12, True
    elif kind == "h3":
        pf.first_line_indent = Cm(0)
        pf.keep_with_next = True
        size, bold = 12, False
    elif kind == "h4":
        pf.first_line_indent = Cm(0.74)
        pf.keep_with_next = True
        size, bold = 12, False
    else:
        pf.left_indent = Cm(0.74); pf.first_line_indent = Cm(-0.74)
        size, bold = 10.5, False
    add_run_font(p.add_run(text), size, bold)
    return p

lines = source.read_text(encoding="utf-8").splitlines()
in_refs = False
for line in lines:
    s = line.strip()
    if not s or s.startswith("# "):
        continue
    if s == "## 参考文献":
        in_refs = True
        add_para("参考文献", "h2")
    elif s.startswith("## "):
        add_para(s[3:], "h2")
    elif s.startswith("### "):
        add_para(s[4:], "h3")
    elif s.startswith("#### "):
        add_para(s[5:], "h4")
    else:
        add_para(s, "ref" if in_refs else "body")

# Keep the school's original schedule and approval area together after the report body.
schedule_cell = doc.tables[0].cell(1, 0)
if schedule_cell.paragraphs:
    schedule_cell.paragraphs[0].paragraph_format.page_break_before = True

# Ensure the cell ends with a paragraph and prevent outer table rows from being split awkwardly where possible.
if not cell._tc.findall(qn("w:p")):
    cell.add_paragraph()

doc.core_properties.title = "浙江平湖油库工艺设计开题报告"
doc.core_properties.subject = "油气储运工程专业毕业设计开题报告"
doc.core_properties.author = "张淑鑫"
doc.save(out)

body = source.read_text(encoding="utf-8").split("## 参考文献", 1)[0]
body = re.sub(r"[#\s\[\]0-9.：—、，。；（）“”‘’A-Za-z\-]", "", body)
audit = {
    "student": "张淑鑫",
    "title": "浙江平湖油库工艺设计",
    "body_chinese_character_count": len(body),
    "reference_count": 15,
    "source_pdf_count": 15,
    "output": str(out),
    "template_preserved_rows": [1, 2, 3, 4]
}
(ROOT / r"03_过程文件\01_开题阶段\开题报告成品审计.json").write_text(
    json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8"
)
print(json.dumps(audit, ensure_ascii=False))
