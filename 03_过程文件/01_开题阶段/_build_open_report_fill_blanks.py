from pathlib import Path
from copy import deepcopy
import json, re, shutil
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"D:\毕业论文")
template = ROOT / r"01_原始资料\02_学生当前稿件\开题报告.docx"
source = ROOT / r"03_过程文件\01_开题阶段\开题报告正文_张淑鑫.md"
out = ROOT / r"04_最终成品\01_开题报告\张淑鑫_浙江平湖油库工艺设计_开题报告.docx"
out.parent.mkdir(parents=True, exist_ok=True)
shutil.copy2(template, out)

raw_lines = source.read_text(encoding="utf-8").splitlines()

def section(start, end):
    a = next(i for i, x in enumerate(raw_lines) if x.strip() == start) + 1
    b = next(i for i, x in enumerate(raw_lines[a:], a) if x.strip() == end)
    return [x.strip() for x in raw_lines[a:b] if x.strip()]

purpose = [x for x in section("## 1. 研究目的与意义", "## 2. 国内外研究现状") if not x.startswith("#")]
review_raw = section("## 2. 国内外研究现状", "## 3. 主要研究内容与拟解决的问题")
review = []
for x in review_raw:
    if x == "### 2.1 国内研究现状": review.append("（一）国内研究现状")
    elif x == "### 2.2 国外研究现状": review.append("（二）国外研究现状")
    elif not x.startswith("#"): review.append(x)
main_body = [x for x in section("## 3. 主要研究内容与拟解决的问题", "## 4. 研究提纲") if not x.startswith("#")]

outline = [
    "（二）写作提纲",
    "第一章 设计总论",
    "1.1 项目概况与设计范围",
    "1.1.1 项目背景与建设条件",
    "1.1.2 设计任务、范围及边界",
    "1.2 设计依据与基础数据",
    "1.2.1 设计规范与技术原则",
    "1.2.2 油品性质及危险性分析",
    "第二章 库容计算与储罐方案",
    "2.1 库容计算",
    "2.1.1 年周转量与储备天数",
    "2.1.2 设计库容及备用容量",
    "2.2 储罐方案",
    "2.2.1 储罐型式与数量比选",
    "2.2.2 罐组划分及附件配置",
    "第三章 收发油工艺流程设计",
    "3.1 作业工况与物流组织",
    "3.1.1 收油、发油及倒罐工况",
    "3.1.2 同时作业与混油控制",
    "3.2 工艺流程方案",
    "3.2.1 正常收发油流程",
    "3.2.2 计量、清扫及异常流程",
    "第四章 管网水力计算与设备选型",
    "4.1 设计流量、经济流速与管径",
    "4.2 沿程阻力、局部阻力与系统扬程",
    "4.3 输油泵、阀门及计量设备选型",
    "第五章 总图与安全消防设计",
    "5.1 功能分区、罐组布置与安全间距",
    "5.2 防火堤、消防道路与应急疏散",
    "5.3 消防冷却、泡沫灭火与事故水控制",
    "第六章 环境保护与油气回收",
    "6.1 VOCs源项与蒸发损耗分析",
    "6.2 油气密闭收集及回收方案比选",
    "6.3 含油污水、雨污分流与运行监测",
    "第七章 自动控制与运行管理",
    "7.1 液位、压力、流量和可燃气体监测",
    "7.2 联锁切断、报警与控制逻辑",
    "7.3 操作、检维修和应急管理建议",
    "第八章 方案综合评价与结论",
    "8.1 技术、安全、环保和经济比较",
    "8.2 设计成果汇总、结论与建议",
]
main = ["（一）主要研究内容"] + main_body + outline

refs = [x for x in section("## 参考文献", "__END__") if x.startswith("[")] if "__END__" in raw_lines else []
if not refs:
    start = next(i for i, x in enumerate(raw_lines) if x.strip() == "## 参考文献") + 1
    refs = [x.strip() for x in raw_lines[start:] if x.strip().startswith("[")]

doc = Document(out)
table = doc.tables[0]
cell = table.cell(0, 0)
original_texts = [p.text for p in cell.paragraphs]

# Long filled content must be allowed to continue across pages; no template wording is changed.
row0_pr = table.rows[0]._tr.get_or_add_trPr()
for node in row0_pr.findall(qn("w:cantSplit")):
    row0_pr.remove(node)

def set_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"; run.font.size = Pt(size); run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")

def format_fill(p, kind="body"):
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if kind == "body": pf.first_line_indent = Cm(0.74)
    elif kind == "label":
        pf.first_line_indent = Cm(0); pf.keep_with_next = True
    elif kind == "outline":
        pf.first_line_indent = Cm(0); pf.left_indent = Cm(0.74)
    elif kind == "ref":
        pf.left_indent = Cm(0.74); pf.first_line_indent = Cm(-0.74)

def fill_existing(p, text, kind="body"):
    # Only blank template paragraphs are filled.
    assert p.text == ""
    run = p.add_run(text)
    set_font(run, 10.5 if kind == "ref" else 12, kind == "label")
    format_fill(p, kind)
    return p

def insert_after(anchor, text, kind="body"):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    p = Paragraph(new_p, anchor._parent)
    run = p.add_run(text)
    set_font(run, 10.5 if kind == "ref" else 12, kind == "label")
    format_fill(p, kind)
    return p

def add_sequence(anchor, items, first_blank=None, kind_fn=None):
    last = anchor
    for i, text in enumerate(items):
        kind = kind_fn(text) if kind_fn else "body"
        if i == 0 and first_blank is not None:
            last = fill_existing(first_blank, text, kind)
        else:
            last = insert_after(last, text, kind)
    return last

# Stable locators are the exact, preserved template headings and red instruction blocks.
paras = cell.paragraphs
h2 = next(p for p in paras if p.text == "2.本选题国内外研究状况综述")
h3 = next(p for p in paras if p.text == "3.本选题研究的主要内容及写作提纲")
h4 = next(p for p in paras if p.text == "4.主要参考文献")

blank1 = cell.paragraphs[3]
add_sequence(blank1, purpose, first_blank=blank1)

# Insert the review after the template's final review instruction and before heading 3.
review_instruction = next(p for p in cell.paragraphs if p.text.startswith("基于你调查到的国内外"))
add_sequence(review_instruction, review, kind_fn=lambda x: "label" if x.startswith("（") else "body")

# Fill the first available blank after the section-3 instruction, then extend within that slot.
section3_instruction = next(p for p in cell.paragraphs if p.text.startswith("写作提纲就是目录"))
section3_index = next(i for i, p in enumerate(cell.paragraphs) if p.text.startswith("写作提纲就是目录"))
blank3 = next(p for p in cell.paragraphs[section3_index+1:] if p.text == "")
add_sequence(blank3, main, first_blank=blank3,
             kind_fn=lambda x: "label" if x.startswith("（") else ("outline" if re.match(r"^(第|\d+\.)", x) else "body"))

# Fill the template's reference blanks; add more only if the template provides too few.
ref_instruction = next(p for p in cell.paragraphs if p.text == "至少15篇")
ref_index = next(i for i, p in enumerate(cell.paragraphs) if p.text == "至少15篇")
ref_blanks = [p for p in cell.paragraphs[ref_index+1:] if p.text == ""]
last = ref_instruction
for i, text in enumerate(refs):
    if i < len(ref_blanks): last = fill_existing(ref_blanks[i], text, "ref")
    else: last = insert_after(last, text, "ref")
for p in ref_blanks[len(refs):]:
    # These are unused visual filler lines in the blank slot; collapse them so they do not create an empty page.
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)

# Keep item 5 on a fresh page without changing any of its text or tables.
schedule_cell = table.cell(1, 0)
if schedule_cell.paragraphs:
    schedule_cell.paragraphs[0].paragraph_format.page_break_before = True

doc.core_properties.title = "浙江平湖油库工艺设计开题报告"
doc.core_properties.subject = "油气储运工程专业毕业设计开题报告"
doc.core_properties.author = "张淑鑫"
doc.save(out)

# Audit: every original template paragraph must survive in order and item 5 must remain present.
final_doc = Document(out)
final_texts = [p.text for p in final_doc.tables[0].cell(0, 0).paragraphs]
pos = 0
for text in (x for x in original_texts if x != ""):
    while pos < len(final_texts) and final_texts[pos] != text:
        pos += 1
    if pos == len(final_texts):
        raise RuntimeError(f"Template paragraph missing or changed: {text!r}")
    pos += 1

body_text = "".join(purpose + review + main_body)
audit = {
    "template_headings_preserved": all(x in final_texts for x in [
        "1.本选题研究的目的及意义", "2.本选题国内外研究状况综述",
        "3.本选题研究的主要内容及写作提纲", "4.主要参考文献"]),
    "template_instruction_texts_preserved": all(x in final_texts for x in [
        "开题报告字数不少于3500字，艺术类专业不少于2500字", "至少15篇"]),
    "body_chinese_character_count": len(re.sub(r"[^\u4e00-\u9fff]", "", body_text)),
    "reference_count": len(refs),
    "output": str(out)
}
(ROOT / r"03_过程文件\01_开题阶段\开题报告填空版审计.json").write_text(
    json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
print(json.dumps(audit, ensure_ascii=False))
