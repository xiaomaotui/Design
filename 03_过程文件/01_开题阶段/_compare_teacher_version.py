from pathlib import Path
from zipfile import ZipFile
from difflib import SequenceMatcher, unified_diff
from docx import Document
from docx.oxml.ns import qn
import hashlib, json, re

ROOT = Path(r"D:\毕业论文")
original = ROOT / r"04_最终成品\01_开题报告\张淑鑫_浙江平湖油库工艺设计_开题报告.docx"
teacher = ROOT / r"04_最终成品\01_开题报告\老师修改版张淑鑫_浙江平湖油库工艺设计_开题报告-2026.07.16.docx"
out_dir = ROOT / r"03_过程文件\01_开题阶段\老师修改版对比"
out_dir.mkdir(parents=True, exist_ok=True)

def sha(p): return hashlib.sha256(p.read_bytes()).hexdigest()

def iter_block_paragraphs(doc):
    # Paragraphs in document order, including nested tables, without duplicating merged cells.
    seen_cells = set()
    def walk_cell(cell, prefix):
        key = id(cell._tc)
        if key in seen_cells: return
        seen_cells.add(key)
        for i, p in enumerate(cell.paragraphs):
            yield f"{prefix}/p{i}", p
        for ti, table in enumerate(cell.tables):
            for ri, row in enumerate(table.rows):
                for ci, c in enumerate(row.cells):
                    yield from walk_cell(c, f"{prefix}/t{ti}/r{ri}c{ci}")
    for i, p in enumerate(doc.paragraphs):
        yield f"body/p{i}", p
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                yield from walk_cell(cell, f"table{ti}/r{ri}c{ci}")

def fmt_sig(p):
    pf = p.paragraph_format
    runs = []
    for r in p.runs:
        runs.append({
            "text": r.text,
            "font": r.font.name,
            "size": r.font.size.pt if r.font.size else None,
            "bold": r.bold, "italic": r.italic,
            "color": str(r.font.color.rgb) if r.font.color and r.font.color.rgb else None,
        })
    return {
        "style": p.style.name if p.style else None,
        "align": str(p.alignment),
        "first_indent": pf.first_line_indent.pt if pf.first_line_indent else None,
        "left_indent": pf.left_indent.pt if pf.left_indent else None,
        "line_spacing": str(pf.line_spacing),
        "before": pf.space_before.pt if pf.space_before else None,
        "after": pf.space_after.pt if pf.space_after else None,
        "runs": runs,
    }

def package_info(path):
    with ZipFile(path) as z:
        names = set(z.namelist())
        xml = z.read("word/document.xml").decode("utf-8", "ignore")
        return {
            "parts": len(names),
            "has_comments": "word/comments.xml" in names,
            "comment_count": (z.read("word/comments.xml").count(b"<w:comment ") if "word/comments.xml" in names else 0),
            "tracked_insertions": xml.count("<w:ins"),
            "tracked_deletions": xml.count("<w:del"),
            "section_count": xml.count("<w:sectPr"),
            "table_count_xml": xml.count("<w:tbl>"),
        }

def inspect(path):
    d = Document(path)
    recs = [(loc, p.text, fmt_sig(p)) for loc, p in iter_block_paragraphs(d)]
    return d, recs

od, orecs = inspect(original)
td, trecs = inspect(teacher)
otext = [x[1] for x in orecs]
ttext = [x[1] for x in trecs]
matcher = SequenceMatcher(a=otext, b=ttext, autojunk=False)
ops = []
for tag, i1, i2, j1, j2 in matcher.get_opcodes():
    if tag != "equal":
        ops.append({
            "type": tag,
            "original_range": [i1, i2], "teacher_range": [j1, j2],
            "original": [{"loc": orecs[i][0], "text": orecs[i][1]} for i in range(i1, i2)],
            "teacher": [{"loc": trecs[j][0], "text": trecs[j][1]} for j in range(j1, j2)],
        })

# Formatting changes only for paragraphs whose text and relative alignment stayed equal.
format_changes = []
for tag, i1, i2, j1, j2 in matcher.get_opcodes():
    if tag == "equal":
        for oi, tj in zip(range(i1, i2), range(j1, j2)):
            if orecs[oi][2] != trecs[tj][2]:
                format_changes.append({"text": orecs[oi][1], "original_loc": orecs[oi][0], "teacher_loc": trecs[tj][0],
                                       "original_format": orecs[oi][2], "teacher_format": trecs[tj][2]})

report = {
    "original": {"path": str(original), "sha256": sha(original), "bytes": original.stat().st_size,
                 "paragraph_records": len(orecs), "sections": len(od.sections), "tables": len(od.tables), **package_info(original)},
    "teacher": {"path": str(teacher), "sha256": sha(teacher), "bytes": teacher.stat().st_size,
                "paragraph_records": len(trecs), "sections": len(td.sections), "tables": len(td.tables), **package_info(teacher)},
    "content_operations": ops,
    "format_change_count": len(format_changes),
    "format_changes": format_changes,
}
(out_dir / "结构与文字差异.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
(out_dir / "纯文字统一差异.txt").write_text("\n".join(unified_diff(otext, ttext, fromfile="原版", tofile="老师修改版", lineterm="")), encoding="utf-8")

print(json.dumps({
    "original_bytes": report["original"]["bytes"], "teacher_bytes": report["teacher"]["bytes"],
    "content_operation_count": len(ops), "format_change_count": len(format_changes),
    "original_package": package_info(original), "teacher_package": package_info(teacher),
}, ensure_ascii=False))
for n, op in enumerate(ops, 1):
    print(f"\nDIFF {n} {op['type']}")
    for x in op["original"]: print("-", x["text"])
    for x in op["teacher"]: print("+", x["text"])
