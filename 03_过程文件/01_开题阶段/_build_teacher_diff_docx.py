from pathlib import Path
from difflib import SequenceMatcher
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

ROOT = Path(r"D:\毕业论文")
original_path = ROOT / r"04_最终成品\01_开题报告\张淑鑫_浙江平湖油库工艺设计_开题报告.docx"
teacher_path = ROOT / r"04_最终成品\01_开题报告\老师修改版张淑鑫_浙江平湖油库工艺设计_开题报告-2026.07.16.docx"
output = ROOT / r"04_最终成品\01_开题报告\老师修改版_vs_原版_差异对照.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
RED = RGBColor(192, 0, 0)
GRAY = RGBColor(90, 90, 90)
LIGHT_RED = "FCE8E6"
LIGHT_BLUE = "E8F1FA"
LIGHT_GRAY = "F3F5F7"

def set_run_font(run, size=11, bold=False, color=None, strike=False, underline=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.strike = strike
    run.font.underline = underline
    if color: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "微软雅黑")

def set_cell_fill(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for name, val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}"); tcMar.append(node)
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths=(4680,4680), indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None: tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None: tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent)); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc=OxmlElement("w:gridCol"); gc.set(qn("w:w"),str(w)); grid.append(gc)
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        for node in trPr.findall(qn("w:cantSplit")): trPr.remove(node)
        for cell,w in zip(row.cells,widths):
            tcW=cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcW is None: tcW=OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"),str(w)); tcW.set(qn("w:type"),"dxa")
            set_cell_margins(cell)

def repeat_header(row):
    trPr=row._tr.get_or_add_trPr()
    node=trPr.find(qn("w:tblHeader"))
    if node is None:
        node=OxmlElement("w:tblHeader"); trPr.append(node)
    node.set(qn("w:val"),"true")

def clear_para(p):
    for r in list(p.runs): p._p.remove(r._r)

def add_text_with_breaks(p, text, **font_kwargs):
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i: p.add_run().add_break()
        if part: set_run_font(p.add_run(part), **font_kwargs)

def paragraph_texts(path):
    d=Document(path)
    return [p.text.strip() for p in d.tables[0].cell(0,0).paragraphs if p.text.strip()]

old = paragraph_texts(original_path)
new = paragraph_texts(teacher_path)
sm = SequenceMatcher(a=old,b=new,autojunk=False)
changes=[]
for tag,i1,i2,j1,j2 in sm.get_opcodes():
    if tag != "equal":
        changes.append((tag,i1,i2,j1,j2,"\n".join(old[i1:i2]),"\n".join(new[j1:j2])))

def section_at(seq, idx):
    for p in reversed(seq[:max(idx,1)]):
        if re.match(r"^[1-4]\.", p) and len(p) < 30:
            return p
    return "正文"

doc=Document()
sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
sec.top_margin=Inches(0.75); sec.bottom_margin=Inches(0.75)
sec.left_margin=Inches(1); sec.right_margin=Inches(1)
sec.header_distance=Inches(0.3); sec.footer_distance=Inches(0.35)

styles=doc.styles
normal=styles["Normal"]
normal.font.name="Calibri"; normal.font.size=Pt(11)
normal._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [
    ("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,DARK_BLUE,10,5)]:
    st=styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=color
    st._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

# Running header/footer.
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
set_run_font(hp.add_run("浙江平湖油库工艺设计｜开题报告差异对照"),9,color=GRAY)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_run_font(fp.add_run("第 "),9,color=GRAY)
for fld in ["PAGE", "NUMPAGES"]:
    if fld=="NUMPAGES": set_run_font(fp.add_run(" / "),9,color=GRAY)
    begin=OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"),"begin")
    instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text=fld
    sep=OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"),"separate")
    text=OxmlElement("w:t"); text.text="1"
    end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"),"end")
    r=OxmlElement("w:r"); r.extend([begin,instr,sep,text,end]); fp._p.append(r)
set_run_font(fp.add_run(" 页"),9,color=GRAY)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
set_run_font(p.add_run("开题报告差异对照"),23,bold=True,color=RGBColor(0,0,0))
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(14)
set_run_font(p.add_run("老师修改版 vs 原版｜生成日期：2026-07-16"),12,color=GRAY)

legend=doc.add_table(rows=2,cols=2)
set_table_geometry(legend)
repeat_header(legend.rows[0])
for j,(label,fill,color) in enumerate([("原版删除/被替换内容",LIGHT_RED,RED),("老师新增/改写内容",LIGHT_BLUE,BLUE)]):
    set_cell_fill(legend.cell(0,j),fill); set_cell_fill(legend.cell(1,j),fill)
    q=legend.cell(0,j).paragraphs[0]; clear_para(q); set_run_font(q.add_run(label),10,bold=True,color=color)
    q=legend.cell(1,j).paragraphs[0]; clear_para(q)
    if j==0: set_run_font(q.add_run("红色删除线"),10,color=RED,strike=True)
    else: set_run_font(q.add_run("蓝色下划线"),10,color=BLUE,underline=True)

doc.add_heading("一、版式与结构差异", level=1)
for title,body in [
    ("分页修复", "原版当前共62页，正文前存在49页空白；老师版整理为11页，封面、正文、进度表和审核栏连续排布。"),
    ("分节清理", "原版含2个分节并保留末尾空白页；老师版合并为1个分节，删除多余空段和空白页。"),
    ("格式规范化", "老师版为正文中的中英文字符显式指定字体，标题层级和正文视觉效果基本保持一致。"),
    ("修订状态", "两份文件均无批注，也没有可见的Word修订标记；老师的改动已直接写入正文。"),
]:
    p=doc.add_paragraph()
    set_run_font(p.add_run(title+"："),11,bold=True,color=DARK_BLUE)
    set_run_font(p.add_run(body),11)

doc.add_heading("二、逐项文字差异", level=1)
for n,(tag,i1,i2,j1,j2,ot,nt) in enumerate(changes,1):
    context=section_at(new,j1) if nt else section_at(old,i1)
    doc.add_heading(f"变更 {n}｜{context}",level=2)
    table=doc.add_table(rows=2,cols=2)
    set_table_geometry(table)
    repeat_header(table.rows[0])
    for j,(label,fill,color) in enumerate([("原版",LIGHT_RED,RED),("老师修改版",LIGHT_BLUE,BLUE)]):
        set_cell_fill(table.cell(0,j),fill)
        q=table.cell(0,j).paragraphs[0]; clear_para(q); q.alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(q.add_run(label),10,bold=True,color=color)
        table.cell(1,j).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
    po=table.cell(1,0).paragraphs[0]; pn=table.cell(1,1).paragraphs[0]
    clear_para(po); clear_para(pn)
    po.paragraph_format.line_spacing=1.15; pn.paragraph_format.line_spacing=1.15
    char=SequenceMatcher(a=ot,b=nt,autojunk=False)
    for ctag,a1,a2,b1,b2 in char.get_opcodes():
        if ctag=="equal":
            add_text_with_breaks(po,ot[a1:a2],size=10)
            add_text_with_breaks(pn,nt[b1:b2],size=10)
        elif ctag in ("delete","replace"):
            add_text_with_breaks(po,ot[a1:a2],size=10,color=RED,strike=True)
        if ctag in ("insert","replace"):
            add_text_with_breaks(pn,nt[b1:b2],size=10,color=BLUE,underline=True)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

doc.add_heading("三、未发生实质变化的部分",level=1)
p=doc.add_paragraph()
set_run_font(p.add_run("封面基本信息、15篇参考文献、毕业设计进度安排、指导教师意见及学院审核栏未发现实质文字改动。"),11)

settings=doc.settings._element
update=settings.find(qn("w:updateFields"))
if update is None: update=OxmlElement("w:updateFields"); settings.append(update)
update.set(qn("w:val"),"true")
doc.core_properties.title="开题报告差异对照：老师修改版 vs 原版"
doc.core_properties.author="Codex"
doc.save(output)
print(output)
print("meaningful_changes",len(changes))
